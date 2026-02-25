"""Word (.docx) report generator for speed analysis."""

import base64
import io

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from .config import RoadConfig


def _b64_to_stream(b64_str):
    return io.BytesIO(base64.b64decode(b64_str))


def _add_table_to_doc(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(8)
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = str(row_data[col])
            for par in row_cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.size = Pt(8)


def _add_method_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)


def _add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def generate_word_report(charts, tables, static_map_paths, config: RoadConfig):
    """Generate a Word (.docx) report."""
    doc = Document()
    road = config.road_name
    limit = config.speed_limit
    date_desc = config.date_description or "N/D"
    primary_dt = config.day_types[0] if config.day_types else "Feriali"

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    img_width = Inches(6.5)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Analisi della Distribuzione delle Velocit\u00e0")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run(road)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x28, 0x35, 0x93)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Dati TomTom Speed Profiles \u2022 Periodo: {date_desc}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # Table of contents
    doc.add_heading("Indice", level=1)
    toc_items = [
        "1. Introduzione e Metodologia",
        "2. Panoramica dei Dati",
        "3. Profili Temporali di Velocit\u00e0",
        "4. Superamento del Limite di Velocit\u00e0",
        "5. Analisi V85 (85\u00b0 Percentile)",
        "6. Variabilit\u00e0 delle Velocit\u00e0",
        "7. Analisi delle Velocit\u00e0 Notturne",
        "8. Confronto tra Periodi",
        "9. Mappe Statiche delle Velocit\u00e0",
        "10. Conclusioni e Raccomandazioni",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading("1. Introduzione e Metodologia", level=1)
    doc.add_paragraph(
        f"Il presente report analizza la distribuzione delle velocit\u00e0 veicolari lungo "
        f"{road}, utilizzando i dati TomTom Speed Profiles "
        f"relativi al periodo {date_desc}."
    )
    doc.add_paragraph("L\u2019analisi si basa sui seguenti dataset:")
    for dt in config.day_types:
        dirs_str = ", ".join(f"Direzione {d}" for d in config.directions)
        doc.add_paragraph(f"{dt} \u2014 {dirs_str}", style="List Bullet")

    doc.add_paragraph("Metriche principali:")
    for name, desc in [
        ("Velocit\u00e0 media armonica", "calcolata da TomTom come media armonica "
         "ponderata sulla lunghezza di tutti i segmenti."),
        ("V85 (85\u00b0 percentile)", "velocit\u00e0 non superata dall\u201985% "
         "dei veicoli. Indice 16 dell\u2019array di 19 percentili TomTom."),
        ("Deviazione standard", "dispersione delle velocit\u00e0 individuali."),
        ("Tasso di superamento", "percentuale dei km di percorso "
         "con velocit\u00e0 superiore al limite."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(f"Limite di velocit\u00e0: {limit} km/h.")

    # 2. Overview
    doc.add_heading("2. Panoramica dei Dati", level=1)
    _add_method_box(doc,
        "Nota metodologica: le velocit\u00e0 sono le medie armoniche a livello "
        "di percorso calcolate da TomTom.")
    _add_table_to_doc(doc, tables["overview"])
    doc.add_paragraph()

    doc.add_heading("Riepilogo Superamento Limiti (pesato per km)", level=2)
    _add_table_to_doc(doc, tables["exceedance"])

    # 3. Temporal profiles
    doc.add_page_break()
    doc.add_heading("3. Profili Temporali di Velocit\u00e0", level=1)

    doc.add_heading("Velocit\u00e0 Media Armonica per Ora", level=2)
    doc.add_picture(_b64_to_stream(charts["speed_by_hour"]), width=img_width)
    _add_caption(doc, "Velocit\u00e0 media armonica dell\u2019intero itinerario per ora.")

    doc.add_heading("V85 Medio per Ora", level=2)
    doc.add_picture(_b64_to_stream(charts["v85_by_hour"]), width=img_width)
    _add_caption(doc, "85\u00b0 percentile mediato su tutti i segmenti per ora.")

    doc.add_heading("Mappe di Calore: Velocit\u00e0 per Progressiva e Ora", level=2)
    for label, b64 in charts.get("heatmaps", {}).items():
        doc.add_heading(label, level=3)
        doc.add_picture(_b64_to_stream(b64), width=img_width)

    # 4. Exceedance
    doc.add_page_break()
    doc.add_heading("4. Superamento del Limite di Velocit\u00e0", level=1)
    doc.add_paragraph(
        f"Per ogni ora si calcola la percentuale della lunghezza del percorso in cui "
        f"la velocit\u00e0 supera {limit} km/h.")
    doc.add_picture(_b64_to_stream(charts["exceedance"]), width=img_width)

    # 5. V85
    doc.add_page_break()
    doc.add_heading("5. Analisi V85 (85\u00b0 Percentile)", level=1)
    _add_method_box(doc,
        "Il V85 \u00e8 la velocit\u00e0 al di sotto della quale viaggia l\u201985% dei veicoli. "
        "Indice 16 dell\u2019array di 19 percentili TomTom.")

    for direction, b64 in charts.get("v85_spatial", {}).items():
        doc.add_heading(f"Profilo Spaziale V85 \u2014 Dir. {direction}", level=2)
        doc.add_picture(_b64_to_stream(b64), width=img_width)

    for direction in config.directions:
        dir_slug = direction.lower().replace(" ", "_")
        table_key = f"top_fast_{dir_slug}"
        if table_key in tables:
            doc.add_heading(
                f"Top 10 Tratti con V85 pi\u00f9 Elevato \u2014 Dir. {direction} ({primary_dt})",
                level=2)
            _add_table_to_doc(doc, tables[table_key])
            doc.add_paragraph()

    # 6. Variability
    doc.add_page_break()
    doc.add_heading("6. Variabilit\u00e0 delle Velocit\u00e0", level=1)
    doc.add_picture(_b64_to_stream(charts["variability"]), width=img_width)

    # 7. Night
    doc.add_page_break()
    doc.add_heading("7. Analisi delle Velocit\u00e0 Notturne", level=1)
    doc.add_picture(_b64_to_stream(charts["night"]), width=img_width)

    # 8. Comparison
    if charts.get("weekday_weekend"):
        doc.add_page_break()
        dt_label = " vs ".join(config.day_types[:2])
        doc.add_heading(f"8. Confronto {dt_label}", level=1)
        doc.add_picture(_b64_to_stream(charts["weekday_weekend"]), width=img_width)

    # 9. Static maps
    doc.add_page_break()
    doc.add_heading("9. Mappe Statiche delle Velocit\u00e0", level=1)
    doc.add_paragraph(
        f"Le mappe mostrano la distribuzione spaziale delle velocit\u00e0 lungo {road}. "
        f"Il colore indica la velocit\u00e0: verde = ridotta, giallo = limite, rosso = elevata."
    )

    from pathlib import Path
    static_dir = config.output_dir / "static_maps"
    if static_dir.exists():
        for p in sorted(static_dir.glob("*.png")):
            doc.add_heading(p.stem.replace("_", " ").title(), level=3)
            doc.add_picture(str(p), width=img_width)

    # 10. Conclusions
    doc.add_page_break()
    doc.add_heading("10. Conclusioni e Raccomandazioni", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Risultati Principali:")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    for item in [
        f"L\u2019analisi copre l\u2019intero corridoio di {road}",
        f"Il limite di {limit} km/h viene superato in porzioni significative del percorso",
        "Il V85 supera il limite in molte fasce orarie",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Report generato automaticamente dai dati TomTom Speed Profiles. "
        "Si raccomanda un\u2019interpretazione contestualizzata da parte di tecnici qualificati."
    )
    run.italic = True
    run.font.size = Pt(9)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"TomTom Speed Profiles \u2014 {road} \u2014 {date_desc}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Save
    road_slug = road.lower().replace(" ", "_")[:30]
    docx_path = config.output_dir / f"report_{road_slug}.docx"
    doc.save(str(docx_path))
    return str(docx_path)
