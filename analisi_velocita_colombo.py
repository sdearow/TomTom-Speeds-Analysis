#!/usr/bin/env python3
"""
Analisi della Distribuzione delle Velocità — Via Cristoforo Colombo, Roma
Dati TomTom Speed Profiles
"""

# ================================================================
# SECTION 1 — IMPORTS & CONFIGURATION
# ================================================================

import json
import base64
import io
import warnings
from pathlib import Path

import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import folium
import branca.colormap as cm
import contextily as cx
from shapely.geometry import shape
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

warnings.filterwarnings("ignore")

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "Via Cristoforo Colombo"
OUTPUT_DIR = BASE_DIR / "output_colombo"
MAPS_DIR = OUTPUT_DIR / "maps"
STATIC_MAPS_DIR = OUTPUT_DIR / "static_maps"

# ---------- Date-range ID → day_type name ----------
DATE_RANGE_MAP = {
    1: "Inv. Feriale",
    2: "Est. Feriale",
    3: "Est. Festivo",
    4: "Inv. Festiva",
}
# Short labels for seasons
SEASON_FERIALE = ["Inv. Feriale", "Est. Feriale"]
SEASON_FESTIVO = ["Est. Festivo", "Inv. Festiva"]
ALL_DAY_TYPES = list(DATE_RANGE_MAP.values())
TIMESET_OFFSET = 5  # timeSet ID 5 → hour 0

# ---------- File definitions ----------
# Each entry: (carriage, direction, file_path)
# carriage: "Centrale", "Lat. Ostia", "Lat. Centro"
# For Centrale: direction = "Ostia" or "Centro"
# For Laterals: direction = tratto name (will be concatenated by direction)
FILE_DEFS = [
    ("Centrale", "Ostia",  DATA_DIR / "Centrale_Dir. Ostia_1.geojson"),
    ("Centrale", "Centro", DATA_DIR / "Centrale_Dir.Centro_2.geojson"),
    ("Lat. Ostia", "Pontina-Mezzocammino",
     DATA_DIR / "Laterale_Dir. Ostia_Tratto Pontina-Mezzocammino_3.geojson"),
    ("Lat. Ostia", "Via di Acilia-Piazzale",
     DATA_DIR / "Laterale_Dir. Ostia_Tratto Via di Acilia-Piazzale_4.geojson"),
    ("Lat. Centro", "Piazzale-Acilia",
     DATA_DIR / "Laterale_Dir. Centro_Tratto Piazzale-Acilia_5.geojson"),
    ("Lat. Centro", "Svincolo Malafede",
     DATA_DIR / "Laterale_Dir. Centro_Svincolo Via di Malafede_6.geojson"),
    ("Lat. Centro", "Alfonsine-Brandeliero",
     DATA_DIR / "Laterale_Dir. Centro_Tratto Alfonsine-Brandeliero_7.geojson"),
]

# Tratto ordering for concatenation (geographic order along the route)
LAT_OSTIA_TRATTI = ["Pontina-Mezzocammino", "Via di Acilia-Piazzale"]
LAT_CENTRO_TRATTI = ["Piazzale-Acilia", "Svincolo Malafede", "Alfonsine-Brandeliero"]

# ---------- Carriageway groups ----------
CARRIAGE_GROUPS = {
    "Centrale": {
        "label": "Carreggiata Centrale",
        "directions": ["Ostia", "Centro"],
        "dir_labels": {"Ostia": "Dir. Ostia", "Centro": "Dir. Centro"},
    },
    "Lat. Ostia": {
        "label": "Carreggiata Laterale Dir. Ostia",
        "directions": LAT_OSTIA_TRATTI,
        "dir_labels": {t: t for t in LAT_OSTIA_TRATTI},
        "combined_label": "Laterale Dir. Ostia (intero percorso)",
    },
    "Lat. Centro": {
        "label": "Carreggiata Laterale Dir. Centro",
        "directions": LAT_CENTRO_TRATTI,
        "dir_labels": {t: t for t in LAT_CENTRO_TRATTI},
        "combined_label": "Laterale Dir. Centro (intero percorso)",
    },
}

# ---------- Time periods ----------
NIGHT_HOURS = list(range(0, 6)) + [22, 23]
AM_PEAK = [7, 8]
PM_PEAK = [17, 18]
MIDDAY = [12, 13]
P85_IDX = 16
P95_IDX = 18

# ---------- Plotting ----------
plt.rcParams.update({
    "figure.dpi": 140, "savefig.dpi": 140,
    "font.family": "sans-serif", "font.size": 10,
    "axes.titlesize": 12, "axes.labelsize": 10,
})
HOUR_SHORT = [f"{h}" for h in range(24)]

# Colours per carriage
CARRIAGE_COLORS = {
    "Centrale": {"Ostia": "#1f77b4", "Centro": "#ff7f0e"},
    "Lat. Ostia": {
        "Pontina-Mezzocammino": "#2ca02c",
        "Via di Acilia-Piazzale": "#d62728",
    },
    "Lat. Centro": {
        "Piazzale-Acilia": "#9467bd",
        "Svincolo Malafede": "#8c564b",
        "Alfonsine-Brandeliero": "#e377c2",
    },
}
# Colours for day_type overlays
DAY_TYPE_COLORS = {
    "Inv. Feriale": "#1f77b4",
    "Est. Feriale": "#ff7f0e",
    "Est. Festivo": "#2ca02c",
    "Inv. Festiva": "#d62728",
}
DAY_TYPE_LABELS = {
    "Inv. Feriale": "Invernale Feriale",
    "Est. Feriale": "Estiva Feriale",
    "Est. Festivo": "Estiva Festivo",
    "Inv. Festiva": "Invernale Festiva",
}


# ================================================================
# SECTION 2 — DATA LOADING
# ================================================================

def load_geojson(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def parse_segments(data, carriage, direction):
    """Parse a single GeoJSON file → (segments_df, summaries_df, header).

    Unlike Corso Francia (which has 1 dateRange per file), Via Colombo files
    contain 4 dateRanges.  Each segment has 24 hours × 4 dateRanges = 96 rows.
    """
    features = data["features"]
    header = features[0]["properties"]
    segments = features[1:]

    rows = []
    cumul = 0.0
    for idx, feat in enumerate(segments):
        props = feat["properties"]
        geom = shape(feat["geometry"])
        seg_dist = props["distance"]
        for tr in props["segmentTimeResults"]:
            hour = tr["timeSet"] - TIMESET_OFFSET
            day_type = DATE_RANGE_MAP.get(tr["dateRange"], f"DR{tr['dateRange']}")
            sp = tr["speedPercentiles"]
            rows.append({
                "carriage": carriage, "direction": direction,
                "day_type": day_type, "seg_idx": idx,
                "segmentId": props["segmentId"],
                "streetName": props.get("streetName", "Via Cristoforo Colombo"),
                "frc": props["frc"], "speedLimit": props["speedLimit"],
                "seg_distance": seg_dist,
                "cum_dist_start": cumul, "cum_dist_mid": cumul + seg_dist / 2,
                "cum_dist_end": cumul + seg_dist,
                "hour": hour,
                "harm_avg_speed": tr["harmonicAverageSpeed"],
                "avg_speed": tr["averageSpeed"],
                "median_speed": tr["medianSpeed"],
                "std_speed": tr.get("standardDeviationSpeed", np.nan),
                "avg_tt": tr["averageTravelTime"],
                "median_tt": tr["medianTravelTime"],
                "std_tt": tr.get("travelTimeStandardDeviation", np.nan),
                "tt_ratio": tr["travelTimeRatio"],
                "sample_size": tr["sampleSize"],
                "norm_sample": tr["normalizedSampleSize"],
                "p5": sp[0], "p15": sp[2], "p25": sp[4], "p50": sp[9],
                "p75": sp[14], "p85": sp[16], "p90": sp[17], "p95": sp[18],
                "geometry": geom,
            })
        cumul += seg_dist

    # Route-level summaries (per dateRange × hour)
    sum_rows = []
    for s in header.get("summaries", []):
        hour = s["timeSet"] - TIMESET_OFFSET
        day_type = DATE_RANGE_MAP.get(s["dateRange"], f"DR{s['dateRange']}")
        ssp = s.get("speedPercentiles", [])
        sum_rows.append({
            "carriage": carriage, "direction": direction,
            "day_type": day_type, "hour": hour,
            "route_dist": s["distance"],
            "harm_avg_speed": s["harmonicAverageSpeed"],
            "avg_tt": s["averageTravelTime"],
            "median_tt": s["medianTravelTime"],
            "pti": s["planningTimeIndex"],
            "route_p85": ssp[P85_IDX] if len(ssp) > P85_IDX else np.nan,
            "route_p95": ssp[P95_IDX] if len(ssp) > P95_IDX else np.nan,
        })
    return pd.DataFrame(rows), pd.DataFrame(sum_rows), header


def load_all_data():
    """Load all 7 files, assign carriage groups.

    For lateral carriages, tratti within each direction are concatenated
    so that progressive distance is continuous across tratti.
    """
    all_seg, all_sum = [], []
    headers = {}

    for carriage, direction, path in FILE_DEFS:
        data = load_geojson(path)
        seg_df, sum_df, hdr = parse_segments(data, carriage, direction)
        all_seg.append(seg_df)
        all_sum.append(sum_df)
        headers[(carriage, direction)] = hdr

    segments = pd.concat(all_seg, ignore_index=True)
    summaries = pd.concat(all_sum, ignore_index=True)

    # --- Fix cumulative distances for concatenated lateral tratti ---
    # For each lateral direction, the tratti are sequential.  We need to
    # offset the cum_dist of the second (and third) tratto so progressive
    # distance is continuous.
    for carriage, tratti_list in [("Lat. Ostia", LAT_OSTIA_TRATTI),
                                   ("Lat. Centro", LAT_CENTRO_TRATTI)]:
        cumul_offset = 0.0
        for tratto in tratti_list:
            mask = ((segments["carriage"] == carriage)
                    & (segments["direction"] == tratto))
            if mask.sum() == 0:
                continue
            segments.loc[mask, "cum_dist_start"] += cumul_offset
            segments.loc[mask, "cum_dist_mid"] += cumul_offset
            segments.loc[mask, "cum_dist_end"] += cumul_offset
            # Update offset for next tratto = max end of current tratto
            max_end = segments.loc[mask, "cum_dist_end"].max()
            cumul_offset = max_end

    # Re-index seg_idx to be unique within each (carriage, direction, day_type)
    # For laterals, make seg_idx continuous across tratti
    for carriage, tratti_list in [("Lat. Ostia", LAT_OSTIA_TRATTI),
                                   ("Lat. Centro", LAT_CENTRO_TRATTI)]:
        idx_offset = 0
        for tratto in tratti_list:
            mask = ((segments["carriage"] == carriage)
                    & (segments["direction"] == tratto))
            if mask.sum() == 0:
                continue
            current_max = segments.loc[mask, "seg_idx"].max()
            segments.loc[mask, "seg_idx"] += idx_offset
            idx_offset += current_max + 1

    return segments, summaries, headers


# ================================================================
# SECTION 3 — ANALYSIS FUNCTIONS
# ================================================================

def get_speed_limit_for_group(segments, carriage):
    """Return the predominant speed limit (by km) for a carriage group."""
    sub = segments[segments["carriage"] == carriage].drop_duplicates("seg_idx")
    if sub.empty:
        return 50
    grp = sub.groupby("speedLimit")["seg_distance"].sum()
    return int(grp.idxmax())


def compute_exceedance_by_km(df, speed_col, limit=None):
    """% of ROUTE LENGTH (km) where speed_col > limit, weighted by seg_distance.

    If limit is None, use each segment's own speedLimit.
    """
    df = df.copy()
    if limit is not None:
        df["exceed_dist"] = np.where(df[speed_col] > limit, df["seg_distance"], 0.0)
    else:
        df["exceed_dist"] = np.where(df[speed_col] > df["speedLimit"],
                                      df["seg_distance"], 0.0)
    grp = df.groupby(["carriage", "direction", "day_type", "hour"]).agg(
        exceed_m=("exceed_dist", "sum"), total_m=("seg_distance", "sum"),
    ).reset_index()
    grp["pct_km_exceed"] = grp["exceed_m"] / grp["total_m"] * 100
    return grp


def hourly_means(df, value_col):
    return df.groupby(["carriage", "direction", "day_type", "hour"]
                      )[value_col].mean().reset_index()


def segment_peak_stats(df):
    """Compute per-segment peak statistics across all hours."""
    def _agg(sub):
        return pd.Series({
            "night_avg_speed": sub.loc[sub["hour"].isin(NIGHT_HOURS), "avg_speed"].mean(),
            "night_v85":       sub.loc[sub["hour"].isin(NIGHT_HOURS), "p85"].mean(),
            "am_avg_speed":    sub.loc[sub["hour"].isin(AM_PEAK), "avg_speed"].mean(),
            "am_v85":          sub.loc[sub["hour"].isin(AM_PEAK), "p85"].mean(),
            "pm_avg_speed":    sub.loc[sub["hour"].isin(PM_PEAK), "avg_speed"].mean(),
            "pm_v85":          sub.loc[sub["hour"].isin(PM_PEAK), "p85"].mean(),
            "all_avg_speed":   sub["avg_speed"].mean(),
            "all_v85":         sub["p85"].mean(),
            "all_std":         sub["std_speed"].mean(),
            "max_v85":         sub["p85"].max(),
            "max_avg_speed":   sub["avg_speed"].max(),
            "cum_dist_mid":    sub["cum_dist_mid"].iloc[0],
            "cum_dist_start":  sub["cum_dist_start"].iloc[0],
            "seg_distance":    sub["seg_distance"].iloc[0],
            "streetName":      sub["streetName"].iloc[0],
            "speedLimit":      sub["speedLimit"].iloc[0],
        })
    return df.groupby(["carriage", "direction", "day_type", "seg_idx"]).apply(
        _agg, include_groups=False
    ).reset_index()


# ================================================================
# SECTION 4 — CHART GENERATION
# ================================================================

def fig_to_base64(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


def chart_speed_by_hour(summaries, carriage):
    """Speed by hour — 4 panels (one per day_type), lines per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    colors = CARRIAGE_COLORS[carriage]
    fig, axes = plt.subplots(1, len(ALL_DAY_TYPES), figsize=(5 * len(ALL_DAY_TYPES), 5),
                             sharey=True)
    if len(ALL_DAY_TYPES) == 1:
        axes = [axes]
    for ax, dt in zip(axes, ALL_DAY_TYPES):
        for d in dirs:
            sub = summaries[(summaries["carriage"] == carriage)
                            & (summaries["direction"] == d)
                            & (summaries["day_type"] == dt)].sort_values("hour")
            if sub.empty:
                continue
            ax.plot(sub["hour"], sub["harm_avg_speed"], color=colors[d],
                    label=cg["dir_labels"][d], linewidth=2, marker="o", markersize=3)
        speed_limit = 50  # reference line
        ax.axhline(speed_limit, color="red", ls="--", lw=1, alpha=.7)
        ax.set_xlabel("Ora")
        ax.set_title(f"{DAY_TYPE_LABELS[dt]}", fontsize=10)
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=6)
        ax.legend(fontsize=7); ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
    axes[0].set_ylabel("Velocità media armonica (km/h)")
    fig.suptitle(f"Velocità Media per Ora — {cg['label']}", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_v85_by_hour(segments, carriage):
    """V85 by hour — 4 panels (one per day_type), lines per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    colors = CARRIAGE_COLORS[carriage]
    hmeans = hourly_means(segments[segments["carriage"] == carriage], "p85")
    fig, axes = plt.subplots(1, len(ALL_DAY_TYPES), figsize=(5 * len(ALL_DAY_TYPES), 5),
                             sharey=True)
    if len(ALL_DAY_TYPES) == 1:
        axes = [axes]
    for ax, dt in zip(axes, ALL_DAY_TYPES):
        for d in dirs:
            sub = hmeans[(hmeans["direction"] == d)
                         & (hmeans["day_type"] == dt)].sort_values("hour")
            if sub.empty:
                continue
            ax.plot(sub["hour"], sub["p85"], color=colors[d],
                    label=cg["dir_labels"][d], linewidth=2, marker="o", markersize=3)
        ax.axhline(50, color="red", ls="--", lw=1, alpha=.7)
        ax.set_xlabel("Ora")
        ax.set_title(f"{DAY_TYPE_LABELS[dt]}", fontsize=10)
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=6)
        ax.legend(fontsize=7); ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
    axes[0].set_ylabel("V85 medio dei segmenti (km/h)")
    fig.suptitle(f"V85 per Ora — {cg['label']}", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def _make_heatmap_progressive(sub, value_col, title, cmap, norm, fig_ax=None):
    """Heatmap with progressive km on x-axis using pcolormesh."""
    seg_first = sub.drop_duplicates("seg_idx").sort_values("seg_idx")
    x_edges = np.concatenate([
        seg_first["cum_dist_start"].values,
        [seg_first["cum_dist_end"].values[-1]]
    ]) / 1000.0
    y_edges = np.arange(-0.5, 24.5, 1.0)

    pivot = sub.pivot_table(index="hour", columns="seg_idx",
                             values=value_col, aggfunc="mean")
    pivot = pivot.reindex(index=range(24), columns=seg_first["seg_idx"].values)

    if fig_ax is None:
        fig, ax = plt.subplots(figsize=(12, 7))
    else:
        fig, ax = fig_ax

    mesh = ax.pcolormesh(x_edges, y_edges, pivot.values,
                          cmap=cmap, norm=norm, shading="flat")
    ax.invert_yaxis()
    ax.set_xlabel("Progressiva (km)")
    ax.set_ylabel("Ora")
    ax.set_yticks(range(24)); ax.set_yticklabels(HOUR_SHORT, fontsize=7)
    max_km = x_edges[-1]
    tick_iv = 0.25 if max_km < 2.0 else (0.5 if max_km < 5.0 else 1.0)
    xticks = np.arange(0, max_km + tick_iv, tick_iv)
    ax.set_xticks(xticks)
    ax.set_xticklabels([f"{x:.2f}" for x in xticks], fontsize=7)
    ax.set_title(title, fontsize=10)
    fig.colorbar(mesh, ax=ax, shrink=0.75, label="km/h")
    return fig


def chart_heatmaps(segments, carriage, day_type="Inv. Feriale"):
    """Speed heatmaps: hour (y) x progressive distance (x), one per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    results = {}
    norm = mcolors.TwoSlopeNorm(vmin=15, vcenter=50, vmax=90)
    cmap = plt.cm.RdYlGn
    for d in cg["directions"]:
        sub = segments[(segments["carriage"] == carriage)
                       & (segments["direction"] == d)
                       & (segments["day_type"] == day_type)]
        if sub.empty:
            continue
        label = cg["dir_labels"][d]
        title = f"Velocità Media — {DAY_TYPE_LABELS[day_type]} {label}"
        fig = _make_heatmap_progressive(sub, "avg_speed", title, cmap, norm)
        results[label] = fig_to_base64(fig)
    return results


def chart_exceedance(segments, carriage, day_type="Inv. Feriale"):
    """% of route-km exceeding each segment's own speed limit, by hour."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    colors = CARRIAGE_COLORS[carriage]
    sub_all = segments[(segments["carriage"] == carriage)
                       & (segments["day_type"] == day_type)]
    exc_avg = compute_exceedance_by_km(sub_all, "avg_speed")
    exc_v85 = compute_exceedance_by_km(sub_all, "p85")

    fig, axes = plt.subplots(1, 2, figsize=(14, 5), sharey=True)
    for d in dirs:
        sub = exc_avg[exc_avg["direction"] == d].sort_values("hour")
        if sub.empty:
            continue
        axes[0].plot(sub["hour"], sub["pct_km_exceed"], color=colors[d],
                     label=cg["dir_labels"][d], linewidth=2, marker="o", markersize=3)
    axes[0].set_title("Km con Vel. Media > Limite")
    axes[0].set_xlabel("Ora")
    axes[0].set_ylabel("% del percorso (km) oltre il limite")
    axes[0].set_xticks(range(24)); axes[0].set_xticklabels(HOUR_SHORT, fontsize=7)
    axes[0].legend(fontsize=7); axes[0].grid(True, alpha=.3)
    axes[0].set_xlim(-.5, 23.5); axes[0].set_ylim(0, 105)

    for d in dirs:
        sub = exc_v85[exc_v85["direction"] == d].sort_values("hour")
        if sub.empty:
            continue
        axes[1].plot(sub["hour"], sub["pct_km_exceed"], color=colors[d],
                     label=cg["dir_labels"][d], linewidth=2, marker="o", markersize=3)
    axes[1].set_title("Km con V85 > Limite")
    axes[1].set_xlabel("Ora")
    axes[1].set_xticks(range(24)); axes[1].set_xticklabels(HOUR_SHORT, fontsize=7)
    axes[1].legend(fontsize=7); axes[1].grid(True, alpha=.3)
    axes[1].set_xlim(-.5, 23.5)

    title = f"Superamento Limite (pesato per km) — {cg['label']} {DAY_TYPE_LABELS[day_type]}"
    fig.suptitle(title, fontsize=12, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_v85_spatial(segments, carriage, day_type="Inv. Feriale"):
    """V85 along corridor at key time periods, one chart per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    results = {}
    periods = {
        "Notte (22–05)": NIGHT_HOURS,
        "Punta AM (07–08)": AM_PEAK,
        "Mezzog. (12–13)": MIDDAY,
        "Punta PM (17–18)": PM_PEAK,
    }
    pcols = ["#7570b3", "#d95f02", "#1b9e77", "#e7298a"]
    for d in cg["directions"]:
        sub = segments[(segments["carriage"] == carriage)
                       & (segments["direction"] == d)
                       & (segments["day_type"] == day_type)]
        if sub.empty:
            continue
        fig, ax = plt.subplots(figsize=(12, 5))
        for (pname, phours), pcol in zip(periods.items(), pcols):
            psub = sub[sub["hour"].isin(phours)]
            prof = psub.groupby("seg_idx").agg(
                v85=("p85", "mean"), dist=("cum_dist_mid", "first"),
            ).sort_values("dist")
            ax.plot(prof["dist"] / 1000, prof["v85"], color=pcol,
                    label=pname, linewidth=1.8)
        # Draw speed limit line using predominant limit
        predom_limit = get_speed_limit_for_group(
            segments[(segments["carriage"] == carriage)
                     & (segments["direction"] == d)], carriage)
        ax.axhline(predom_limit, color="red", ls="--", lw=1, alpha=.7,
                   label=f"Limite {predom_limit} km/h")
        ax.set_xlabel("Progressiva (km)")
        ax.set_ylabel("V85 (km/h)")
        label = cg["dir_labels"][d]
        ax.set_title(f"Profilo Spaziale V85 — {DAY_TYPE_LABELS[day_type]} {label}")
        ax.legend(fontsize=8); ax.grid(True, alpha=.3)
        results[label] = fig_to_base64(fig)
    return results


def chart_speed_variability(segments, carriage, day_type="Inv. Feriale"):
    """Std dev heatmaps (progressive km) + spatial profiles per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    ncols = len(dirs)
    fig, axes = plt.subplots(2, ncols, figsize=(7 * ncols, 10))
    if ncols == 1:
        axes = axes.reshape(2, 1)
    norm = mcolors.Normalize(vmin=0, vmax=25)
    cmap = plt.cm.YlOrRd
    for ci, d in enumerate(dirs):
        sub = segments[(segments["carriage"] == carriage)
                       & (segments["direction"] == d)
                       & (segments["day_type"] == day_type)]
        if sub.empty:
            continue
        label = cg["dir_labels"][d]
        _make_heatmap_progressive(
            sub, "std_speed", f"Dev. Std. — {label}",
            cmap, norm, fig_ax=(fig, axes[0, ci]))
        ax2 = axes[1, ci]
        prof = sub.groupby("seg_idx").agg(
            std_mean=("std_speed", "mean"), std_max=("std_speed", "max"),
            dist=("cum_dist_mid", "first"),
        ).sort_values("dist")
        dkm = prof["dist"] / 1000
        ax2.fill_between(dkm, 0, prof["std_max"], alpha=.2, color="orange",
                         label="Max orario")
        ax2.plot(dkm, prof["std_mean"], color="darkorange", lw=2,
                 label="Media giornaliera")
        ax2.set_xlabel("Progressiva (km)"); ax2.set_ylabel("Dev. Std. (km/h)")
        ax2.set_title(f"Variabilità Spaziale — {label}")
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle(f"Variabilità Velocità — {cg['label']} {DAY_TYPE_LABELS[day_type]}",
                 fontsize=12, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_night_analysis(segments, carriage, day_type="Inv. Feriale"):
    """Night analysis: histograms + spatial profiles per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    ncols = len(dirs)
    fig, axes = plt.subplots(2, ncols, figsize=(7 * ncols, 10))
    if ncols == 1:
        axes = axes.reshape(2, 1)
    sub_all = segments[(segments["carriage"] == carriage)
                       & (segments["day_type"] == day_type)]
    for ci, d in enumerate(dirs):
        sub_dir = sub_all[sub_all["direction"] == d]
        if sub_dir.empty:
            continue
        night = sub_dir[sub_dir["hour"].isin(NIGHT_HOURS)]["avg_speed"]
        day = sub_dir[~sub_dir["hour"].isin(NIGHT_HOURS)]["avg_speed"]
        ax = axes[0, ci]
        ax.hist(day, bins=40, alpha=.6, color="#1f77b4",
                label="Diurno (06–21)", density=True)
        ax.hist(night, bins=40, alpha=.6, color="#9467bd",
                label="Notturno (22–05)", density=True)
        ax.axvline(50, color="red", ls="--", lw=1, label="Limite 50 km/h")
        ax.set_xlabel("Velocità media (km/h)"); ax.set_ylabel("Densità")
        label = cg["dir_labels"][d]
        ax.set_title(f"Distribuzione — {label}")
        ax.legend(fontsize=8)

        ax2 = axes[1, ci]
        sub_n = sub_dir[sub_dir["hour"].isin(NIGHT_HOURS)]
        prof = sub_n.groupby("seg_idx").agg(
            v85=("p85", "mean"), avg=("avg_speed", "mean"),
            dist=("cum_dist_mid", "first"),
        ).sort_values("dist")
        dkm = prof["dist"] / 1000
        ax2.plot(dkm, prof["v85"], color="#9467bd", lw=2, label="V85 notturno")
        ax2.plot(dkm, prof["avg"], color="#1f77b4", lw=2, label="Vel. media notturna")
        ax2.axhline(50, color="red", ls="--", lw=1, alpha=.7)
        ax2.set_xlabel("Progressiva (km)"); ax2.set_ylabel("Velocità (km/h)")
        ax2.set_title(f"Profilo Notturno — {label}")
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle(f"Analisi Notturna — {cg['label']} {DAY_TYPE_LABELS[day_type]}",
                 fontsize=12, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_season_comparison(segments, carriage):
    """Compare Invernale vs Estiva (Feriale) — avg speed + V85 per direction."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    colors_dt = {"Inv. Feriale": "#1f77b4", "Est. Feriale": "#ff7f0e"}
    ncols = len(dirs)
    fig, axes = plt.subplots(2, ncols, figsize=(7 * ncols, 10))
    if ncols == 1:
        axes = axes.reshape(2, 1)
    for ci, d in enumerate(dirs):
        ax = axes[0, ci]
        for dt, ls_style in [("Inv. Feriale", "-"), ("Est. Feriale", "--")]:
            sub = segments[(segments["carriage"] == carriage)
                           & (segments["direction"] == d)
                           & (segments["day_type"] == dt)]
            if sub.empty:
                continue
            hm = sub.groupby("hour")["avg_speed"].mean().reset_index().sort_values("hour")
            ax.plot(hm["hour"], hm["avg_speed"], color=colors_dt[dt],
                    ls=ls_style, lw=2, label=DAY_TYPE_LABELS[dt],
                    marker="o", markersize=3)
        ax.axhline(50, color="red", ls="--", lw=1, alpha=.7)
        label = cg["dir_labels"][d]
        ax.set_title(f"Vel. Media — {label}")
        ax.set_xlabel("Ora"); ax.set_ylabel("km/h")
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=7)
        ax.legend(fontsize=8); ax.grid(True, alpha=.3)

        ax2 = axes[1, ci]
        for dt, ls_style in [("Inv. Feriale", "-"), ("Est. Feriale", "--")]:
            sub = segments[(segments["carriage"] == carriage)
                           & (segments["direction"] == d)
                           & (segments["day_type"] == dt)]
            if sub.empty:
                continue
            hm = sub.groupby("hour")["p85"].mean().reset_index().sort_values("hour")
            ax2.plot(hm["hour"], hm["p85"], color=colors_dt[dt],
                     ls=ls_style, lw=2, label=DAY_TYPE_LABELS[dt],
                     marker="o", markersize=3)
        ax2.axhline(50, color="red", ls="--", lw=1, alpha=.7)
        ax2.set_title(f"V85 — {label}")
        ax2.set_xlabel("Ora"); ax2.set_ylabel("km/h")
        ax2.set_xticks(range(24)); ax2.set_xticklabels(HOUR_SHORT, fontsize=7)
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle(f"Confronto Invernale vs Estiva (Feriali) — {cg['label']}",
                 fontsize=12, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


def generate_all_charts(segments, summaries):
    """Generate all charts for all carriage groups. Returns nested dict."""
    charts = {}
    for carriage in CARRIAGE_GROUPS:
        c = {}
        c["speed_by_hour"] = chart_speed_by_hour(summaries, carriage)
        c["v85_by_hour"] = chart_v85_by_hour(segments, carriage)
        c["heatmaps"] = chart_heatmaps(segments, carriage, "Inv. Feriale")
        c["exceedance"] = chart_exceedance(segments, carriage, "Inv. Feriale")
        c["v85_spatial"] = chart_v85_spatial(segments, carriage, "Inv. Feriale")
        c["variability"] = chart_speed_variability(segments, carriage, "Inv. Feriale")
        c["night"] = chart_night_analysis(segments, carriage, "Inv. Feriale")
        c["season_comparison"] = chart_season_comparison(segments, carriage)
        charts[carriage] = c
    return charts


# ================================================================
# SECTION 5 — STATIC MAP GENERATION
# ================================================================

def _build_speed_gdf(segments, carriage, direction, day_type, value_col,
                     hour_filter=None, agg_func="mean"):
    """Aggregate segment speeds → GeoDataFrame in EPSG:4326."""
    sub = segments[(segments["carriage"] == carriage)
                   & (segments["direction"] == direction)
                   & (segments["day_type"] == day_type)]
    if hour_filter is not None:
        sub = sub[sub["hour"].isin(hour_filter)]
    agg = sub.groupby("seg_idx").agg({
        value_col: agg_func, "geometry": "first",
        "cum_dist_start": "first", "cum_dist_mid": "first",
        "seg_distance": "first", "streetName": "first",
        "speedLimit": "first",
    }).reset_index()
    return gpd.GeoDataFrame(agg, geometry="geometry", crs="EPSG:4326")


def _add_basemap(ax, crs="EPSG:4326"):
    """Add contextily basemap tiles."""
    try:
        cx.add_basemap(ax, crs=crs, source=cx.providers.CartoDB.Positron,
                        zoom=16, attribution=False)
        return True
    except Exception:
        pass
    try:
        cx.add_basemap(ax, crs=crs, source=cx.providers.OpenStreetMap.Mapnik,
                        zoom=16, attribution=False)
        return True
    except Exception:
        pass
    ax.set_facecolor("#eef0f4")
    ax.grid(True, alpha=0.25, linestyle="--", color="#888", zorder=0)
    ax.tick_params(labelsize=7, colors="#555")
    ax.ticklabel_format(useOffset=False, style="plain")
    for spine in ax.spines.values():
        spine.set_edgecolor("#999"); spine.set_linewidth(0.8)
    return False


def _annotate_segments(ax, gdf, value_col):
    """Progressive markers, start/end, stats box."""
    total_dist = gdf["cum_dist_start"].max() + gdf["seg_distance"].max()
    interval = 500 if total_dist > 3000 else (250 if total_dist > 1000 else 100)
    for target_m in range(0, int(total_dist) + interval, interval):
        candidates = gdf.loc[(gdf["cum_dist_start"] <= target_m)
                              & (gdf["cum_dist_start"] + gdf["seg_distance"] >= target_m)]
        if candidates.empty:
            continue
        row = candidates.iloc[0]
        d = row["seg_distance"]
        frac = (target_m - row["cum_dist_start"]) / d if d > 0 else 0.0
        frac = max(0.0, min(1.0, frac))
        pt = row["geometry"].interpolate(frac, normalized=True)
        label = f"{target_m} m" if target_m < 1000 else f"{target_m/1000:.1f} km"
        ax.annotate(label, xy=(pt.x, pt.y), fontsize=6, fontweight="bold",
                    color="#1a237e", ha="left", va="bottom",
                    xytext=(5, 5), textcoords="offset points",
                    bbox=dict(boxstyle="round,pad=0.15", fc="white",
                              ec="#1a237e", alpha=0.85, lw=0.5), zorder=5)
        ax.plot(pt.x, pt.y, "o", color="#1a237e", markersize=3, zorder=5)

    start_pt = list(gdf.iloc[0]["geometry"].coords)[0]
    end_pt = list(gdf.iloc[-1]["geometry"].coords)[-1]
    ax.plot(start_pt[0], start_pt[1], "^", color="green", markersize=10,
            zorder=6, label="Inizio")
    ax.plot(end_pt[0], end_pt[1], "v", color="red", markersize=10,
            zorder=6, label="Fine")

    vals = gdf[value_col]
    metric_label = "V85" if "p85" in value_col else "Vel. media"
    stats_text = (f"{metric_label}: media={vals.mean():.1f}, "
                  f"max={vals.max():.1f}, min={vals.min():.1f} km/h")
    ax.text(0.02, 0.02, stats_text, transform=ax.transAxes, fontsize=8,
            va="bottom", ha="left",
            bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="gray", alpha=0.9))
    ax.legend(loc="upper right", fontsize=7)


def generate_static_speed_map(segments, carriage, direction, day_type, value_col,
                               hour_filter=None, title="", fname="map.png",
                               agg_func="mean"):
    """Render a static matplotlib map with coloured segments."""
    gdf = _build_speed_gdf(segments, carriage, direction, day_type, value_col,
                            hour_filter, agg_func=agg_func)
    if gdf.empty:
        return None

    fig, ax = plt.subplots(figsize=(10, 14))
    norm = mcolors.TwoSlopeNorm(vmin=20, vcenter=50, vmax=80)
    cmap = plt.cm.RdYlGn_r

    gdf.plot(ax=ax, color="#d0d0d0", linewidth=9, zorder=1, alpha=0.8)
    gdf.plot(ax=ax, column=value_col, cmap=cmap, norm=norm,
             linewidth=5, legend=False, zorder=2)
    has_tiles = _add_basemap(ax, crs="EPSG:4326")
    if has_tiles:
        ax.set_axis_off()
    else:
        ax.set_xlabel("Longitudine", fontsize=9)
        ax.set_ylabel("Latitudine", fontsize=9)
    _annotate_segments(ax, gdf, value_col)

    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=ax, shrink=0.4, pad=0.03, aspect=30)
    cbar.set_label("km/h", fontsize=10)
    cbar.ax.axhline(y=50, color="black", linewidth=1.5, linestyle="--")
    cbar.ax.text(1.3, 50, "Limite 50", transform=cbar.ax.get_yaxis_transform(),
                 va="center", fontsize=8, color="black")
    ax.set_title(title, fontsize=12, fontweight="bold", pad=12)
    fig.tight_layout()

    out_path = STATIC_MAPS_DIR / fname
    fig.savefig(str(out_path), dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(out_path)


def generate_static_speed_map_sidebyside(segments, carriage, day_type, value_col,
                                          dir_a, dir_b, label_a, label_b,
                                          hour_filter=None, title="",
                                          fname="map_sbs.png", agg_func="mean"):
    """Side-by-side static map for two directions."""
    gdf_a = _build_speed_gdf(segments, carriage, dir_a, day_type, value_col,
                              hour_filter, agg_func=agg_func)
    gdf_b = _build_speed_gdf(segments, carriage, dir_b, day_type, value_col,
                              hour_filter, agg_func=agg_func)
    if gdf_a.empty or gdf_b.empty:
        return None

    fig, axes = plt.subplots(1, 2, figsize=(20, 14))
    norm = mcolors.TwoSlopeNorm(vmin=20, vcenter=50, vmax=80)
    cmap = plt.cm.RdYlGn_r

    for ax, gdf_orig, lbl in [(axes[0], gdf_a, label_a), (axes[1], gdf_b, label_b)]:
        gdf_orig.plot(ax=ax, color="#d0d0d0", linewidth=9, zorder=1, alpha=0.8)
        gdf_orig.plot(ax=ax, column=value_col, cmap=cmap, norm=norm,
                      linewidth=5, legend=False, zorder=2)
        has_tiles = _add_basemap(ax, crs="EPSG:4326")
        _annotate_segments(ax, gdf_orig, value_col)
        if has_tiles:
            ax.set_axis_off()
        ax.set_title(lbl, fontsize=12, fontweight="bold")

    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=axes.tolist(), shrink=0.4, pad=0.02, aspect=30)
    cbar.set_label("km/h", fontsize=10)
    cbar.ax.axhline(y=50, color="black", linewidth=1.5, linestyle="--")

    fig.suptitle(title, fontsize=13, fontweight="bold", y=1.01)
    fig.tight_layout()

    out_path = STATIC_MAPS_DIR / fname
    fig.savefig(str(out_path), dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(out_path)


def generate_all_static_maps(segments):
    """Generate static speed maps for all carriage groups."""
    STATIC_MAPS_DIR.mkdir(parents=True, exist_ok=True)
    paths = []
    day_type = "Inv. Feriale"

    configs = [
        ("avg_speed", None,        "Velocità Media (24h)",      "avg_speed_24h", "mean"),
        ("p85",       None,        "V85 (24h)",                 "v85_24h",       "mean"),
        ("avg_speed", NIGHT_HOURS, "Velocità Media Notturna",   "avg_speed_night", "mean"),
        ("p85",       NIGHT_HOURS, "V85 Notturno",              "v85_night",     "mean"),
    ]

    for carriage in CARRIAGE_GROUPS:
        cg = CARRIAGE_GROUPS[carriage]
        dirs = cg["directions"]
        slug = carriage.lower().replace(". ", "_").replace(" ", "_")

        # Individual maps per direction
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            for vcol, hfilter, label, prefix, agg in configs:
                period = "22–05" if hfilter is not None else "tutte le ore"
                title = f"{label} — {cg['label']} {cg['dir_labels'][d]} ({period})"
                fname = f"{prefix}_{slug}_{d_slug}.png"
                p = generate_static_speed_map(
                    segments, carriage, d, day_type, vcol,
                    hour_filter=hfilter, title=title, fname=fname, agg_func=agg)
                if p:
                    paths.append(p)
                    print(f"  - {fname}")

        # Side-by-side maps (only for Centrale with 2 directions)
        if carriage == "Centrale":
            d_a, d_b = dirs[0], dirs[1]
            la, lb = cg["dir_labels"][d_a], cg["dir_labels"][d_b]
            for vcol, hfilter, label, prefix, agg in configs:
                period = "22–05" if hfilter is not None else "tutte le ore"
                title = f"{label} — {cg['label']} ({period})"
                fname = f"{prefix}_{slug}_sidebyside.png"
                p = generate_static_speed_map_sidebyside(
                    segments, carriage, day_type, vcol, d_a, d_b, la, lb,
                    hour_filter=hfilter, title=title, fname=fname, agg_func=agg)
                if p:
                    paths.append(p)
                    print(f"  - {fname}")

        # Max V85 maps
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            title = f"V85 Massimo (ora peggiore) — {cg['label']} {cg['dir_labels'][d]}"
            fname = f"max_v85_{slug}_{d_slug}.png"
            p = generate_static_speed_map(
                segments, carriage, d, day_type, "p85",
                hour_filter=None, title=title, fname=fname, agg_func="max")
            if p:
                paths.append(p)
                print(f"  - {fname}")

        # Max V85 side-by-side (Centrale only)
        if carriage == "Centrale":
            d_a, d_b = dirs[0], dirs[1]
            la, lb = cg["dir_labels"][d_a], cg["dir_labels"][d_b]
            title = f"V85 Massimo (ora peggiore) — {cg['label']}"
            fname = f"max_v85_{slug}_sidebyside.png"
            p = generate_static_speed_map_sidebyside(
                segments, carriage, day_type, "p85", d_a, d_b, la, lb,
                hour_filter=None, title=title, fname=fname, agg_func="max")
            if p:
                paths.append(p)
                print(f"  - {fname}")

    return paths


# ================================================================
# SECTION 6 — INTERACTIVE MAP GENERATION
# ================================================================

def _seg_geodataframe(segments, carriage, direction, day_type, value_col, agg="mean"):
    """Build a GeoDataFrame aggregated over hours for one direction."""
    sub = segments[(segments["carriage"] == carriage)
                   & (segments["direction"] == direction)
                   & (segments["day_type"] == day_type)]
    agg_dict = {
        value_col: agg, "geometry": "first",
        "cum_dist_mid": "first", "cum_dist_start": "first",
        "streetName": "first", "speedLimit": "first", "seg_distance": "first",
        "avg_speed": "mean", "p85": "mean", "std_speed": "mean",
    }
    grouped = sub.groupby("seg_idx").agg(agg_dict).reset_index()
    return gpd.GeoDataFrame(grouped, geometry="geometry", crs="EPSG:4326")


def _add_folium_progressive(m, gdf, interval_m=500):
    """Add progressive distance markers + start/end flags to a Folium map."""
    total_dist = gdf["cum_dist_start"].max() + gdf["seg_distance"].max()
    if interval_m is None:
        if total_dist > 10000:
            interval_m = 1000
        elif total_dist > 3000:
            interval_m = 500
        else:
            interval_m = 250
    for target_m in range(0, int(total_dist) + interval_m, interval_m):
        candidates = gdf.loc[(gdf["cum_dist_start"] <= target_m)
                              & (gdf["cum_dist_start"] + gdf["seg_distance"] >= target_m)]
        if candidates.empty:
            continue
        row = candidates.iloc[0]
        d = row["seg_distance"]
        frac = (target_m - row["cum_dist_start"]) / d if d > 0 else 0.0
        frac = max(0.0, min(1.0, frac))
        pt = row["geometry"].interpolate(frac, normalized=True)
        label = f"{target_m} m" if target_m < 1000 else f"{target_m/1000:.1f} km"
        folium.Marker(
            [pt.y, pt.x],
            icon=folium.DivIcon(
                html=(f'<div style="font-size:10px;font-weight:bold;background:white;'
                      f'padding:2px 5px;border:1px solid #1a237e;border-radius:3px;'
                      f'white-space:nowrap;box-shadow:1px 1px 3px rgba(0,0,0,.3);'
                      f'color:#1a237e;">{label}</div>'),
                icon_size=(70, 22), icon_anchor=(35, 11)),
        ).add_to(m)
    sc = list(gdf.iloc[0]["geometry"].coords)[0]
    ec = list(gdf.iloc[-1]["geometry"].coords)[-1]
    folium.Marker([sc[1], sc[0]], icon=folium.Icon(color="green"),
                  popup="Inizio (0 m)").add_to(m)
    folium.Marker([ec[1], ec[0]], icon=folium.Icon(color="red"),
                  popup=f"Fine ({total_dist:.0f} m)").add_to(m)


def _add_folium_stats(m, gdf, value_col, title):
    """Add a floating stats box to a Folium map."""
    vals = gdf[value_col]
    metric = "V85" if "p85" in value_col else ("Dev. Std." if "std" in value_col else "Vel. media")
    html = (
        f'<div style="position:fixed;bottom:30px;left:10px;z-index:9999;'
        f'background:white;padding:10px 14px;border-radius:6px;'
        f'box-shadow:0 2px 8px rgba(0,0,0,.3);font-family:sans-serif;'
        f'font-size:12px;max-width:320px;line-height:1.5;">'
        f'<b style="color:#1a237e;">{title}</b><br>'
        f'{metric}: media={vals.mean():.1f}, max={vals.max():.1f}, '
        f'min={vals.min():.1f} km/h'
        f'</div>'
    )
    m.get_root().html.add_child(folium.Element(html))


def make_folium_map(gdf, value_col, title, vmin=20, vmax=80,
                     reverse_cmap=False, add_progressive=True,
                     add_stats=True):
    """Folium map with coloured segments, legend, markers, stats."""
    clat = gdf.geometry.centroid.y.mean()
    clon = gdf.geometry.centroid.x.mean()
    # Adapt zoom to route length
    total_dist = gdf["cum_dist_start"].max() + gdf["seg_distance"].max()
    zoom = 12 if total_dist > 10000 else (13 if total_dist > 5000 else 15)
    m = folium.Map(location=[clat, clon], zoom_start=zoom,
                   tiles="CartoDB positron")

    colors = (["green", "yellow", "red"] if not reverse_cmap
              else ["red", "yellow", "green"])
    cmap = cm.LinearColormap(colors=colors, vmin=vmin, vmax=vmax,
                              caption=f"{title} (km/h)")
    for _, row in gdf.iterrows():
        coords = [[c[1], c[0]] for c in row.geometry.coords]
        popup_html = (
            f"<b>{row['streetName']}</b><br>"
            f"Progr.: {row['cum_dist_start']:.0f} m<br>"
            f"Vel. media: {row['avg_speed']:.1f} km/h<br>"
            f"V85: {row['p85']:.1f} km/h<br>"
            f"Dev. std: {row['std_speed']:.1f} km/h<br>"
            f"Limite: {row['speedLimit']} km/h<br>"
            f"Lungh.: {row['seg_distance']:.0f} m"
        )
        folium.PolyLine(coords, weight=10, color=cmap(row[value_col]),
                        opacity=0.9,
                        popup=folium.Popup(popup_html, max_width=280)).add_to(m)
    cmap.add_to(m)
    if add_progressive:
        _add_folium_progressive(m, gdf)
    if add_stats:
        _add_folium_stats(m, gdf, value_col, title)
    return m


def make_folium_map_combined(gdf_a, gdf_b, value_col, title,
                               label_a="Dir. A", label_b="Dir. B",
                               vmin=20, vmax=80, reverse_cmap=False):
    """Folium map with BOTH directions on one map, each in a layer group."""
    all_lats = list(gdf_a.geometry.centroid.y) + list(gdf_b.geometry.centroid.y)
    all_lons = list(gdf_a.geometry.centroid.x) + list(gdf_b.geometry.centroid.x)
    clat, clon = np.mean(all_lats), np.mean(all_lons)
    total_dist = max(
        gdf_a["cum_dist_start"].max() + gdf_a["seg_distance"].max(),
        gdf_b["cum_dist_start"].max() + gdf_b["seg_distance"].max(),
    )
    zoom = 12 if total_dist > 10000 else (13 if total_dist > 5000 else 14)
    m = folium.Map(location=[clat, clon], zoom_start=zoom,
                   tiles="CartoDB positron")

    colors = (["green", "yellow", "red"] if not reverse_cmap
              else ["red", "yellow", "green"])
    cmap = cm.LinearColormap(colors=colors, vmin=vmin, vmax=vmax,
                              caption=f"{title} (km/h)")

    for gdf_cur, lbl in [(gdf_a, label_a), (gdf_b, label_b)]:
        fg = folium.FeatureGroup(name=lbl)
        for _, row in gdf_cur.iterrows():
            coords = [[c[1], c[0]] for c in row.geometry.coords]
            popup_html = (
                f"<b>{lbl} — {row['streetName']}</b><br>"
                f"Progr.: {row['cum_dist_start']:.0f} m<br>"
                f"Vel. media: {row['avg_speed']:.1f} km/h<br>"
                f"V85: {row['p85']:.1f} km/h<br>"
                f"Limite: {row['speedLimit']} km/h"
            )
            folium.PolyLine(coords, weight=8, color=cmap(row[value_col]),
                            opacity=0.85,
                            popup=folium.Popup(popup_html, max_width=280)).add_to(fg)
        fg.add_to(m)
    cmap.add_to(m)
    folium.LayerControl(collapsed=False).add_to(m)

    # Combined stats box
    vals_all = pd.concat([gdf_a[value_col], gdf_b[value_col]])
    metric = "V85" if "p85" in value_col else ("Dev. Std." if "std" in value_col else "Vel. media")
    stats_html = (
        f'<div style="position:fixed;bottom:30px;left:10px;z-index:9999;'
        f'background:white;padding:10px 14px;border-radius:6px;'
        f'box-shadow:0 2px 8px rgba(0,0,0,.3);font-family:sans-serif;'
        f'font-size:12px;max-width:350px;line-height:1.5;">'
        f'<b style="color:#1a237e;">{title}</b><br>'
        f'{label_a}: {metric} media={gdf_a[value_col].mean():.1f} km/h<br>'
        f'{label_b}: {metric} media={gdf_b[value_col].mean():.1f} km/h'
        f'</div>'
    )
    m.get_root().html.add_child(folium.Element(stats_html))
    return m


def generate_maps(segments, seg_stats):
    """Generate all interactive Folium maps for all carriage groups."""
    MAPS_DIR.mkdir(parents=True, exist_ok=True)
    map_files = {}
    day_type = "Inv. Feriale"

    configs = [
        ("avg_speed", None,        "Velocità Media (24h)",      "avg_speed_24h", "mean", 20, 80, False),
        ("p85",       None,        "V85 (24h)",                 "v85_24h",       "mean", 20, 80, False),
        ("avg_speed", NIGHT_HOURS, "Velocità Media Notturna",   "avg_speed_night", "mean", 20, 80, False),
        ("p85",       NIGHT_HOURS, "V85 Notturno",              "v85_night",     "mean", 20, 80, False),
        ("p85",       None,        "V85 Massimo (ora peggiore)","max_v85",       "max",  20, 90, False),
    ]

    for carriage in CARRIAGE_GROUPS:
        cg = CARRIAGE_GROUPS[carriage]
        dirs = cg["directions"]
        slug = carriage.lower().replace(". ", "_").replace(" ", "_")

        # --- Individual maps per direction ---
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            for vcol, hfilter, label, prefix, agg, vmin, vmax, rev in configs:
                period = "22–05" if hfilter is not None else "tutte le ore"
                title = f"{label} — {cg['label']} {cg['dir_labels'][d]} ({period})"
                fname = f"{prefix}_{slug}_{d_slug}.html"
                sub = segments[segments["carriage"] == carriage]
                if hfilter is not None:
                    sub = sub[sub["hour"].isin(hfilter)]
                gdf = _seg_geodataframe(sub, carriage, d, day_type, vcol, agg)
                if gdf.empty:
                    continue
                m = make_folium_map(gdf, vcol, title, vmin, vmax, rev)
                m.save(str(MAPS_DIR / fname))
                map_files[fname] = title

        # --- Combined both-directions maps (Centrale only) ---
        if carriage == "Centrale":
            d_a, d_b = dirs[0], dirs[1]
            la, lb = cg["dir_labels"][d_a], cg["dir_labels"][d_b]
            for vcol, hfilter, label, prefix, agg, vmin, vmax, rev in configs:
                period = "22–05" if hfilter is not None else "tutte le ore"
                title = f"{label} — {cg['label']} ({period})"
                fname = f"{prefix}_{slug}_combined.html"
                sub = segments[segments["carriage"] == carriage]
                if hfilter is not None:
                    sub = sub[sub["hour"].isin(hfilter)]
                gdf_a = _seg_geodataframe(sub, carriage, d_a, day_type, vcol, agg)
                gdf_b = _seg_geodataframe(sub, carriage, d_b, day_type, vcol, agg)
                if gdf_a.empty or gdf_b.empty:
                    continue
                m = make_folium_map_combined(gdf_a, gdf_b, vcol, title, la, lb,
                                             vmin, vmax, rev)
                m.save(str(MAPS_DIR / fname))
                map_files[fname] = title

        # --- Variability maps ---
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            label = cg["dir_labels"][d]
            title = f"Variabilità — {cg['label']} {label}"
            fname = f"std_{slug}_{d_slug}.html"
            gdf = _seg_geodataframe(segments, carriage, d, day_type, "std_speed", "mean")
            if gdf.empty:
                continue
            m = make_folium_map(gdf, "std_speed", title, 5, 22, True)
            m.save(str(MAPS_DIR / fname))
            map_files[fname] = title

        # --- Progressive distance maps ---
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            sub = segments[(segments["carriage"] == carriage)
                           & (segments["direction"] == d)
                           & (segments["day_type"] == day_type)]
            si = sub.drop_duplicates("seg_idx").sort_values("seg_idx")
            if si.empty:
                continue
            clat = si.geometry.apply(lambda g: g.centroid.y).mean()
            clon = si.geometry.apply(lambda g: g.centroid.x).mean()
            total_dist = si["cum_dist_start"].max() + si["seg_distance"].max()
            zoom = 12 if total_dist > 10000 else (13 if total_dist > 5000 else 15)
            m = folium.Map(location=[clat, clon], zoom_start=zoom,
                           tiles="CartoDB positron")
            for _, row in si.iterrows():
                coords = [[c[1], c[0]] for c in row["geometry"].coords]
                folium.PolyLine(coords, weight=6, color="#1a237e", opacity=.8).add_to(m)
            _add_folium_progressive(m, si)
            fname = f"progressive_{slug}_{d_slug}.html"
            m.save(str(MAPS_DIR / fname))
            map_files[fname] = f"Progressive — {cg['dir_labels'][d]}"

        # --- Top-10 V85 maps ---
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            sub = segments[(segments["carriage"] == carriage)
                           & (segments["direction"] == d)
                           & (segments["day_type"] == day_type)]
            si = sub.drop_duplicates("seg_idx").sort_values("seg_idx")
            if si.empty:
                continue
            top = seg_stats[
                (seg_stats["carriage"] == carriage)
                & (seg_stats["direction"] == d)
                & (seg_stats["day_type"] == day_type)
            ].nlargest(10, "max_v85")
            top_set = set(top["seg_idx"].values)
            clat = si.geometry.apply(lambda g: g.centroid.y).mean()
            clon = si.geometry.apply(lambda g: g.centroid.x).mean()
            total_dist = si["cum_dist_start"].max() + si["seg_distance"].max()
            zoom = 12 if total_dist > 10000 else (13 if total_dist > 5000 else 15)
            m = folium.Map(location=[clat, clon], zoom_start=zoom,
                           tiles="CartoDB positron")
            for _, row in si.iterrows():
                coords = [[c[1], c[0]] for c in row["geometry"].coords]
                is_top = row["seg_idx"] in top_set
                folium.PolyLine(
                    coords,
                    weight=10 if is_top else 4,
                    color="#e41a1c" if is_top else "#aaaaaa",
                    opacity=0.9 if is_top else 0.4,
                ).add_to(m)
            for rank, (_, row) in enumerate(top.iterrows(), 1):
                match = si[si["seg_idx"] == row["seg_idx"]]
                if match.empty:
                    continue
                c = match.iloc[0]["geometry"].centroid
                folium.Marker(
                    [c.y, c.x],
                    icon=folium.DivIcon(
                        html=(f'<div style="font-size:11px;background:#e41a1c;color:white;'
                              f'padding:2px 6px;border-radius:50%;text-align:center;'
                              f'width:24px;height:24px;line-height:20px;font-weight:bold;'
                              f'box-shadow:1px 1px 3px rgba(0,0,0,.4);">{rank}</div>'),
                        icon_size=(24, 24), icon_anchor=(12, 12)),
                    popup=(f"<b>#{rank}</b><br>Progr.: {row['cum_dist_start']:.0f} m<br>"
                           f"V85 max: {row['max_v85']:.0f} km/h<br>"
                           f"Vel. media: {row['all_avg_speed']:.1f} km/h"),
                ).add_to(m)
            fname = f"top10_v85_{slug}_{d_slug}.html"
            m.save(str(MAPS_DIR / fname))
            map_files[fname] = f"Top 10 V85 — {cg['dir_labels'][d]}"

    return map_files


# ================================================================
# SECTION 7 — TABLES & HTML REPORT
# ================================================================

def build_summary_tables(segments, summaries, seg_stats):
    """Build summary tables per carriage group."""
    tables = {}

    for carriage in CARRIAGE_GROUPS:
        cg = CARRIAGE_GROUPS[carriage]
        dirs = cg["directions"]
        slug = carriage.lower().replace(". ", "_").replace(" ", "_")

        # Overview table
        overview_rows = []
        for dt in ALL_DAY_TYPES:
            for d in dirs:
                ss = summaries[(summaries["carriage"] == carriage)
                               & (summaries["direction"] == d)
                               & (summaries["day_type"] == dt)]
                if ss.empty:
                    continue
                overview_rows.append({
                    "Periodo": DAY_TYPE_LABELS[dt],
                    "Direzione": cg["dir_labels"][d],
                    "Vel. Media 24h": f"{ss['harm_avg_speed'].mean():.1f}",
                    "Vel. Punta AM": f"{ss[ss['hour'].isin(AM_PEAK)]['harm_avg_speed'].mean():.1f}",
                    "Vel. Punta PM": f"{ss[ss['hour'].isin(PM_PEAK)]['harm_avg_speed'].mean():.1f}",
                    "Vel. Notturna": f"{ss[ss['hour'].isin(NIGHT_HOURS)]['harm_avg_speed'].mean():.1f}",
                })
        tables[f"overview_{slug}"] = pd.DataFrame(overview_rows)

        # Exceedance table
        exc_rows = []
        for dt in ["Inv. Feriale", "Est. Feriale"]:
            for d in dirs:
                sub = segments[(segments["carriage"] == carriage)
                               & (segments["direction"] == d)
                               & (segments["day_type"] == dt)]
                if sub.empty:
                    continue
                total_m = sub.drop_duplicates("seg_idx")["seg_distance"].sum()

                def _pct(col):
                    vals = []
                    for h in range(24):
                        hs = sub[sub["hour"] == h]
                        vals.append(hs.loc[hs[col] > hs["speedLimit"],
                                           "seg_distance"].sum() / total_m * 100)
                    return np.mean(vals)

                exc_rows.append({
                    "Periodo": DAY_TYPE_LABELS[dt],
                    "Direzione": cg["dir_labels"][d],
                    "% km Vel.Media > Limite (media 24h)": f"{_pct('avg_speed'):.1f}%",
                    "% km V85 > Limite (media 24h)":       f"{_pct('p85'):.1f}%",
                    "V85 massimo (km/h)":                  f"{sub['p85'].max():.0f}",
                })
        tables[f"exceedance_{slug}"] = pd.DataFrame(exc_rows)

        # Top 10
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            st = seg_stats[
                (seg_stats["carriage"] == carriage)
                & (seg_stats["direction"] == d)
                & (seg_stats["day_type"] == "Inv. Feriale")
            ].nlargest(10, "max_v85").reset_index(drop=True)
            if st.empty:
                continue
            tbl = pd.DataFrame({
                "#": range(1, len(st) + 1),
                "Progressiva (m)": st["cum_dist_start"].apply(lambda x: f"{x:.0f}"),
                "V85 Max (km/h)":  st["max_v85"].apply(lambda x: f"{x:.0f}"),
                "Vel. Media (km/h)": st["all_avg_speed"].apply(lambda x: f"{x:.1f}"),
                "Limite (km/h)":     st["speedLimit"].apply(lambda x: f"{x:.0f}"),
            })
            tables[f"top_fast_{slug}_{d_slug}"] = tbl

    return tables


def df_to_html_table(df):
    return df.to_html(index=False, classes="data-table", border=0)


def _iframe(fname, title):
    return (f'<h3>{title}</h3>'
            f'<iframe src="maps/{fname}" width="100%" height="500" '
            f'frameborder="0" style="border:1px solid #ddd;border-radius:4px;'
            f'margin-bottom:20px;"></iframe>')


def _build_carriage_section(carriage, charts, tables, map_files, section_num_base):
    """Build the HTML for a single carriage group's analysis."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    slug = carriage.lower().replace(". ", "_").replace(" ", "_")
    cdata = charts[carriage]
    s = section_num_base

    # Heatmaps
    heatmap_imgs = "".join(
        f'<h3>{lab}</h3><img src="data:image/png;base64,{b64}" style="max-width:100%;">'
        for lab, b64 in cdata["heatmaps"].items()
    )

    # V85 spatial
    v85_spatial_imgs = "".join(
        f'<h3>{lab}</h3><img src="data:image/png;base64,{b64}" style="max-width:100%;">'
        for lab, b64 in cdata["v85_spatial"].items()
    )

    # Progressive iframes
    progressive_iframes = ""
    for d in dirs:
        d_slug = d.lower().replace(" ", "_").replace("-", "_")
        fname = f"progressive_{slug}_{d_slug}.html"
        if fname in map_files:
            progressive_iframes += _iframe(fname, f"Progressive — {cg['dir_labels'][d]}")

    # Top 10 iframes + tables
    top10_sections = ""
    for d in dirs:
        d_slug = d.lower().replace(" ", "_").replace("-", "_")
        tkey = f"top_fast_{slug}_{d_slug}"
        fname = f"top10_v85_{slug}_{d_slug}.html"
        dlabel = cg["dir_labels"][d]
        if tkey in tables:
            top10_sections += f"<h3>Top 10 Tratti V85 — {dlabel}</h3>"
            top10_sections += df_to_html_table(tables[tkey])
        if fname in map_files:
            top10_sections += _iframe(fname, f"Mappa Top 10 — {dlabel}")

    # Interactive map iframes (main speed maps)
    map_iframes = ""
    map_configs = [
        ("avg_speed_24h", "Velocità Media (24h)"),
        ("v85_24h", "V85 (24h)"),
        ("avg_speed_night", "Velocità Media Notturna"),
        ("v85_night", "V85 Notturno"),
        ("max_v85", "V85 Massimo (ora peggiore)"),
    ]
    for d in dirs:
        d_slug = d.lower().replace(" ", "_").replace("-", "_")
        for prefix, label in map_configs:
            fname = f"{prefix}_{slug}_{d_slug}.html"
            if fname in map_files:
                map_iframes += _iframe(fname, f"{label} — {cg['dir_labels'][d]}")
    # Combined maps (Centrale)
    if carriage == "Centrale":
        for prefix, label in map_configs:
            fname = f"{prefix}_{slug}_combined.html"
            if fname in map_files:
                map_iframes += _iframe(fname, f"{label} — Entrambe le direzioni")
    # Variability maps
    for d in dirs:
        d_slug = d.lower().replace(" ", "_").replace("-", "_")
        fname = f"std_{slug}_{d_slug}.html"
        if fname in map_files:
            map_iframes += _iframe(fname, f"Variabilità — {cg['dir_labels'][d]}")

    # Overview / exceedance tables
    overview_key = f"overview_{slug}"
    exceedance_key = f"exceedance_{slug}"
    overview_html = df_to_html_table(tables[overview_key]) if overview_key in tables else ""
    exceedance_html = df_to_html_table(tables[exceedance_key]) if exceedance_key in tables else ""

    return f"""
<!-- ===== {cg['label'].upper()} ===== -->
<section id="carriage_{slug}">
<h2 style="color:#fff;background:#1a237e;padding:15px 20px;border-radius:6px;">
{cg['label']}</h2>

<h3>{s}.1 Panoramica</h3>
{overview_html}

<h3>{s}.2 Superamento Limiti (pesato per km)</h3>
<div class="method-box">
<strong>Nota:</strong> il superamento &egrave; calcolato rispetto al limite di velocit&agrave;
specifico di ciascun segmento (variabile: 30, 40, 50, 60, 80 km/h).
</div>
{exceedance_html}

<h3>{s}.3 Profili Temporali</h3>
<h4>Velocit&agrave; Media per Ora</h4>
<img src="data:image/png;base64,{cdata['speed_by_hour']}" style="max-width:100%;">
<h4>V85 per Ora</h4>
<img src="data:image/png;base64,{cdata['v85_by_hour']}" style="max-width:100%;">

<h3>{s}.4 Mappe di Calore (Invernale Feriale)</h3>
{heatmap_imgs}

<h3>{s}.5 Superamento Limiti per Ora</h3>
<img src="data:image/png;base64,{cdata['exceedance']}" style="max-width:100%;">

<h3>{s}.6 Profili Spaziali V85</h3>
{v85_spatial_imgs}
{top10_sections}

<h3>{s}.7 Variabilit&agrave;</h3>
<img src="data:image/png;base64,{cdata['variability']}" style="max-width:100%;">

<h3>{s}.8 Analisi Notturna</h3>
<img src="data:image/png;base64,{cdata['night']}" style="max-width:100%;">

<h3>{s}.9 Confronto Invernale vs Estiva</h3>
<img src="data:image/png;base64,{cdata['season_comparison']}" style="max-width:100%;">

<h3>{s}.10 Progressive e Mappe Interattive</h3>
{progressive_iframes}
{map_iframes}
</section>
"""


def generate_html_report(charts, map_files, tables):
    """Build the full HTML report."""

    # Build carriage sections
    carriage_sections = ""
    for i, carriage in enumerate(CARRIAGE_GROUPS, start=2):
        carriage_sections += _build_carriage_section(
            carriage, charts, tables, map_files, section_num_base=i)

    html = f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Analisi Velocit&agrave; &mdash; Via Cristoforo Colombo, Roma</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',Tahoma,Geneva,Verdana,sans-serif;line-height:1.6;
  color:#333;background:#f9f9f9;max-width:1200px;margin:0 auto;padding:20px}}
header{{background:linear-gradient(135deg,#1a237e,#283593);color:#fff;
  padding:40px;border-radius:8px;margin-bottom:30px;text-align:center}}
header h1{{font-size:2em;margin-bottom:8px}}
header h2{{font-size:1.3em;font-weight:300;opacity:.9}}
header p{{margin-top:12px;opacity:.8;font-size:.95em}}
nav{{background:#fff;padding:20px 30px;border-radius:8px;margin-bottom:30px;
  box-shadow:0 2px 4px rgba(0,0,0,.1)}}
nav h3{{margin-bottom:10px;color:#1a237e}}
nav ol{{padding-left:20px}} nav li{{margin-bottom:4px}}
nav a{{color:#1565c0;text-decoration:none}} nav a:hover{{text-decoration:underline}}
section{{background:#fff;padding:30px;border-radius:8px;margin-bottom:20px;
  box-shadow:0 2px 4px rgba(0,0,0,.1)}}
h2{{color:#1a237e;border-bottom:2px solid #e8eaf6;padding-bottom:8px;margin-bottom:20px}}
h3{{color:#283593;margin:20px 0 10px 0}}
h4{{color:#3949ab;margin:15px 0 8px 0}}
img{{max-width:100%;height:auto;border-radius:4px;margin:10px 0}}
.data-table{{width:100%;border-collapse:collapse;margin:15px 0;font-size:.9em}}
.data-table th{{background:#1a237e;color:#fff;padding:10px 12px;text-align:left;font-weight:600}}
.data-table td{{padding:8px 12px;border-bottom:1px solid #e0e0e0}}
.data-table tr:nth-child(even){{background:#f5f5f5}}
.data-table tr:hover{{background:#e8eaf6}}
.method-box{{background:#e3f2fd;border-left:4px solid #1976d2;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}}
.insight-box{{background:#e8f5e9;border-left:4px solid #4caf50;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}}
.warning-box{{background:#fff3e0;border-left:4px solid #ff9800;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}}
footer{{text-align:center;padding:20px;color:#999;font-size:.85em}}
</style>
</head>
<body>

<header>
<h1>Analisi della Distribuzione delle Velocit&agrave;</h1>
<h2>Via Cristoforo Colombo &mdash; Roma</h2>
<p>Dati TomTom Speed Profiles &bull; Periodo: Inverno feb. 2026 / Estate lug. 2025</p>
</header>

<nav>
<h3>Indice</h3>
<ol>
<li><a href="#intro">Introduzione e Metodologia</a></li>
<li><a href="#carriage_centrale">Carreggiata Centrale</a></li>
<li><a href="#carriage_lat_ostia">Carreggiata Laterale Dir. Ostia</a></li>
<li><a href="#carriage_lat_centro">Carreggiata Laterale Dir. Centro</a></li>
<li><a href="#conclusions">Conclusioni e Raccomandazioni</a></li>
</ol>
</nav>

<!-- 1 INTRODUCTION -->
<section id="intro">
<h2>1. Introduzione e Metodologia</h2>
<p>Il presente report analizza la distribuzione delle velocit&agrave; veicolari lungo
<strong>Via Cristoforo Colombo</strong> a Roma, utilizzando i dati TomTom Speed Profiles.</p>

<p>L&rsquo;arteria &egrave; composta da tre carreggiate distinte:</p>
<ul>
<li><strong>Carreggiata Centrale</strong> (Dir. Ostia e Dir. Centro) &mdash; ~16 km,
    limiti 50/60/80 km/h</li>
<li><strong>Laterale Dir. Ostia</strong> (Tratto Pontina-Mezzocammino + Tratto Via di Acilia-Piazzale)
    &mdash; ~13 km, limiti 30/50 km/h</li>
<li><strong>Laterale Dir. Centro</strong> (Tratto Piazzale-Acilia + Svincolo Malafede +
    Tratto Alfonsine-Brandeliero) &mdash; ~11 km, limiti 30/50 km/h</li>
</ul>

<p>I dati coprono <strong>4 periodi</strong>:</p>
<ul>
<li><strong>Invernale Feriale</strong>: 1&ndash;15 febbraio 2026 (lun&ndash;ven)</li>
<li><strong>Estiva Feriale</strong>: 15&ndash;31 luglio 2025 (lun&ndash;ven)</li>
<li><strong>Estiva Festivo</strong>: 15&ndash;31 luglio 2025 (sab&ndash;dom)</li>
<li><strong>Invernale Festiva</strong>: 1&ndash;15 febbraio 2026 (sab&ndash;dom)</li>
</ul>

<div class="method-box">
<strong>Metriche principali:</strong>
<ul>
<li><strong>Velocit&agrave; media armonica</strong> &mdash; media ponderata per lunghezza segmento</li>
<li><strong>V85 (85&deg; percentile)</strong> &mdash; velocit&agrave; non superata dall&rsquo;85% dei veicoli</li>
<li><strong>Deviazione standard</strong> &mdash; dispersione delle velocit&agrave; individuali</li>
<li><strong>Tasso di superamento</strong> &mdash; % dei km con velocit&agrave; superiore al limite
    specifico del segmento</li>
</ul>
</div>

<p><strong>Nota:</strong> i limiti di velocit&agrave; variano per segmento (30, 40, 50, 60, 80 km/h).
Il superamento &egrave; calcolato rispetto al limite specifico di ciascun tratto.</p>
</section>

{carriage_sections}

<!-- CONCLUSIONS -->
<section id="conclusions">
<h2>5. Conclusioni e Raccomandazioni</h2>
<div class="insight-box">
<strong>Risultati Principali:</strong>
<ul>
<li>L&rsquo;analisi copre ~57 km di Via Cristoforo Colombo su 3 carreggiate</li>
<li>La carreggiata centrale presenta velocit&agrave; significativamente superiori
    alle laterali, con limiti fino a 80 km/h</li>
<li>Le carreggiate laterali hanno limiti pi&ugrave; bassi (30-50 km/h) ma presentano
    comunque superamenti frequenti</li>
<li>Le velocit&agrave; estive e invernali mostrano profili diversi</li>
</ul>
</div>
<div class="warning-box">
<strong>Aree di Attenzione:</strong>
<ul>
<li>I tratti con elevata variabilit&agrave; richiedono attenzione per la sicurezza</li>
<li>Le velocit&agrave; notturne sono significativamente pi&ugrave; elevate</li>
<li>Le laterali presentano punti critici con V85 ben oltre il limite</li>
</ul>
</div>
<p><em>Report generato automaticamente dai dati TomTom Speed Profiles.
Si raccomanda un&rsquo;interpretazione contestualizzata da parte di tecnici qualificati.</em></p>
</section>

<footer>
<p>TomTom Speed Profiles &mdash; Via Cristoforo Colombo, Roma</p>
</footer>
</body></html>"""
    return html


# ================================================================
# SECTION 8 — WORD REPORT
# ================================================================

def _b64_to_stream(b64_str):
    return io.BytesIO(base64.b64decode(b64_str))


def _add_table_to_doc(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(8)
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = str(row_data[col])
            for par in row_cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.size = Pt(8)


def _add_method_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)


def _add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_carriage_to_doc(doc, carriage, charts, tables, static_map_paths,
                          section_num, img_width):
    """Add a carriage group section to the Word document."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]
    slug = carriage.lower().replace(". ", "_").replace(" ", "_")
    cdata = charts[carriage]

    doc.add_heading(f"{section_num}. {cg['label']}", level=1)

    # Overview
    doc.add_heading(f"{section_num}.1 Panoramica", level=2)
    overview_key = f"overview_{slug}"
    if overview_key in tables:
        _add_table_to_doc(doc, tables[overview_key])

    # Exceedance
    doc.add_heading(f"{section_num}.2 Superamento Limiti", level=2)
    _add_method_box(doc, "Il superamento è calcolato rispetto al limite specifico di ciascun segmento.")
    exceedance_key = f"exceedance_{slug}"
    if exceedance_key in tables:
        _add_table_to_doc(doc, tables[exceedance_key])

    # Temporal
    doc.add_heading(f"{section_num}.3 Profili Temporali", level=2)
    doc.add_picture(_b64_to_stream(cdata["speed_by_hour"]), width=img_width)
    _add_caption(doc, "Velocità media armonica per ora (4 periodi).")
    doc.add_picture(_b64_to_stream(cdata["v85_by_hour"]), width=img_width)
    _add_caption(doc, "V85 medio per ora (4 periodi).")

    # Heatmaps
    doc.add_heading(f"{section_num}.4 Mappe di Calore", level=2)
    for lab, b64 in cdata["heatmaps"].items():
        doc.add_picture(_b64_to_stream(b64), width=img_width)
        _add_caption(doc, f"Mappa di calore — {lab}")

    # Exceedance chart
    doc.add_heading(f"{section_num}.5 Superamento Limiti per Ora", level=2)
    doc.add_picture(_b64_to_stream(cdata["exceedance"]), width=img_width)

    # V85 spatial
    doc.add_heading(f"{section_num}.6 Profili Spaziali V85", level=2)
    for lab, b64 in cdata["v85_spatial"].items():
        doc.add_picture(_b64_to_stream(b64), width=img_width)
        _add_caption(doc, f"Profilo V85 — {lab}")

    # Top 10
    for d in dirs:
        d_slug = d.lower().replace(" ", "_").replace("-", "_")
        tkey = f"top_fast_{slug}_{d_slug}"
        if tkey in tables:
            doc.add_heading(f"Top 10 V85 — {cg['dir_labels'][d]}", level=3)
            _add_table_to_doc(doc, tables[tkey])

    # Variability
    doc.add_heading(f"{section_num}.7 Variabilità", level=2)
    doc.add_picture(_b64_to_stream(cdata["variability"]), width=img_width)

    # Night
    doc.add_heading(f"{section_num}.8 Analisi Notturna", level=2)
    doc.add_picture(_b64_to_stream(cdata["night"]), width=img_width)

    # Season comparison
    doc.add_heading(f"{section_num}.9 Confronto Invernale vs Estiva", level=2)
    doc.add_picture(_b64_to_stream(cdata["season_comparison"]), width=img_width)

    # Static maps
    doc.add_heading(f"{section_num}.10 Mappe Statiche", level=2)
    for sp in static_map_paths:
        if slug in sp:
            try:
                doc.add_picture(sp, width=img_width)
                _add_caption(doc, Path(sp).stem.replace("_", " ").title())
            except Exception:
                pass

    doc.add_page_break()


def generate_word_report(charts, tables, static_map_paths):
    """Generate Word (.docx) report."""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    img_width = Inches(6.5)

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)

    # Title page
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Analisi della Distribuzione delle Velocità")
    run.bold = True; run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Via Cristoforo Colombo — Roma")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x28, 0x35, 0x93)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Dati TomTom Speed Profiles")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # TOC
    doc.add_heading("Indice", level=1)
    toc_items = [
        "1. Introduzione e Metodologia",
        "2. Carreggiata Centrale",
        "3. Carreggiata Laterale Dir. Ostia",
        "4. Carreggiata Laterale Dir. Centro",
        "5. Conclusioni e Raccomandazioni",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading("1. Introduzione e Metodologia", level=1)
    doc.add_paragraph(
        "Il presente report analizza la distribuzione delle velocità veicolari lungo "
        "Via Cristoforo Colombo a Roma, utilizzando i dati TomTom Speed Profiles."
    )
    doc.add_paragraph("L'arteria è composta da tre carreggiate distinte:")
    for item in [
        "Carreggiata Centrale (Dir. Ostia e Dir. Centro) — ~16 km, limiti 50/60/80 km/h",
        "Laterale Dir. Ostia (2 tratti) — ~13 km, limiti 30/50 km/h",
        "Laterale Dir. Centro (3 tratti) — ~11 km, limiti 30/50 km/h",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph("I dati coprono 4 periodi:")
    for item in [
        "Invernale Feriale: 1–15 febbraio 2026 (lun–ven)",
        "Estiva Feriale: 15–31 luglio 2025 (lun–ven)",
        "Estiva Festivo: 15–31 luglio 2025 (sab–dom)",
        "Invernale Festiva: 1–15 febbraio 2026 (sab–dom)",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    _add_method_box(doc,
        "Metriche: velocità media armonica, V85 (85° percentile), deviazione standard, "
        "tasso di superamento (% km oltre il limite specifico del segmento)."
    )
    doc.add_page_break()

    # Carriage sections
    for i, carriage in enumerate(CARRIAGE_GROUPS, start=2):
        _add_carriage_to_doc(doc, carriage, charts, tables,
                              static_map_paths, i, img_width)

    # 5. Conclusions
    doc.add_heading("5. Conclusioni e Raccomandazioni", level=1)
    doc.add_paragraph(
        "L'analisi copre ~57 km di Via Cristoforo Colombo su 3 carreggiate. "
        "La carreggiata centrale presenta velocità superiori alle laterali. "
        "Le laterali, nonostante limiti più bassi, mostrano superamenti frequenti."
    )
    items = [
        "I tratti con elevata variabilità richiedono attenzione per la sicurezza",
        "Le velocità notturne sono significativamente più elevate",
        "Le velocità estive e invernali mostrano profili diversi",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    p = doc.add_paragraph()
    run = p.add_run(
        "Report generato automaticamente dai dati TomTom Speed Profiles. "
        "Si raccomanda un'interpretazione contestualizzata da parte di tecnici qualificati."
    )
    run.italic = True; run.font.size = Pt(9)

    docx_path = OUTPUT_DIR / "report_colombo.docx"
    doc.save(str(docx_path))
    return str(docx_path)


# ================================================================
# SECTION 9 — MAIN
# ================================================================

def main():
    print("Caricamento dati Via Cristoforo Colombo...")
    segments, summaries, headers = load_all_data()
    print(f"  Segmenti caricati: {len(segments)} righe, "
          f"{segments['seg_idx'].nunique()} segmenti unici")

    print("Calcolo statistiche per segmento...")
    seg_stats = segment_peak_stats(segments)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    MAPS_DIR.mkdir(parents=True, exist_ok=True)

    print("Generazione grafici...")
    charts = generate_all_charts(segments, summaries)
    for carriage in charts:
        print(f"  - {CARRIAGE_GROUPS[carriage]['label']}: "
              f"{len(charts[carriage])} gruppi di grafici")

    print("Generazione mappe statiche...")
    static_map_paths = generate_all_static_maps(segments)
    print(f"  - {len(static_map_paths)} mappe statiche in {STATIC_MAPS_DIR}")

    print("Generazione mappe interattive...")
    map_files = generate_maps(segments, seg_stats)
    print(f"  - {len(map_files)} mappe interattive in {MAPS_DIR}")

    print("Costruzione tabelle...")
    tables = build_summary_tables(segments, summaries, seg_stats)

    print("Assemblaggio report HTML...")
    html = generate_html_report(charts, map_files, tables)
    report_path = OUTPUT_DIR / "report_colombo.html"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"  - HTML: {report_path}")

    print("Assemblaggio report Word (.docx)...")
    docx_path = generate_word_report(charts, tables, static_map_paths)
    print(f"  - DOCX: {docx_path}")

    csv_path = OUTPUT_DIR / "dati_segmenti_colombo.csv"
    segments.drop(columns=["geometry"]).to_csv(csv_path, index=False)
    print(f"  - CSV: {csv_path}")
    print(f"\nReport HTML: {report_path}")
    print(f"Report DOCX: {docx_path}")
    print(f"Mappe:       {MAPS_DIR}")


if __name__ == "__main__":
    main()
