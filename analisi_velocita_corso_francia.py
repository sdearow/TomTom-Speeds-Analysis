#!/usr/bin/env python3
"""
Analisi della Distribuzione delle Velocità — Corso Francia, Roma
Dati TomTom Speed Profiles — Febbraio 2026
"""

# ================================================================
# SECTION 1 — IMPORTS & CONFIGURATION
# ================================================================

import json
import base64
import io
import warnings
from pathlib import Path

import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import folium
import branca.colormap as cm
import contextily as cx
from shapely.geometry import shape
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

warnings.filterwarnings("ignore")

BASE_DIR = Path(__file__).resolve().parent
DATA_FILES = {
    ("Feriali", "Centro"): BASE_DIR / "Corso Francia_Feriali" / "Corso Francia_Dir. Centro_1.geojson",
    ("Feriali", "GRA"):    BASE_DIR / "Corso Francia_Feriali" / "Corso Francia_Dir. GRA_2.geojson",
    ("Festivi", "Centro"): BASE_DIR / "Corso Francia_Festivi" / "Corso Francia_Dir.Centro_1.geojson",
    ("Festivi", "GRA"):    BASE_DIR / "Corso Francia_Festivi" / "Corso Francia_Dir. GRA_2.geojson",
}
OUTPUT_DIR = BASE_DIR / "output"
MAPS_DIR = OUTPUT_DIR / "maps"

SPEED_LIMIT = 50
NIGHT_HOURS = list(range(0, 6)) + [22, 23]
AM_PEAK = [7, 8]
PM_PEAK = [17, 18]
MIDDAY = [12, 13]
P85_IDX = 16
P95_IDX = 18

plt.rcParams.update({
    "figure.dpi": 140, "savefig.dpi": 140,
    "font.family": "sans-serif", "font.size": 10,
    "axes.titlesize": 12, "axes.labelsize": 10,
})
COLORS = {
    ("Feriali", "Centro"): "#1f77b4",
    ("Feriali", "GRA"):    "#ff7f0e",
    ("Festivi", "Centro"): "#2ca02c",
    ("Festivi", "GRA"):    "#d62728",
}
LABELS = {
    ("Feriali", "Centro"): "Feriali \u2192 Centro",
    ("Feriali", "GRA"):    "Feriali \u2192 GRA",
    ("Festivi", "Centro"): "Festivi \u2192 Centro",
    ("Festivi", "GRA"):    "Festivi \u2192 GRA",
}
HOUR_SHORT = [f"{h}" for h in range(24)]


# ================================================================
# SECTION 2 — DATA LOADING (unchanged)
# ================================================================

def load_geojson(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return json.load(f)


def parse_segments(data, day_type, direction):
    features = data["features"]
    header = features[0]["properties"]
    segments = features[1:]
    rows = []
    cumul = 0.0
    for idx, feat in enumerate(segments):
        props = feat["properties"]
        geom = shape(feat["geometry"])
        seg_dist = props["distance"]
        for tr in props["segmentTimeResults"]:
            hour = tr["timeSet"] - 2
            sp = tr["speedPercentiles"]
            rows.append({
                "day_type": day_type, "direction": direction, "seg_idx": idx,
                "segmentId": props["segmentId"], "streetName": props["streetName"],
                "frc": props["frc"], "speedLimit": props["speedLimit"],
                "seg_distance": seg_dist,
                "cum_dist_start": cumul, "cum_dist_mid": cumul + seg_dist / 2,
                "cum_dist_end": cumul + seg_dist,
                "hour": hour,
                "harm_avg_speed": tr["harmonicAverageSpeed"],
                "avg_speed": tr["averageSpeed"], "median_speed": tr["medianSpeed"],
                "std_speed": tr["standardDeviationSpeed"],
                "avg_tt": tr["averageTravelTime"], "median_tt": tr["medianTravelTime"],
                "std_tt": tr["travelTimeStandardDeviation"],
                "tt_ratio": tr["travelTimeRatio"],
                "sample_size": tr["sampleSize"], "norm_sample": tr["normalizedSampleSize"],
                "p5": sp[0], "p15": sp[2], "p25": sp[4], "p50": sp[9],
                "p75": sp[14], "p85": sp[16], "p90": sp[17], "p95": sp[18],
                "geometry": geom,
            })
        cumul += seg_dist

    sum_rows = []
    for s in header["summaries"]:
        hour = s["timeSet"] - 2
        ssp = s.get("speedPercentiles", [])
        sum_rows.append({
            "day_type": day_type, "direction": direction, "hour": hour,
            "route_dist": s["distance"],
            "harm_avg_speed": s["harmonicAverageSpeed"],
            "avg_tt": s["averageTravelTime"], "median_tt": s["medianTravelTime"],
            "pti": s["planningTimeIndex"],
            "route_p85": ssp[P85_IDX] if len(ssp) > P85_IDX else np.nan,
            "route_p95": ssp[P95_IDX] if len(ssp) > P95_IDX else np.nan,
        })
    return pd.DataFrame(rows), pd.DataFrame(sum_rows), header


def load_all_data():
    all_seg, all_sum = [], []
    headers = {}
    for (dt, dr), path in DATA_FILES.items():
        data = load_geojson(path)
        seg_df, sum_df, hdr = parse_segments(data, dt, dr)
        all_seg.append(seg_df)
        all_sum.append(sum_df)
        headers[(dt, dr)] = hdr
    return pd.concat(all_seg, ignore_index=True), pd.concat(all_sum, ignore_index=True), headers


# ================================================================
# SECTION 3 — ANALYSIS FUNCTIONS
# ================================================================

def compute_exceedance_by_km(df, speed_col, limit=SPEED_LIMIT):
    """% of ROUTE LENGTH (km) where speed_col > limit, weighted by seg_distance."""
    df = df.copy()
    df["exceed_dist"] = np.where(df[speed_col] > limit, df["seg_distance"], 0.0)
    grp = df.groupby(["day_type", "direction", "hour"]).agg(
        exceed_m=("exceed_dist", "sum"), total_m=("seg_distance", "sum"),
    ).reset_index()
    grp["pct_km_exceed"] = grp["exceed_m"] / grp["total_m"] * 100
    return grp[["day_type", "direction", "hour", "pct_km_exceed"]]


def hourly_means(df, value_col):
    return df.groupby(["day_type", "direction", "hour"])[value_col].mean().reset_index()


def segment_peak_stats(df):
    def _agg(sub):
        return pd.Series({
            "night_avg_speed": sub.loc[sub["hour"].isin(NIGHT_HOURS), "avg_speed"].mean(),
            "night_v85":       sub.loc[sub["hour"].isin(NIGHT_HOURS), "p85"].mean(),
            "am_avg_speed":    sub.loc[sub["hour"].isin(AM_PEAK), "avg_speed"].mean(),
            "am_v85":          sub.loc[sub["hour"].isin(AM_PEAK), "p85"].mean(),
            "pm_avg_speed":    sub.loc[sub["hour"].isin(PM_PEAK), "avg_speed"].mean(),
            "pm_v85":          sub.loc[sub["hour"].isin(PM_PEAK), "p85"].mean(),
            "all_avg_speed":   sub["avg_speed"].mean(),
            "all_v85":         sub["p85"].mean(),
            "all_std":         sub["std_speed"].mean(),
            "max_v85":         sub["p85"].max(),
            "max_avg_speed":   sub["avg_speed"].max(),
            "cum_dist_mid":    sub["cum_dist_mid"].iloc[0],
            "cum_dist_start":  sub["cum_dist_start"].iloc[0],
            "seg_distance":    sub["seg_distance"].iloc[0],
            "streetName":      sub["streetName"].iloc[0],
            "speedLimit":      sub["speedLimit"].iloc[0],
        })
    return df.groupby(["day_type", "direction", "seg_idx"]).apply(
        _agg, include_groups=False
    ).reset_index()


# ================================================================
# SECTION 4 — CHART GENERATION
# ================================================================

def fig_to_base64(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return base64.b64encode(buf.read()).decode("utf-8")


# -- Mod 4/5: split Feriali / Festivi into two panels --

def chart_speed_by_hour(summaries):
    """Two-panel: Feriali (left), Festivi (right)."""
    fig, axes = plt.subplots(1, 2, figsize=(14, 5), sharey=True)
    for ax, day_type in zip(axes, ["Feriali", "Festivi"]):
        for direction in ["Centro", "GRA"]:
            key = (day_type, direction)
            sub = summaries[(summaries["day_type"] == day_type)
                            & (summaries["direction"] == direction)].sort_values("hour")
            ax.plot(sub["hour"], sub["harm_avg_speed"], color=COLORS[key],
                    label=f"Dir. {direction}", linewidth=2, marker="o", markersize=3)
        ax.axhline(SPEED_LIMIT, color="red", ls="--", lw=1, alpha=.7,
                    label=f"Limite {SPEED_LIMIT} km/h")
        ax.set_xlabel("Ora")
        ax.set_title(f"Velocit\u00e0 Media Armonica \u2014 {day_type}")
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=7)
        ax.legend(fontsize=8); ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
    axes[0].set_ylabel("Velocit\u00e0 media armonica (km/h)")
    fig.suptitle("Velocit\u00e0 Media di Percorrenza per Ora", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_v85_by_hour(segments):
    """Two-panel V85: Feriali (left), Festivi (right)."""
    hmeans = hourly_means(segments, "p85")
    fig, axes = plt.subplots(1, 2, figsize=(14, 5), sharey=True)
    for ax, day_type in zip(axes, ["Feriali", "Festivi"]):
        for direction in ["Centro", "GRA"]:
            key = (day_type, direction)
            sub = hmeans[(hmeans["day_type"] == day_type)
                         & (hmeans["direction"] == direction)].sort_values("hour")
            ax.plot(sub["hour"], sub["p85"], color=COLORS[key],
                    label=f"Dir. {direction}", linewidth=2, marker="o", markersize=3)
        ax.axhline(SPEED_LIMIT, color="red", ls="--", lw=1, alpha=.7,
                    label=f"Limite {SPEED_LIMIT} km/h")
        ax.set_xlabel("Ora")
        ax.set_title(f"V85 Medio \u2014 {day_type}")
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=7)
        ax.legend(fontsize=8); ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
    axes[0].set_ylabel("V85 medio dei segmenti (km/h)")
    fig.suptitle("85\u00b0 Percentile della Velocit\u00e0 (V85) per Ora", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


# -- Mod 6: heatmaps with progressive distance (pcolormesh) --

def _make_heatmap_progressive(sub, value_col, title, cmap, norm, fig_ax=None):
    """Heatmap with progressive km on x-axis using pcolormesh."""
    seg_first = sub.drop_duplicates("seg_idx").sort_values("seg_idx")
    x_edges = np.concatenate([
        seg_first["cum_dist_start"].values,
        [seg_first["cum_dist_end"].values[-1]]
    ]) / 1000.0
    y_edges = np.arange(-0.5, 24.5, 1.0)

    pivot = sub.pivot_table(index="hour", columns="seg_idx",
                             values=value_col, aggfunc="mean")
    pivot = pivot.reindex(index=range(24), columns=seg_first["seg_idx"].values)

    if fig_ax is None:
        fig, ax = plt.subplots(figsize=(12, 7))
    else:
        fig, ax = fig_ax

    mesh = ax.pcolormesh(x_edges, y_edges, pivot.values,
                          cmap=cmap, norm=norm, shading="flat")
    ax.invert_yaxis()
    ax.set_xlabel("Progressiva (km)")
    ax.set_ylabel("Ora")
    ax.set_yticks(range(24)); ax.set_yticklabels(HOUR_SHORT, fontsize=7)
    max_km = x_edges[-1]
    tick_iv = 0.25 if max_km < 2.0 else 0.5
    xticks = np.arange(0, max_km + tick_iv, tick_iv)
    ax.set_xticks(xticks)
    ax.set_xticklabels([f"{x:.2f}" for x in xticks], fontsize=7)
    ax.set_title(title)
    fig.colorbar(mesh, ax=ax, shrink=0.75, label="km/h")
    return fig


def chart_heatmaps(segments):
    """Speed heatmaps: hour (y) x progressive distance (x)."""
    results = {}
    norm = mcolors.TwoSlopeNorm(vmin=15, vcenter=SPEED_LIMIT, vmax=90)
    cmap = plt.cm.RdYlGn
    for key in LABELS:
        sub = segments[(segments["day_type"] == key[0])
                       & (segments["direction"] == key[1])]
        title = f"Velocit\u00e0 Media per Progressiva e Ora \u2014 {LABELS[key]}"
        fig = _make_heatmap_progressive(sub, "avg_speed", title, cmap, norm)
        results[LABELS[key]] = fig_to_base64(fig)
    return results


# -- Mod 1/7: exceedance by km --

def chart_exceedance(segments):
    """% of route-km exceeding speed limit by hour."""
    exc_avg = compute_exceedance_by_km(segments, "avg_speed")
    exc_v85 = compute_exceedance_by_km(segments, "p85")

    fig, axes = plt.subplots(1, 2, figsize=(14, 5), sharey=True)
    for key in LABELS:
        sub = exc_avg[(exc_avg["day_type"] == key[0])
                      & (exc_avg["direction"] == key[1])].sort_values("hour")
        axes[0].plot(sub["hour"], sub["pct_km_exceed"], color=COLORS[key],
                     label=LABELS[key], linewidth=2, marker="o", markersize=3)
    axes[0].set_title("Km con Velocit\u00e0 Media > 50 km/h")
    axes[0].set_xlabel("Ora")
    axes[0].set_ylabel("% del percorso (km) oltre il limite")
    axes[0].set_xticks(range(24)); axes[0].set_xticklabels(HOUR_SHORT, fontsize=7)
    axes[0].legend(fontsize=7); axes[0].grid(True, alpha=.3)
    axes[0].set_xlim(-.5, 23.5); axes[0].set_ylim(0, 105)

    for key in LABELS:
        sub = exc_v85[(exc_v85["day_type"] == key[0])
                      & (exc_v85["direction"] == key[1])].sort_values("hour")
        axes[1].plot(sub["hour"], sub["pct_km_exceed"], color=COLORS[key],
                     label=LABELS[key], linewidth=2, marker="o", markersize=3)
    axes[1].set_title("Km con V85 > 50 km/h")
    axes[1].set_xlabel("Ora")
    axes[1].set_xticks(range(24)); axes[1].set_xticklabels(HOUR_SHORT, fontsize=7)
    axes[1].legend(fontsize=7); axes[1].grid(True, alpha=.3)
    axes[1].set_xlim(-.5, 23.5)

    fig.suptitle("Superamento del Limite (pesato per km)", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_v85_spatial(segments):
    """V85 along corridor at key time periods, one chart per direction."""
    results = {}
    periods = {
        "Notte (22\u201305)": NIGHT_HOURS,
        "Punta mattina (07\u201308)": AM_PEAK,
        "Mezzog. (12\u201313)": MIDDAY,
        "Punta sera (17\u201318)": PM_PEAK,
    }
    pcols = ["#7570b3", "#d95f02", "#1b9e77", "#e7298a"]
    for direction in ["Centro", "GRA"]:
        fig, ax = plt.subplots(figsize=(12, 5))
        sub = segments[(segments["day_type"] == "Feriali")
                       & (segments["direction"] == direction)]
        for (pname, phours), pcol in zip(periods.items(), pcols):
            psub = sub[sub["hour"].isin(phours)]
            prof = psub.groupby("seg_idx").agg(
                v85=("p85", "mean"), dist=("cum_dist_mid", "first"),
            ).sort_values("dist")
            ax.plot(prof["dist"] / 1000, prof["v85"], color=pcol,
                    label=pname, linewidth=1.8)
        ax.axhline(SPEED_LIMIT, color="red", ls="--", lw=1, alpha=.7)
        ax.set_xlabel("Progressiva (km)")
        ax.set_ylabel("V85 (km/h)")
        ax.set_title(f"Profilo Spaziale V85 \u2014 Feriali Dir. {direction}")
        ax.legend(fontsize=8); ax.grid(True, alpha=.3)
        results[direction] = fig_to_base64(fig)
    return results


# -- Mod 10: variability with progressive distance --

def chart_speed_variability(segments):
    """Std dev heatmaps (progressive km) + spatial profiles."""
    fig, axes = plt.subplots(2, 2, figsize=(14, 10))
    norm = mcolors.Normalize(vmin=0, vmax=25)
    cmap = plt.cm.YlOrRd
    for ci, direction in enumerate(["Centro", "GRA"]):
        sub = segments[(segments["day_type"] == "Feriali")
                       & (segments["direction"] == direction)]
        _make_heatmap_progressive(
            sub, "std_speed",
            f"Dev. Std. Velocit\u00e0 \u2014 Dir. {direction}",
            cmap, norm, fig_ax=(fig, axes[0, ci]))
        # spatial profile
        ax2 = axes[1, ci]
        prof = sub.groupby("seg_idx").agg(
            std_mean=("std_speed", "mean"), std_max=("std_speed", "max"),
            dist=("cum_dist_mid", "first"),
        ).sort_values("dist")
        dkm = prof["dist"] / 1000
        ax2.fill_between(dkm, 0, prof["std_max"], alpha=.2, color="orange",
                         label="Max orario")
        ax2.plot(dkm, prof["std_mean"], color="darkorange", lw=2,
                 label="Media giornaliera")
        ax2.set_xlabel("Progressiva (km)"); ax2.set_ylabel("Dev. Std. (km/h)")
        ax2.set_title(f"Variabilit\u00e0 Spaziale \u2014 Dir. {direction}")
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle("Variabilit\u00e0 delle Velocit\u00e0 (Dev. Std.) \u2014 Feriali",
                 fontsize=13, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


# -- Mod 11: night analysis separated by direction --

def chart_night_analysis(segments):
    """Night analysis: separate histograms and spatial profiles per direction."""
    fig, axes = plt.subplots(2, 2, figsize=(14, 10))
    fer = segments[segments["day_type"] == "Feriali"]
    for ci, direction in enumerate(["Centro", "GRA"]):
        sub_dir = fer[fer["direction"] == direction]
        night = sub_dir[sub_dir["hour"].isin(NIGHT_HOURS)]["avg_speed"]
        day = sub_dir[~sub_dir["hour"].isin(NIGHT_HOURS)]["avg_speed"]
        # histogram
        ax = axes[0, ci]
        ax.hist(day, bins=40, alpha=.6, color="#1f77b4",
                label="Diurno (06\u201321)", density=True)
        ax.hist(night, bins=40, alpha=.6, color="#9467bd",
                label="Notturno (22\u201305)", density=True)
        ax.axvline(SPEED_LIMIT, color="red", ls="--", lw=1,
                   label=f"Limite {SPEED_LIMIT}")
        ax.set_xlabel("Velocit\u00e0 media (km/h)"); ax.set_ylabel("Densit\u00e0")
        ax.set_title(f"Distribuzione Notte vs Giorno \u2014 Dir. {direction}")
        ax.legend(fontsize=8)
        # spatial profile
        ax2 = axes[1, ci]
        sub_n = sub_dir[sub_dir["hour"].isin(NIGHT_HOURS)]
        prof = sub_n.groupby("seg_idx").agg(
            v85=("p85", "mean"), avg=("avg_speed", "mean"),
            dist=("cum_dist_mid", "first"),
        ).sort_values("dist")
        dkm = prof["dist"] / 1000
        ax2.plot(dkm, prof["v85"], color="#9467bd", lw=2, label="V85 notturno")
        ax2.plot(dkm, prof["avg"], color="#1f77b4", lw=2, label="Vel. media notturna")
        ax2.axhline(SPEED_LIMIT, color="red", ls="--", lw=1, alpha=.7)
        ax2.set_xlabel("Progressiva (km)"); ax2.set_ylabel("Velocit\u00e0 (km/h)")
        ax2.set_title(f"Profilo Notturno \u2014 Dir. {direction}")
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle("Analisi delle Velocit\u00e0 Notturne (Feriali)", fontsize=13, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_weekday_weekend(segments):
    fig, axes = plt.subplots(2, 2, figsize=(14, 10))
    for ci, direction in enumerate(["Centro", "GRA"]):
        ax = axes[0, ci]
        for dt, ls in [("Feriali", "-"), ("Festivi", "--")]:
            sub = segments[(segments["day_type"] == dt)
                           & (segments["direction"] == direction)]
            hm = sub.groupby("hour")["avg_speed"].mean().reset_index().sort_values("hour")
            ax.plot(hm["hour"], hm["avg_speed"], color=COLORS[(dt, direction)],
                    ls=ls, lw=2, label=dt, marker="o", markersize=3)
        ax.axhline(SPEED_LIMIT, color="red", ls="--", lw=1, alpha=.7)
        ax.set_title(f"Velocit\u00e0 Media \u2014 Dir. {direction}")
        ax.set_xlabel("Ora"); ax.set_ylabel("km/h")
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=7)
        ax.legend(fontsize=8); ax.grid(True, alpha=.3)

        ax2 = axes[1, ci]
        for dt, ls in [("Feriali", "-"), ("Festivi", "--")]:
            sub = segments[(segments["day_type"] == dt)
                           & (segments["direction"] == direction)]
            hm = sub.groupby("hour")["p85"].mean().reset_index().sort_values("hour")
            ax2.plot(hm["hour"], hm["p85"], color=COLORS[(dt, direction)],
                     ls=ls, lw=2, label=dt, marker="o", markersize=3)
        ax2.axhline(SPEED_LIMIT, color="red", ls="--", lw=1, alpha=.7)
        ax2.set_title(f"V85 \u2014 Dir. {direction}")
        ax2.set_xlabel("Ora"); ax2.set_ylabel("km/h")
        ax2.set_xticks(range(24)); ax2.set_xticklabels(HOUR_SHORT, fontsize=7)
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle("Confronto Feriali vs Festivi", fontsize=13, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


# ================================================================
# SECTION 4b — STATIC MAP GENERATION
# ================================================================

STATIC_MAPS_DIR = OUTPUT_DIR / "static_maps"


def _build_speed_gdf(segments, day_type, direction, value_col,
                     hour_filter=None, agg_func="mean"):
    """Aggregate segment speeds and return a GeoDataFrame in EPSG:4326."""
    sub = segments[(segments["day_type"] == day_type)
                   & (segments["direction"] == direction)]
    if hour_filter is not None:
        sub = sub[sub["hour"].isin(hour_filter)]
    agg = sub.groupby("seg_idx").agg({
        value_col: agg_func, "geometry": "first",
        "cum_dist_start": "first", "cum_dist_mid": "first",
        "seg_distance": "first", "streetName": "first",
        "speedLimit": "first",
    }).reset_index()
    return gpd.GeoDataFrame(agg, geometry="geometry", crs="EPSG:4326")


def _add_basemap(ax, crs="EPSG:4326"):
    """Add contextily basemap tiles. Works with EPSG:4326 or EPSG:3857."""
    try:
        cx.add_basemap(ax, crs=crs, source=cx.providers.CartoDB.Positron,
                        zoom=16, attribution=False)
        return True
    except Exception:
        pass
    try:
        cx.add_basemap(ax, crs=crs, source=cx.providers.OpenStreetMap.Mapnik,
                        zoom=16, attribution=False)
        return True
    except Exception:
        pass
    # Offline fallback: cartographic styling
    ax.set_facecolor("#eef0f4")
    ax.grid(True, alpha=0.25, linestyle="--", color="#888", zorder=0)
    ax.tick_params(labelsize=7, colors="#555")
    ax.ticklabel_format(useOffset=False, style="plain")
    for spine in ax.spines.values():
        spine.set_edgecolor("#999")
        spine.set_linewidth(0.8)
    return False


def _annotate_segments(ax, gdf, value_col):
    """Add progressive markers, start/end markers, and stats to a map axis."""
    total_dist = gdf["cum_dist_start"].max() + gdf["seg_distance"].max()
    interval = 500 if total_dist > 1500 else 250
    for target_m in range(0, int(total_dist) + interval, interval):
        candidates = gdf.loc[(gdf["cum_dist_start"] <= target_m)
                              & (gdf["cum_dist_start"] + gdf["seg_distance"] >= target_m)]
        if candidates.empty:
            continue
        row = candidates.iloc[0]
        d = row["seg_distance"]
        frac = (target_m - row["cum_dist_start"]) / d if d > 0 else 0.0
        frac = max(0.0, min(1.0, frac))
        pt = row["geometry"].interpolate(frac, normalized=True)
        label = f"{target_m} m" if target_m < 1000 else f"{target_m/1000:.1f} km"
        ax.annotate(label, xy=(pt.x, pt.y), fontsize=7, fontweight="bold",
                    color="#1a237e", ha="left", va="bottom",
                    xytext=(5, 5), textcoords="offset points",
                    bbox=dict(boxstyle="round,pad=0.15", fc="white",
                              ec="#1a237e", alpha=0.85, lw=0.5),
                    zorder=5)
        ax.plot(pt.x, pt.y, "o", color="#1a237e", markersize=4, zorder=5)

    start_pt = list(gdf.iloc[0]["geometry"].coords)[0]
    end_pt = list(gdf.iloc[-1]["geometry"].coords)[-1]
    ax.plot(start_pt[0], start_pt[1], "^", color="green", markersize=10,
            zorder=6, label="Inizio")
    ax.plot(end_pt[0], end_pt[1], "v", color="red", markersize=10,
            zorder=6, label="Fine")

    vals = gdf[value_col]
    metric_label = "V85" if "p85" in value_col or "max_p85" in value_col else "Vel. media"
    stats_text = (f"{metric_label}: media={vals.mean():.1f}, "
                  f"max={vals.max():.1f}, min={vals.min():.1f} km/h")
    ax.text(0.02, 0.02, stats_text, transform=ax.transAxes, fontsize=8,
            va="bottom", ha="left",
            bbox=dict(boxstyle="round,pad=0.3", fc="white", ec="gray",
                      alpha=0.9))
    ax.legend(loc="upper right", fontsize=8)


def generate_static_speed_map(segments, day_type, direction, value_col,
                               hour_filter=None, title="", fname="map.png",
                               agg_func="mean"):
    """Render a static matplotlib map with coloured segments on a basemap."""
    gdf = _build_speed_gdf(segments, day_type, direction, value_col,
                            hour_filter, agg_func=agg_func)

    fig, ax = plt.subplots(figsize=(10, 12))

    vmin, vmax = 20, 80
    norm = mcolors.TwoSlopeNorm(vmin=vmin, vcenter=SPEED_LIMIT, vmax=vmax)
    cmap = plt.cm.RdYlGn_r

    # Road-bed shadow + coloured speed data in EPSG:4326
    gdf.plot(ax=ax, color="#d0d0d0", linewidth=9, zorder=1, alpha=0.8)
    gdf.plot(ax=ax, column=value_col, cmap=cmap, norm=norm,
             linewidth=5, legend=False, zorder=2)

    # Add basemap tiles (contextily handles reprojection internally)
    has_tiles = _add_basemap(ax, crs="EPSG:4326")

    if has_tiles:
        ax.set_axis_off()
    else:
        ax.set_xlabel("Longitudine", fontsize=9)
        ax.set_ylabel("Latitudine", fontsize=9)

    # Progressive annotations
    _annotate_segments(ax, gdf, value_col)

    # Colour bar
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=ax, shrink=0.5, pad=0.03, aspect=30)
    cbar.set_label("km/h", fontsize=10)
    cbar.ax.axhline(y=SPEED_LIMIT, color="black", linewidth=1.5, linestyle="--")
    cbar.ax.text(1.3, SPEED_LIMIT, f"Limite {SPEED_LIMIT}",
                 transform=cbar.ax.get_yaxis_transform(),
                 va="center", fontsize=8, color="black")

    ax.set_title(title, fontsize=13, fontweight="bold", pad=12)
    fig.tight_layout()

    out_path = STATIC_MAPS_DIR / fname
    fig.savefig(str(out_path), dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(out_path)


def generate_static_speed_map_sidebyside(segments, day_type, value_col,
                                          hour_filter=None, title="",
                                          fname="map_sidebyside.png",
                                          agg_func="mean"):
    """Side-by-side static map: Dir. Centro (left) and Dir. GRA (right)."""
    gdf_centro = _build_speed_gdf(segments, day_type, "Centro", value_col,
                                   hour_filter, agg_func=agg_func)
    gdf_gra = _build_speed_gdf(segments, day_type, "GRA", value_col,
                                hour_filter, agg_func=agg_func)

    fig, axes = plt.subplots(1, 2, figsize=(20, 12))

    vmin, vmax = 20, 80
    norm = mcolors.TwoSlopeNorm(vmin=vmin, vcenter=SPEED_LIMIT, vmax=vmax)
    cmap = plt.cm.RdYlGn_r

    for ax, gdf_orig, dir_label in [
        (axes[0], gdf_centro, "Centro"),
        (axes[1], gdf_gra, "GRA"),
    ]:
        gdf_orig.plot(ax=ax, color="#d0d0d0", linewidth=9, zorder=1, alpha=0.8)
        gdf_orig.plot(ax=ax, column=value_col, cmap=cmap, norm=norm,
                      linewidth=5, legend=False, zorder=2)
        has_tiles = _add_basemap(ax, crs="EPSG:4326")
        _annotate_segments(ax, gdf_orig, value_col)
        if has_tiles:
            ax.set_axis_off()
        ax.set_title(f"Dir. {dir_label}", fontsize=12, fontweight="bold")

    # Shared colour bar
    sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm)
    sm.set_array([])
    cbar = fig.colorbar(sm, ax=axes.tolist(), shrink=0.5, pad=0.02, aspect=30)
    cbar.set_label("km/h", fontsize=10)
    cbar.ax.axhline(y=SPEED_LIMIT, color="black", linewidth=1.5, linestyle="--")
    cbar.ax.text(1.3, SPEED_LIMIT, f"Limite {SPEED_LIMIT}",
                 transform=cbar.ax.get_yaxis_transform(),
                 va="center", fontsize=8, color="black")

    fig.suptitle(title, fontsize=14, fontweight="bold", y=1.01)
    fig.tight_layout()

    out_path = STATIC_MAPS_DIR / fname
    fig.savefig(str(out_path), dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return str(out_path)


def generate_all_static_maps(segments):
    """Generate static speed maps: individual, side-by-side, and max V85."""
    STATIC_MAPS_DIR.mkdir(parents=True, exist_ok=True)
    configs = [
        # (value_col, hour_filter, label_metric, file_prefix, agg_func)
        ("avg_speed", None,        "Velocit\u00e0 Media (24h)",      "avg_speed_24h", "mean"),
        ("p85",       None,        "V85 (24h)",                      "v85_24h",       "mean"),
        ("avg_speed", NIGHT_HOURS, "Velocit\u00e0 Media Notturna",   "avg_speed_night", "mean"),
        ("p85",       NIGHT_HOURS, "V85 Notturno",                   "v85_night",     "mean"),
    ]
    paths = []

    # --- Individual maps (one per direction) ---
    for direction in ["Centro", "GRA"]:
        for vcol, hfilter, label, prefix, agg in configs:
            period = "22\u201305" if hfilter is not None else "tutte le ore"
            title = f"{label} \u2014 Feriali Dir. {direction} ({period})"
            fname = f"{prefix}_feriali_{direction.lower()}.png"
            p = generate_static_speed_map(
                segments, "Feriali", direction, vcol,
                hour_filter=hfilter, title=title, fname=fname, agg_func=agg)
            paths.append(p)
            print(f"  - {fname}")

    # --- Side-by-side maps (both directions) ---
    for vcol, hfilter, label, prefix, agg in configs:
        period = "22\u201305" if hfilter is not None else "tutte le ore"
        title = f"{label} \u2014 Feriali ({period})"
        fname = f"{prefix}_feriali_sidebyside.png"
        p = generate_static_speed_map_sidebyside(
            segments, "Feriali", vcol,
            hour_filter=hfilter, title=title, fname=fname, agg_func=agg)
        paths.append(p)
        print(f"  - {fname}")

    # --- Max V85 maps (max across all hours per segment) ---
    for direction in ["Centro", "GRA"]:
        title = f"V85 Massimo (ora peggiore) \u2014 Feriali Dir. {direction}"
        fname = f"max_v85_feriali_{direction.lower()}.png"
        p = generate_static_speed_map(
            segments, "Feriali", direction, "p85",
            hour_filter=None, title=title, fname=fname, agg_func="max")
        paths.append(p)
        print(f"  - {fname}")

    # --- Max V85 side-by-side ---
    title = "V85 Massimo (ora peggiore) \u2014 Feriali"
    fname = "max_v85_feriali_sidebyside.png"
    p = generate_static_speed_map_sidebyside(
        segments, "Feriali", "p85",
        hour_filter=None, title=title, fname=fname, agg_func="max")
    paths.append(p)
    print(f"  - {fname}")

    return paths


# ================================================================
# SECTION 5 — INTERACTIVE MAP GENERATION
# ================================================================

def find_point_at_distance(seg_info, target_dist):
    """Find geographic (lat, lon) at a given cumulative distance."""
    for _, row in seg_info.iterrows():
        if target_dist <= row["cum_dist_end"] + 0.01:
            d = row["seg_distance"]
            frac = (target_dist - row["cum_dist_start"]) / d if d > 0 else 0.0
            frac = max(0.0, min(1.0, frac))
            pt = row["geometry"].interpolate(frac, normalized=True)
            return pt.y, pt.x
    last_c = list(seg_info.iloc[-1]["geometry"].coords)[-1]
    return last_c[1], last_c[0]


def _seg_geodataframe(segments, day_type, direction, value_col, agg="mean"):
    sub = segments[(segments["day_type"] == day_type)
                   & (segments["direction"] == direction)]
    agg_dict = {
        value_col: agg, "geometry": "first",
        "cum_dist_mid": "first", "cum_dist_start": "first",
        "streetName": "first", "speedLimit": "first", "seg_distance": "first",
        "avg_speed": "mean", "p85": "mean", "std_speed": "mean",
    }
    grouped = sub.groupby("seg_idx").agg(agg_dict).reset_index()
    return gpd.GeoDataFrame(grouped, geometry="geometry", crs="EPSG:4326")


def make_folium_map(gdf, value_col, title, vmin=None, vmax=None,
                     reverse_cmap=False):
    """Folium map with THICK coloured segments (weight=10)."""
    clat = gdf.geometry.centroid.y.mean()
    clon = gdf.geometry.centroid.x.mean()
    m = folium.Map(location=[clat, clon], zoom_start=15,
                   tiles="CartoDB positron")
    if vmin is None:
        vmin = gdf[value_col].min()
    if vmax is None:
        vmax = gdf[value_col].max()
    colors = (["green", "yellow", "red"] if not reverse_cmap
              else ["red", "yellow", "green"])
    cmap = cm.LinearColormap(colors=colors, vmin=vmin, vmax=vmax,
                              caption=f"{title} (km/h)")
    for _, row in gdf.iterrows():
        coords = [[c[1], c[0]] for c in row.geometry.coords]
        popup_html = (
            f"<b>{row['streetName']}</b><br>"
            f"Progr.: {row['cum_dist_start']:.0f} m<br>"
            f"Vel. media: {row['avg_speed']:.1f} km/h<br>"
            f"V85: {row['p85']:.1f} km/h<br>"
            f"Dev. std: {row['std_speed']:.1f} km/h<br>"
            f"Limite: {row['speedLimit']} km/h<br>"
            f"Lungh.: {row['seg_distance']:.0f} m"
        )
        folium.PolyLine(coords, weight=10, color=cmap(row[value_col]),
                        opacity=0.9,
                        popup=folium.Popup(popup_html, max_width=280)).add_to(m)
    cmap.add_to(m)
    return m


def create_progressive_map(segments, day_type, direction, interval_m=250):
    """Folium map with progressive distance markers every interval_m."""
    sub = segments[(segments["day_type"] == day_type)
                   & (segments["direction"] == direction)]
    si = sub.drop_duplicates("seg_idx").sort_values("seg_idx")
    clat = si.geometry.apply(lambda g: g.centroid.y).mean()
    clon = si.geometry.apply(lambda g: g.centroid.x).mean()
    m = folium.Map(location=[clat, clon], zoom_start=15,
                   tiles="CartoDB positron")
    for _, row in si.iterrows():
        coords = [[c[1], c[0]] for c in row["geometry"].coords]
        folium.PolyLine(coords, weight=6, color="#1a237e", opacity=.8).add_to(m)
    total_dist = si["cum_dist_end"].max()
    for td in range(0, int(total_dist) + interval_m, interval_m):
        if td > total_dist:
            break
        lat, lon = find_point_at_distance(si, td)
        label = f"{td} m" if td < 1000 else f"{td/1000:.2f} km"
        folium.Marker(
            [lat, lon],
            icon=folium.DivIcon(
                html=(f'<div style="font-size:10px;font-weight:bold;background:white;'
                      f'padding:2px 5px;border:1px solid #333;border-radius:3px;'
                      f'white-space:nowrap;box-shadow:1px 1px 3px rgba(0,0,0,.3);">'
                      f'{label}</div>'),
                icon_size=(70, 22), icon_anchor=(35, 11)),
        ).add_to(m)
    sc = list(si.iloc[0]["geometry"].coords)[0]
    ec = list(si.iloc[-1]["geometry"].coords)[-1]
    folium.Marker([sc[1], sc[0]], icon=folium.Icon(color="green"),
                  popup="Inizio (0 m)").add_to(m)
    folium.Marker([ec[1], ec[0]], icon=folium.Icon(color="red"),
                  popup=f"Fine ({total_dist:.0f} m)").add_to(m)
    return m


def create_top_segments_map(segments, seg_stats, day_type, direction, n=10):
    """Folium map highlighting top-N segments by max V85."""
    sub = segments[(segments["day_type"] == day_type)
                   & (segments["direction"] == direction)]
    si = sub.drop_duplicates("seg_idx").sort_values("seg_idx")
    top = seg_stats[
        (seg_stats["day_type"] == day_type)
        & (seg_stats["direction"] == direction)
    ].nlargest(n, "max_v85")
    top_set = set(top["seg_idx"].values)

    clat = si.geometry.apply(lambda g: g.centroid.y).mean()
    clon = si.geometry.apply(lambda g: g.centroid.x).mean()
    m = folium.Map(location=[clat, clon], zoom_start=15,
                   tiles="CartoDB positron")
    for _, row in si.iterrows():
        coords = [[c[1], c[0]] for c in row["geometry"].coords]
        is_top = row["seg_idx"] in top_set
        folium.PolyLine(
            coords,
            weight=10 if is_top else 4,
            color="#e41a1c" if is_top else "#aaaaaa",
            opacity=0.9 if is_top else 0.4,
        ).add_to(m)
    for rank, (_, row) in enumerate(top.iterrows(), 1):
        match = si[si["seg_idx"] == row["seg_idx"]]
        if match.empty:
            continue
        c = match.iloc[0]["geometry"].centroid
        folium.Marker(
            [c.y, c.x],
            icon=folium.DivIcon(
                html=(f'<div style="font-size:11px;background:#e41a1c;color:white;'
                      f'padding:2px 6px;border-radius:50%;text-align:center;'
                      f'width:24px;height:24px;line-height:20px;font-weight:bold;'
                      f'box-shadow:1px 1px 3px rgba(0,0,0,.4);">{rank}</div>'),
                icon_size=(24, 24), icon_anchor=(12, 12)),
            popup=(f"<b>#{rank}</b><br>Progr.: {row['cum_dist_start']:.0f} m<br>"
                   f"V85 max: {row['max_v85']:.0f} km/h<br>"
                   f"Vel. media: {row['all_avg_speed']:.1f} km/h"),
        ).add_to(m)
    return m


def generate_maps(segments, seg_stats):
    MAPS_DIR.mkdir(parents=True, exist_ok=True)
    map_files = {}
    configs = [
        ("v85_feriali_centro.html", "Feriali", "Centro", "p85", "mean",
         "V85 Medio \u2014 Feriali Dir. Centro", 20, 80, False),
        ("v85_feriali_gra.html", "Feriali", "GRA", "p85", "mean",
         "V85 Medio \u2014 Feriali Dir. GRA", 20, 80, False),
        ("v85_festivi_centro.html", "Festivi", "Centro", "p85", "mean",
         "V85 Medio \u2014 Festivi Dir. Centro", 20, 80, False),
        ("v85_festivi_gra.html", "Festivi", "GRA", "p85", "mean",
         "V85 Medio \u2014 Festivi Dir. GRA", 20, 80, False),
        ("std_feriali_centro.html", "Feriali", "Centro", "std_speed", "mean",
         "Variabilit\u00e0 \u2014 Feriali Dir. Centro", 5, 22, True),
        ("std_feriali_gra.html", "Feriali", "GRA", "std_speed", "mean",
         "Variabilit\u00e0 \u2014 Feriali Dir. GRA", 5, 22, True),
        ("night_v85_feriali_centro.html", "Feriali", "Centro", "p85", "mean",
         "V85 Notturno \u2014 Feriali Dir. Centro", 30, 90, False),
        ("night_v85_feriali_gra.html", "Feriali", "GRA", "p85", "mean",
         "V85 Notturno \u2014 Feriali Dir. GRA", 30, 90, False),
    ]
    for cfg in configs:
        fname, dt, dr, vcol, agg, title, vmin, vmax, rev = cfg
        sub = (segments[segments["hour"].isin(NIGHT_HOURS)]
               if "night" in fname else segments)
        gdf = _seg_geodataframe(sub, dt, dr, vcol, agg)
        m = make_folium_map(gdf, vcol, title, vmin, vmax, rev)
        m.save(str(MAPS_DIR / fname))
        map_files[fname] = fname

    # Progressive distance maps (one per direction)
    for direction in ["Centro", "GRA"]:
        m = create_progressive_map(segments, "Feriali", direction)
        fname = f"progressive_{direction.lower()}.html"
        m.save(str(MAPS_DIR / fname))
        map_files[fname] = fname

    # Top-10 V85 maps
    for direction in ["Centro", "GRA"]:
        m = create_top_segments_map(segments, seg_stats, "Feriali", direction)
        fname = f"top10_v85_feriali_{direction.lower()}.html"
        m.save(str(MAPS_DIR / fname))
        map_files[fname] = fname

    # Max V85 maps (maximum V85 across all hours per segment)
    for direction in ["Centro", "GRA"]:
        gdf = _seg_geodataframe(segments, "Feriali", direction, "p85", agg="max")
        m = make_folium_map(gdf, "p85",
                            f"V85 Massimo (ora peggiore) \u2014 Feriali Dir. {direction}",
                            vmin=20, vmax=90, reverse_cmap=False)
        fname = f"max_v85_feriali_{direction.lower()}.html"
        m.save(str(MAPS_DIR / fname))
        map_files[fname] = fname

    return map_files


# ================================================================
# SECTION 6 — HTML REPORT
# ================================================================

def build_summary_tables(segments, summaries, seg_stats):
    tables = {}

    # Overview — removed Distanza and N. Segmenti (mod 2)
    overview_rows = []
    for key in LABELS:
        dt, dr = key
        ss = summaries[(summaries["day_type"] == dt)
                       & (summaries["direction"] == dr)]
        overview_rows.append({
            "Percorso": LABELS[key],
            "Vel. Media 24h (km/h)": f"{ss['harm_avg_speed'].mean():.1f}",
            "Vel. Punta AM (km/h)":  f"{ss[ss['hour'].isin(AM_PEAK)]['harm_avg_speed'].mean():.1f}",
            "Vel. Punta PM (km/h)":  f"{ss[ss['hour'].isin(PM_PEAK)]['harm_avg_speed'].mean():.1f}",
            "Vel. Notturna (km/h)":  f"{ss[ss['hour'].isin(NIGHT_HOURS)]['harm_avg_speed'].mean():.1f}",
        })
    tables["overview"] = pd.DataFrame(overview_rows)

    # Exceedance — weighted by km (mod 3)
    exc_rows = []
    for key in LABELS:
        dt, dr = key
        sub = segments[(segments["day_type"] == dt)
                       & (segments["direction"] == dr)]
        total_m = sub.drop_duplicates("seg_idx")["seg_distance"].sum()

        def _pct(col):
            vals = []
            for h in range(24):
                hs = sub[sub["hour"] == h]
                vals.append(hs.loc[hs[col] > SPEED_LIMIT, "seg_distance"].sum()
                            / total_m * 100)
            return np.mean(vals)

        exc_rows.append({
            "Percorso": LABELS[key],
            "% km Vel.Media > 50 (media 24h)": f"{_pct('avg_speed'):.1f}%",
            "% km V85 > 50 (media 24h)":       f"{_pct('p85'):.1f}%",
            "V85 massimo (km/h)":              f"{sub['p85'].max():.0f}",
            "Vel. Media max (km/h)":           f"{sub['avg_speed'].max():.1f}",
        })
    tables["exceedance"] = pd.DataFrame(exc_rows)

    # Top 10 — with progressive numbering, no Seg./Via (mod 9)
    for direction in ["Centro", "GRA"]:
        st = seg_stats[
            (seg_stats["day_type"] == "Feriali")
            & (seg_stats["direction"] == direction)
        ].nlargest(10, "max_v85").reset_index(drop=True)
        tbl = pd.DataFrame({
            "#": range(1, len(st) + 1),
            "Progressiva (m)": st["cum_dist_start"].apply(lambda x: f"{x:.0f}"),
            "V85 Max (km/h)":  st["max_v85"].apply(lambda x: f"{x:.0f}"),
            "Vel. Media (km/h)": st["all_avg_speed"].apply(lambda x: f"{x:.1f}"),
            "Dev. Std. (km/h)":  st["all_std"].apply(lambda x: f"{x:.1f}"),
            "Limite (km/h)":     st["speedLimit"].apply(lambda x: f"{x:.0f}"),
        })
        tables[f"top_fast_{direction.lower()}"] = tbl

    return tables


def df_to_html_table(df):
    return df.to_html(index=False, classes="data-table", border=0)


def _iframe(fname, title):
    return (f'<h3>{title}</h3>'
            f'<iframe src="maps/{fname}" width="100%" height="500" '
            f'frameborder="0" style="border:1px solid #ddd;border-radius:4px;'
            f'margin-bottom:20px;"></iframe>')


def generate_html_report(charts, map_files, tables):
    heatmap_imgs = "".join(
        f'<h3>{lab}</h3><img src="data:image/png;base64,{b64}" style="max-width:100%;">'
        for lab, b64 in charts["heatmaps"].items()
    )
    v85_spatial_imgs = "".join(
        f'<h3>Dir. {d}</h3><img src="data:image/png;base64,{b64}" style="max-width:100%;">'
        for d, b64 in charts["v85_spatial"].items()
    )
    progressive_iframes = (
        _iframe("progressive_centro.html", "Progressive \u2014 Dir. Centro")
        + _iframe("progressive_gra.html", "Progressive \u2014 Dir. GRA")
    )
    top10_centro_iframe = _iframe("top10_v85_feriali_centro.html",
                                   "Mappa Top 10 \u2014 Dir. Centro")
    top10_gra_iframe = _iframe("top10_v85_feriali_gra.html",
                                "Mappa Top 10 \u2014 Dir. GRA")
    main_map_iframes = "".join(
        _iframe(f, t) for f, t in [
            ("v85_feriali_centro.html", "V85 Medio \u2014 Feriali Dir. Centro"),
            ("v85_feriali_gra.html",    "V85 Medio \u2014 Feriali Dir. GRA"),
            ("v85_festivi_centro.html", "V85 Medio \u2014 Festivi Dir. Centro"),
            ("v85_festivi_gra.html",    "V85 Medio \u2014 Festivi Dir. GRA"),
            ("max_v85_feriali_centro.html", "V85 Massimo (ora peggiore) \u2014 Feriali Dir. Centro"),
            ("max_v85_feriali_gra.html",    "V85 Massimo (ora peggiore) \u2014 Feriali Dir. GRA"),
            ("std_feriali_centro.html", "Variabilit\u00e0 \u2014 Feriali Dir. Centro"),
            ("std_feriali_gra.html",    "Variabilit\u00e0 \u2014 Feriali Dir. GRA"),
            ("night_v85_feriali_centro.html", "V85 Notturno \u2014 Feriali Dir. Centro"),
            ("night_v85_feriali_gra.html",    "V85 Notturno \u2014 Feriali Dir. GRA"),
        ]
    )

    html = f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Analisi Velocit\u00e0 \u2014 Corso Francia, Roma</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',Tahoma,Geneva,Verdana,sans-serif;line-height:1.6;
  color:#333;background:#f9f9f9;max-width:1200px;margin:0 auto;padding:20px}}
header{{background:linear-gradient(135deg,#1a237e,#283593);color:#fff;
  padding:40px;border-radius:8px;margin-bottom:30px;text-align:center}}
header h1{{font-size:2em;margin-bottom:8px}}
header h2{{font-size:1.3em;font-weight:300;opacity:.9}}
header p{{margin-top:12px;opacity:.8;font-size:.95em}}
nav{{background:#fff;padding:20px 30px;border-radius:8px;margin-bottom:30px;
  box-shadow:0 2px 4px rgba(0,0,0,.1)}}
nav h3{{margin-bottom:10px;color:#1a237e}}
nav ol{{padding-left:20px}} nav li{{margin-bottom:4px}}
nav a{{color:#1565c0;text-decoration:none}} nav a:hover{{text-decoration:underline}}
section{{background:#fff;padding:30px;border-radius:8px;margin-bottom:20px;
  box-shadow:0 2px 4px rgba(0,0,0,.1)}}
h2{{color:#1a237e;border-bottom:2px solid #e8eaf6;padding-bottom:8px;margin-bottom:20px}}
h3{{color:#283593;margin:20px 0 10px 0}}
img{{max-width:100%;height:auto;border-radius:4px;margin:10px 0}}
.data-table{{width:100%;border-collapse:collapse;margin:15px 0;font-size:.9em}}
.data-table th{{background:#1a237e;color:#fff;padding:10px 12px;text-align:left;font-weight:600}}
.data-table td{{padding:8px 12px;border-bottom:1px solid #e0e0e0}}
.data-table tr:nth-child(even){{background:#f5f5f5}}
.data-table tr:hover{{background:#e8eaf6}}
.insight-box{{background:#e8f5e9;border-left:4px solid #4caf50;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}}
.warning-box{{background:#fff3e0;border-left:4px solid #ff9800;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}}
.method-box{{background:#e3f2fd;border-left:4px solid #1976d2;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}}
.chart-caption{{text-align:center;font-style:italic;color:#666;font-size:.9em;
  margin-top:-5px;margin-bottom:20px}}
footer{{text-align:center;padding:20px;color:#999;font-size:.85em}}
</style>
</head>
<body>

<header>
<h1>Analisi della Distribuzione delle Velocit&agrave;</h1>
<h2>Corso Francia &mdash; Roma</h2>
<p>Dati TomTom Speed Profiles &bull; Periodo: 1&ndash;15 Febbraio 2026</p>
</header>

<nav>
<h3>Indice</h3>
<ol>
<li><a href="#intro">Introduzione e Metodologia</a></li>
<li><a href="#overview">Panoramica dei Dati</a></li>
<li><a href="#temporal">Profili Temporali di Velocit&agrave;</a></li>
<li><a href="#exceedance">Superamento del Limite di Velocit&agrave;</a></li>
<li><a href="#v85">Analisi V85 (85&deg; Percentile)</a></li>
<li><a href="#variability">Variabilit&agrave; delle Velocit&agrave;</a></li>
<li><a href="#night">Analisi delle Velocit&agrave; Notturne</a></li>
<li><a href="#comparison">Confronto Feriali vs Festivi</a></li>
<li><a href="#maps">Mappe Interattive</a></li>
<li><a href="#conclusions">Conclusioni e Raccomandazioni</a></li>
</ol>
</nav>

<!-- 1 INTRODUCTION -->
<section id="intro">
<h2>1. Introduzione e Metodologia</h2>
<p>Il presente report analizza la distribuzione delle velocit&agrave; veicolari lungo
<strong>Corso di Francia</strong> a Roma, utilizzando i dati TomTom Speed Profiles
relativi al periodo <strong>1&ndash;15 febbraio 2026</strong>.</p>
<p>L&rsquo;analisi si basa su quattro dataset:</p>
<ul>
<li><strong>Giorni Feriali</strong> (lun&ndash;ven) &mdash; Direzione Centro e Direzione GRA</li>
<li><strong>Giorni Festivi</strong> (sab&ndash;dom) &mdash; Direzione Centro e Direzione GRA</li>
</ul>
<p>Metriche principali:</p>
<ul>
<li><strong>Velocit&agrave; media armonica</strong> &mdash; calcolata da TomTom come media armonica
    ponderata sulla lunghezza di tutti i segmenti, rappresentativa della velocit&agrave;
    effettiva di percorrenza.</li>
<li><strong>V85 (85&deg; percentile)</strong> &mdash; velocit&agrave; non superata dall&rsquo;85%
    dei veicoli. TomTom fornisce 19 percentili (dal 5&deg; al 95&deg;); il V85 corrisponde
    al 17&deg; valore dell&rsquo;array.</li>
<li><strong>Deviazione standard</strong> &mdash; dispersione delle velocit&agrave; individuali,
    fornita direttamente da TomTom per ogni segmento e ora.</li>
<li><strong>Tasso di superamento</strong> &mdash; percentuale dei <em>km di percorso</em>
    (non dei segmenti) con velocit&agrave; superiore al limite. Ogni segmento &egrave; pesato
    in proporzione alla propria lunghezza.</li>
</ul>
<p>Limite di velocit&agrave;: <strong>50 km/h</strong> (alcuni tratti terminali su Viale
Maresciallo Pilsudski: 40 km/h).</p>

<h3>Riferimento Spaziale: Progressive Chilometriche</h3>
<p>Tutti i grafici spaziali utilizzano la <em>progressiva chilometrica</em>: la distanza
(in km) dall&rsquo;inizio del percorso. Le due direzioni hanno punti di partenza diversi,
quindi le progressive sono specifiche per ciascuna direzione. Le mappe seguenti mostrano
la corrispondenza tra progressive e posizione geografica.</p>
{progressive_iframes}
</section>

<!-- 2 OVERVIEW -->
<section id="overview">
<h2>2. Panoramica dei Dati</h2>
<div class="method-box">
<strong>Nota metodologica:</strong> le velocit&agrave; sono le <em>medie armoniche a livello
di percorso</em> calcolate da TomTom. La media armonica pondera ogni segmento in base alla
propria lunghezza, risultando rappresentativa della velocit&agrave; effettiva di percorrenza.
I valori per le fasce orarie sono la media dei valori orari ricadenti nella fascia
(punta AM: 07&ndash;08, punta PM: 17&ndash;18, notturna: 22&ndash;05).
</div>
{df_to_html_table(tables['overview'])}

<h3>Riepilogo Superamento Limiti (pesato per km)</h3>
<div class="method-box">
<strong>Nota metodologica:</strong> per ogni ora si sommano le lunghezze dei segmenti con
velocit&agrave; superiore a 50 km/h e si dividono per la lunghezza totale del percorso.
Il valore riportato &egrave; la media delle 24 ore.
</div>
{df_to_html_table(tables['exceedance'])}
</section>

<!-- 3 TEMPORAL -->
<section id="temporal">
<h2>3. Profili Temporali di Velocit&agrave;</h2>
<h3>Velocit&agrave; Media Armonica per Ora</h3>
<img src="data:image/png;base64,{charts['speed_by_hour']}">
<p class="chart-caption">Velocit&agrave; media armonica dell&rsquo;intero itinerario per ora.
Pannello sinistro: feriali; pannello destro: festivi.</p>

<h3>V85 Medio per Ora</h3>
<img src="data:image/png;base64,{charts['v85_by_hour']}">
<p class="chart-caption">85&deg; percentile mediato su tutti i segmenti per ora.
Pannello sinistro: feriali; pannello destro: festivi.</p>

<h3>Mappe di Calore: Velocit&agrave; per Progressiva e Ora</h3>
<p>L&rsquo;asse orizzontale riporta la progressiva chilometrica; la larghezza di ciascuna cella
&egrave; proporzionale alla lunghezza effettiva del segmento. Si veda la mappa delle
progressive (Sezione 1) per il riferimento geografico.</p>
{heatmap_imgs}
<p class="chart-caption">Verde = flusso libero (velocit&agrave; elevate);
rosso = congestione (velocit&agrave; ridotte).</p>
</section>

<!-- 4 EXCEEDANCE -->
<section id="exceedance">
<h2>4. Superamento del Limite di Velocit&agrave;</h2>
<p>Per ogni ora si calcola la percentuale della <em>lunghezza del percorso</em> (km) in cui
la velocit&agrave; supera 50 km/h. Ogni segmento &egrave; pesato in proporzione alla propria
lunghezza.</p>
<img src="data:image/png;base64,{charts['exceedance']}">
<p class="chart-caption">% dei km di percorso oltre il limite per ora.
Sinistra: velocit&agrave; media; destra: V85.</p>
</section>

<!-- 5 V85 -->
<section id="v85">
<h2>5. Analisi V85 (85&deg; Percentile)</h2>
<div class="method-box">
<strong>Metodologia V85:</strong><br>
Il V85 &egrave; la velocit&agrave; al di sotto della quale viaggia l&rsquo;85% dei veicoli.
TomTom fornisce per ogni segmento e ora un array di 19 percentili (5&deg;&ndash;95&deg;,
passo 5); il V85 &egrave; il 17&deg; valore (indice 16).<br><br>
<strong>Profili spaziali:</strong> per ogni fascia oraria (notte 22&ndash;05, punta AM
07&ndash;08, mezzogiorno 12&ndash;13, punta PM 17&ndash;18) si calcola la media del V85 delle
ore della fascia e si riporta in funzione della progressiva chilometrica.<br><br>
<strong>Tabelle top 10:</strong> per ogni segmento si prende il V85 massimo tra le 24 ore.
I 10 segmenti con il valore pi&ugrave; elevato sono in tabella, con la progressiva (m) per
la localizzazione sulla mappa.
</div>

{v85_spatial_imgs}
<p class="chart-caption">Profilo spaziale V85 nelle diverse fasce orarie (Feriali).
La linea rossa tratteggiata indica il limite di 50 km/h.</p>

<h3>Top 10 Tratti con V85 pi&ugrave; Elevato &mdash; Dir. Centro (Feriali)</h3>
{df_to_html_table(tables['top_fast_centro'])}
{top10_centro_iframe}

<h3>Top 10 Tratti con V85 pi&ugrave; Elevato &mdash; Dir. GRA (Feriali)</h3>
{df_to_html_table(tables['top_fast_gra'])}
{top10_gra_iframe}
</section>

<!-- 6 VARIABILITY -->
<section id="variability">
<h2>6. Variabilit&agrave; delle Velocit&agrave;</h2>
<div class="method-box">
<strong>Metodologia:</strong> la deviazione standard &egrave; fornita direttamente da TomTom
per ogni segmento e fascia oraria. Misura la dispersione delle velocit&agrave; individuali
attorno alla media.<br><br>
Un&rsquo;elevata deviazione standard indica che nello stesso tratto convivono veicoli a
velocit&agrave; molto diverse, aumentando il rischio di incidenti.<br><br>
<strong>Mappe di calore:</strong> deviazione standard per progressiva (asse x) e ora (asse y).
La larghezza delle celle &egrave; proporzionale alla lunghezza del segmento.
Si veda la mappa delle progressive (Sezione 1) per il riferimento geografico.<br><br>
<strong>Profili spaziali:</strong> linea continua = media giornaliera della dev. std.;
area ombreggiata = valore massimo nell&rsquo;ora peggiore.
</div>
<img src="data:image/png;base64,{charts['variability']}">
<p class="chart-caption">Mappe di calore (sopra) e profili spaziali (sotto) della deviazione
standard per le due direzioni (Feriali).</p>
</section>

<!-- 7 NIGHT -->
<section id="night">
<h2>7. Analisi delle Velocit&agrave; Notturne</h2>
<p>Le ore notturne (22:00&ndash;05:59) presentano volumi ridotti e velocit&agrave; pi&ugrave;
elevate. I grafici sono separati per direzione perch&eacute; le progressive corrispondono
a posizioni geografiche diverse.</p>
<img src="data:image/png;base64,{charts['night']}">
<p class="chart-caption">Riga superiore: distribuzione diurna vs notturna.
Riga inferiore: profilo spaziale V85 e velocit&agrave; media notturna (Feriali).</p>
</section>

<!-- 8 COMPARISON -->
<section id="comparison">
<h2>8. Confronto Feriali vs Festivi</h2>
<img src="data:image/png;base64,{charts['weekday_weekend']}">
<p class="chart-caption">Velocit&agrave; medie (sopra) e V85 (sotto) feriali vs festivi.</p>
</section>

<!-- 9 MAPS -->
<section id="maps">
<h2>9. Mappe Interattive</h2>
<p>Cliccare su ciascun segmento per le statistiche dettagliate.
Lo spessore delle linee &egrave; stato aumentato per evidenziare
i tratti con velocit&agrave; pi&ugrave; elevate.</p>
{main_map_iframes}
</section>

<!-- 10 CONCLUSIONS -->
<section id="conclusions">
<h2>10. Conclusioni e Raccomandazioni</h2>
<div class="insight-box">
<strong>Risultati Principali:</strong>
<ul>
<li>L&rsquo;analisi copre circa 2.2&ndash;2.6 km di Corso Francia in entrambe le direzioni</li>
<li>Il limite di 50 km/h viene frequentemente superato, specialmente di notte</li>
<li>Il V85 supera il limite in una percentuale significativa dei km in tutte le fasce orarie</li>
<li>Le due direzioni mostrano profili asimmetrici nelle ore di punta</li>
</ul>
</div>
<div class="warning-box">
<strong>Aree di Attenzione:</strong>
<ul>
<li>I tratti con elevata variabilit&agrave; di velocit&agrave; richiedono attenzione per la sicurezza</li>
<li>Le velocit&agrave; notturne suggeriscono la necessit&agrave; di misure di moderazione</li>
<li>I festivi presentano velocit&agrave; generalmente pi&ugrave; elevate dei feriali</li>
</ul>
</div>
<p><em>Report generato automaticamente dai dati TomTom Speed Profiles.
Si raccomanda un&rsquo;interpretazione contestualizzata da parte di tecnici qualificati.</em></p>
</section>

<footer>
<p>TomTom Speed Profiles &mdash; Corso Francia, Roma &mdash; Febbraio 2026</p>
</footer>
</body></html>"""
    return html


# ================================================================
# SECTION 6b — WORD REPORT
# ================================================================

def _b64_to_stream(b64_str):
    """Decode a base64 PNG string into a BytesIO stream for python-docx."""
    return io.BytesIO(base64.b64decode(b64_str))


def _add_table_to_doc(doc, df):
    """Insert a pandas DataFrame as a styled Word table."""
    table = doc.add_table(rows=1, cols=len(df.columns), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(8)
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            row_cells[j].text = str(row_data[col])
            for par in row_cells[j].paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in par.runs:
                    run.font.size = Pt(8)


def _add_method_box(doc, text):
    """Add a method-note paragraph with distinct formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x19, 0x76, 0xD2)


def _add_caption(doc, text):
    """Add a centered italic caption below a figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def generate_word_report(charts, tables, static_map_paths):
    """Generate a Word (.docx) version of the report."""
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    img_width = Inches(6.5)

    # --- Styles ---
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)

    # ===== TITLE PAGE =====
    doc.add_paragraph()  # spacer
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("Analisi della Distribuzione delle Velocit\u00e0")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_p.add_run("Corso Francia \u2014 Roma")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x28, 0x35, 0x93)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run("Dati TomTom Speed Profiles \u2022 Periodo: 1\u201315 Febbraio 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS (manual) =====
    doc.add_heading("Indice", level=1)
    toc_items = [
        "1. Introduzione e Metodologia",
        "2. Panoramica dei Dati",
        "3. Profili Temporali di Velocit\u00e0",
        "4. Superamento del Limite di Velocit\u00e0",
        "5. Analisi V85 (85\u00b0 Percentile)",
        "6. Variabilit\u00e0 delle Velocit\u00e0",
        "7. Analisi delle Velocit\u00e0 Notturne",
        "8. Confronto Feriali vs Festivi",
        "9. Mappe Statiche delle Velocit\u00e0",
        "10. Conclusioni e Raccomandazioni",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ===== 1. INTRODUCTION =====
    doc.add_heading("1. Introduzione e Metodologia", level=1)

    doc.add_paragraph(
        "Il presente report analizza la distribuzione delle velocit\u00e0 veicolari lungo "
        "Corso di Francia a Roma, utilizzando i dati TomTom Speed Profiles "
        "relativi al periodo 1\u201315 febbraio 2026."
    )

    doc.add_paragraph("L\u2019analisi si basa su quattro dataset:")
    for item in [
        "Giorni Feriali (lun\u2013ven) \u2014 Direzione Centro e Direzione GRA",
        "Giorni Festivi (sab\u2013dom) \u2014 Direzione Centro e Direzione GRA",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Metriche principali:")
    metrics = [
        ("Velocit\u00e0 media armonica", "calcolata da TomTom come media armonica "
         "ponderata sulla lunghezza di tutti i segmenti, rappresentativa della velocit\u00e0 "
         "effettiva di percorrenza."),
        ("V85 (85\u00b0 percentile)", "velocit\u00e0 non superata dall\u201985% "
         "dei veicoli. TomTom fornisce 19 percentili (dal 5\u00b0 al 95\u00b0); il V85 corrisponde "
         "al 17\u00b0 valore dell\u2019array."),
        ("Deviazione standard", "dispersione delle velocit\u00e0 individuali, "
         "fornita direttamente da TomTom per ogni segmento e ora."),
        ("Tasso di superamento", "percentuale dei km di percorso "
         "(non dei segmenti) con velocit\u00e0 superiore al limite. Ogni segmento \u00e8 pesato "
         "in proporzione alla propria lunghezza."),
    ]
    for name, desc in metrics:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{name}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "Limite di velocit\u00e0: 50 km/h (alcuni tratti terminali su Viale "
        "Maresciallo Pilsudski: 40 km/h)."
    )

    doc.add_heading("Riferimento Spaziale: Progressive Chilometriche", level=2)
    doc.add_paragraph(
        "Tutti i grafici spaziali utilizzano la progressiva chilometrica: la distanza "
        "(in km) dall\u2019inizio del percorso. Le due direzioni hanno punti di partenza diversi, "
        "quindi le progressive sono specifiche per ciascuna direzione."
    )

    # ===== 2. OVERVIEW =====
    doc.add_heading("2. Panoramica dei Dati", level=1)

    _add_method_box(doc,
        "Nota metodologica: le velocit\u00e0 sono le medie armoniche a livello "
        "di percorso calcolate da TomTom. La media armonica pondera ogni segmento in base alla "
        "propria lunghezza. I valori per le fasce orarie sono la media dei valori orari "
        "ricadenti nella fascia (punta AM: 07\u201308, punta PM: 17\u201318, notturna: 22\u201305)."
    )

    _add_table_to_doc(doc, tables["overview"])
    doc.add_paragraph()

    doc.add_heading("Riepilogo Superamento Limiti (pesato per km)", level=2)
    _add_method_box(doc,
        "Nota metodologica: per ogni ora si sommano le lunghezze dei segmenti con "
        "velocit\u00e0 superiore a 50 km/h e si dividono per la lunghezza totale del percorso. "
        "Il valore riportato \u00e8 la media delle 24 ore."
    )
    _add_table_to_doc(doc, tables["exceedance"])

    # ===== 3. TEMPORAL PROFILES =====
    doc.add_page_break()
    doc.add_heading("3. Profili Temporali di Velocit\u00e0", level=1)

    doc.add_heading("Velocit\u00e0 Media Armonica per Ora", level=2)
    doc.add_picture(_b64_to_stream(charts["speed_by_hour"]), width=img_width)
    _add_caption(doc,
        "Velocit\u00e0 media armonica dell\u2019intero itinerario per ora. "
        "Pannello sinistro: feriali; pannello destro: festivi.")

    doc.add_heading("V85 Medio per Ora", level=2)
    doc.add_picture(_b64_to_stream(charts["v85_by_hour"]), width=img_width)
    _add_caption(doc,
        "85\u00b0 percentile mediato su tutti i segmenti per ora. "
        "Pannello sinistro: feriali; pannello destro: festivi.")

    doc.add_heading("Mappe di Calore: Velocit\u00e0 per Progressiva e Ora", level=2)
    doc.add_paragraph(
        "L\u2019asse orizzontale riporta la progressiva chilometrica; la larghezza di ciascuna cella "
        "\u00e8 proporzionale alla lunghezza effettiva del segmento."
    )
    for label, b64 in charts["heatmaps"].items():
        doc.add_heading(label, level=3)
        doc.add_picture(_b64_to_stream(b64), width=img_width)
    _add_caption(doc,
        "Verde = flusso libero (velocit\u00e0 elevate); rosso = congestione (velocit\u00e0 ridotte).")

    # ===== 4. EXCEEDANCE =====
    doc.add_page_break()
    doc.add_heading("4. Superamento del Limite di Velocit\u00e0", level=1)
    doc.add_paragraph(
        "Per ogni ora si calcola la percentuale della lunghezza del percorso (km) in cui "
        "la velocit\u00e0 supera 50 km/h. Ogni segmento \u00e8 pesato in proporzione alla propria "
        "lunghezza."
    )
    doc.add_picture(_b64_to_stream(charts["exceedance"]), width=img_width)
    _add_caption(doc,
        "% dei km di percorso oltre il limite per ora. Sinistra: velocit\u00e0 media; destra: V85.")

    # ===== 5. V85 ANALYSIS =====
    doc.add_page_break()
    doc.add_heading("5. Analisi V85 (85\u00b0 Percentile)", level=1)

    _add_method_box(doc,
        "Metodologia V85: il V85 \u00e8 la velocit\u00e0 al di sotto della quale viaggia l\u201985% dei veicoli. "
        "TomTom fornisce per ogni segmento e ora un array di 19 percentili (5\u00b0\u201395\u00b0, "
        "passo 5); il V85 \u00e8 il 17\u00b0 valore (indice 16).\n\n"
        "Profili spaziali: per ogni fascia oraria (notte 22\u201305, punta AM "
        "07\u201308, mezzogiorno 12\u201313, punta PM 17\u201318) si calcola la media del V85 delle "
        "ore della fascia e si riporta in funzione della progressiva chilometrica.\n\n"
        "Tabelle top 10: per ogni segmento si prende il V85 massimo tra le 24 ore. "
        "I 10 segmenti con il valore pi\u00f9 elevato sono in tabella."
    )

    for direction, b64 in charts["v85_spatial"].items():
        doc.add_heading(f"Profilo Spaziale V85 \u2014 Dir. {direction}", level=2)
        doc.add_picture(_b64_to_stream(b64), width=img_width)
    _add_caption(doc,
        "Profilo spaziale V85 nelle diverse fasce orarie (Feriali). "
        "La linea rossa tratteggiata indica il limite di 50 km/h.")

    doc.add_heading("Top 10 Tratti con V85 pi\u00f9 Elevato \u2014 Dir. Centro (Feriali)", level=2)
    _add_table_to_doc(doc, tables["top_fast_centro"])
    doc.add_paragraph()

    doc.add_heading("Top 10 Tratti con V85 pi\u00f9 Elevato \u2014 Dir. GRA (Feriali)", level=2)
    _add_table_to_doc(doc, tables["top_fast_gra"])

    # ===== 6. VARIABILITY =====
    doc.add_page_break()
    doc.add_heading("6. Variabilit\u00e0 delle Velocit\u00e0", level=1)

    _add_method_box(doc,
        "Metodologia: la deviazione standard \u00e8 fornita direttamente da TomTom "
        "per ogni segmento e fascia oraria. Misura la dispersione delle velocit\u00e0 individuali "
        "attorno alla media. Un\u2019elevata deviazione standard indica che nello stesso tratto "
        "convivono veicoli a velocit\u00e0 molto diverse, aumentando il rischio di incidenti.\n\n"
        "Mappe di calore: deviazione standard per progressiva (asse x) e ora (asse y). "
        "Profili spaziali: linea continua = media giornaliera; area ombreggiata = massimo orario."
    )

    doc.add_picture(_b64_to_stream(charts["variability"]), width=img_width)
    _add_caption(doc,
        "Mappe di calore (sopra) e profili spaziali (sotto) della deviazione "
        "standard per le due direzioni (Feriali).")

    # ===== 7. NIGHT =====
    doc.add_page_break()
    doc.add_heading("7. Analisi delle Velocit\u00e0 Notturne", level=1)
    doc.add_paragraph(
        "Le ore notturne (22:00\u201305:59) presentano volumi ridotti e velocit\u00e0 pi\u00f9 "
        "elevate. I grafici sono separati per direzione perch\u00e9 le progressive corrispondono "
        "a posizioni geografiche diverse."
    )
    doc.add_picture(_b64_to_stream(charts["night"]), width=img_width)
    _add_caption(doc,
        "Riga superiore: distribuzione diurna vs notturna. "
        "Riga inferiore: profilo spaziale V85 e velocit\u00e0 media notturna (Feriali).")

    # ===== 8. WEEKDAY vs WEEKEND =====
    doc.add_page_break()
    doc.add_heading("8. Confronto Feriali vs Festivi", level=1)
    doc.add_picture(_b64_to_stream(charts["weekday_weekend"]), width=img_width)
    _add_caption(doc,
        "Velocit\u00e0 medie (sopra) e V85 (sotto) feriali vs festivi.")

    # ===== 9. STATIC MAPS =====
    doc.add_page_break()
    doc.add_heading("9. Mappe Statiche delle Velocit\u00e0", level=1)
    doc.add_paragraph(
        "Le mappe seguenti mostrano la distribuzione spaziale delle velocit\u00e0 "
        "lungo Corso Francia. Il colore dei segmenti indica la velocit\u00e0: "
        "verde = velocit\u00e0 ridotta, giallo = limite di 50 km/h, rosso = velocit\u00e0 elevata."
    )

    # Categorize static maps
    map_groups = {
        "Velocit\u00e0 Media (24h)": "avg_speed_24h",
        "V85 (24h)": "v85_24h",
        "Velocit\u00e0 Media Notturna": "avg_speed_night",
        "V85 Notturno": "v85_night",
        "V85 Massimo (ora peggiore)": "max_v85",
    }
    for group_label, prefix in map_groups.items():
        doc.add_heading(group_label, level=2)
        # Side-by-side version
        sbs_path = STATIC_MAPS_DIR / f"{prefix}_feriali_sidebyside.png"
        if sbs_path.exists():
            doc.add_picture(str(sbs_path), width=img_width)
            _add_caption(doc, f"{group_label} \u2014 Confronto Dir. Centro e Dir. GRA (Feriali)")
        # Individual maps
        for direction in ["Centro", "GRA"]:
            indiv_path = STATIC_MAPS_DIR / f"{prefix}_feriali_{direction.lower()}.png"
            if indiv_path.exists():
                doc.add_heading(f"Dir. {direction}", level=3)
                doc.add_picture(str(indiv_path), width=img_width)

    # ===== 10. CONCLUSIONS =====
    doc.add_page_break()
    doc.add_heading("10. Conclusioni e Raccomandazioni", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Risultati Principali:")
    run.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    for item in [
        "L\u2019analisi copre circa 2.2\u20132.6 km di Corso Francia in entrambe le direzioni",
        "Il limite di 50 km/h viene frequentemente superato, specialmente di notte",
        "Il V85 supera il limite in una percentuale significativa dei km in tutte le fasce orarie",
        "Le due direzioni mostrano profili asimmetrici nelle ore di punta",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Aree di Attenzione:")
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0x98, 0x00)
    for item in [
        "I tratti con elevata variabilit\u00e0 di velocit\u00e0 richiedono attenzione per la sicurezza",
        "Le velocit\u00e0 notturne suggeriscono la necessit\u00e0 di misure di moderazione",
        "I festivi presentano velocit\u00e0 generalmente pi\u00f9 elevate dei feriali",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(
        "Report generato automaticamente dai dati TomTom Speed Profiles. "
        "Si raccomanda un\u2019interpretazione contestualizzata da parte di tecnici qualificati."
    )
    run.italic = True
    run.font.size = Pt(9)

    # --- Footer line ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TomTom Speed Profiles \u2014 Corso Francia, Roma \u2014 Febbraio 2026")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # --- Save ---
    docx_path = OUTPUT_DIR / "report_corso_francia.docx"
    doc.save(str(docx_path))
    return str(docx_path)


# ================================================================
# SECTION 7 — MAIN
# ================================================================

def main():
    print("Caricamento dati...")
    segments, summaries, headers = load_all_data()

    print("Calcolo statistiche per segmento...")
    seg_stats = segment_peak_stats(segments)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    MAPS_DIR.mkdir(parents=True, exist_ok=True)

    print("Generazione grafici...")
    charts = {}
    charts["speed_by_hour"] = chart_speed_by_hour(summaries)
    print("  - Velocita media per ora (Feriali/Festivi)")
    charts["v85_by_hour"] = chart_v85_by_hour(segments)
    print("  - V85 per ora (Feriali/Festivi)")
    charts["heatmaps"] = chart_heatmaps(segments)
    print("  - Mappe di calore (progressive)")
    charts["exceedance"] = chart_exceedance(segments)
    print("  - Superamento limiti (per km)")
    charts["v85_spatial"] = chart_v85_spatial(segments)
    print("  - Profili spaziali V85")
    charts["variability"] = chart_speed_variability(segments)
    print("  - Variabilita velocita (progressive)")
    charts["night"] = chart_night_analysis(segments)
    print("  - Analisi notturna (per direzione)")
    charts["weekday_weekend"] = chart_weekday_weekend(segments)
    print("  - Confronto feriali/festivi")

    print("Generazione mappe statiche (contextily)...")
    static_map_paths = generate_all_static_maps(segments)
    print(f"  - {len(static_map_paths)} mappe statiche generate in output/static_maps/")

    print("Generazione mappe interattive...")
    map_files = generate_maps(segments, seg_stats)
    print(f"  - {len(map_files)} mappe generate")

    print("Costruzione tabelle...")
    tables = build_summary_tables(segments, summaries, seg_stats)

    print("Assemblaggio report HTML...")
    html = generate_html_report(charts, map_files, tables)
    report_path = OUTPUT_DIR / "report_corso_francia.html"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write(html)

    print("Assemblaggio report Word (.docx)...")
    docx_path = generate_word_report(charts, tables, static_map_paths)
    print(f"  - DOCX: {docx_path}")

    csv_path = OUTPUT_DIR / "dati_segmenti.csv"
    segments.drop(columns=["geometry"]).to_csv(csv_path, index=False)
    print(f"  - CSV: {csv_path}")
    print(f"\nReport HTML: {report_path}")
    print(f"Report DOCX: {docx_path}")
    print(f"Mappe:       {MAPS_DIR}")


if __name__ == "__main__":
    main()
