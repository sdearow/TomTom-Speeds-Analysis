#!/usr/bin/env python3
"""
Report Ridotto — Via Cristoforo Colombo, Roma
Versione sintetica del report completo.
"""

# ================================================================
# BLOCK A — IMPORTS & DATA RE-USE
# ================================================================

import io, base64, warnings
from pathlib import Path

import numpy as np
import pandas as pd
import geopandas as gpd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import folium
import branca.colormap as cm
import contextily as cx
import ruptures as rpt
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

warnings.filterwarnings("ignore")

from analisi_velocita_colombo import (
    load_all_data, segment_peak_stats, fig_to_base64,
    CARRIAGE_GROUPS, CARRIAGE_COLORS, ALL_DAY_TYPES, DAY_TYPE_LABELS,
    HOUR_SHORT, NIGHT_HOURS, AM_PEAK, PM_PEAK, MIDDAY,
    TRATTA_COLORS,
    hourly_means, compute_exceedance_by_km, get_speed_limit_for_group,
    _build_speed_gdf, _add_basemap,
    _seg_geodataframe, make_folium_map, make_folium_tratte_map,
    _b64_to_stream, _add_table_to_doc, _add_method_box, _add_caption,
    OUTPUT_DIR, STATIC_MAPS_DIR,
)

RIDOTTO_DIR = OUTPUT_DIR / "ridotto"
RIDOTTO_MAPS = RIDOTTO_DIR / "maps"
RIDOTTO_STATIC = RIDOTTO_DIR / "static_maps"


# ================================================================
# BLOCK B — CHARTS WITHOUT SPEED-LIMIT LINES
# ================================================================

def chart_speed_by_hour_nolim(summaries, carriage):
    """Speed by hour — NO 50 km/h reference line."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs, colors = cg["directions"], CARRIAGE_COLORS[carriage]
    fig, axes = plt.subplots(1, len(ALL_DAY_TYPES),
                             figsize=(5 * len(ALL_DAY_TYPES), 5), sharey=True)
    if len(ALL_DAY_TYPES) == 1:
        axes = [axes]
    for ax, dt in zip(axes, ALL_DAY_TYPES):
        for d in dirs:
            sub = summaries[(summaries["carriage"] == carriage)
                            & (summaries["direction"] == d)
                            & (summaries["day_type"] == dt)].sort_values("hour")
            if not sub.empty:
                ax.plot(sub["hour"], sub["harm_avg_speed"], color=colors[d],
                        label=cg["dir_labels"][d], lw=2, marker="o", ms=3)
        ax.set_xlabel("Ora"); ax.set_title(DAY_TYPE_LABELS[dt], fontsize=10)
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=6)
        ax.legend(fontsize=7); ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
    axes[0].set_ylabel("Velocità media armonica (km/h)")
    fig.suptitle(f"Velocità Media per Ora — {cg['label']}", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_v85_by_hour_nolim(segments, carriage):
    """V85 by hour — NO 50 km/h reference line."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs, colors = cg["directions"], CARRIAGE_COLORS[carriage]
    hmeans = hourly_means(segments[segments["carriage"] == carriage], "p85")
    fig, axes = plt.subplots(1, len(ALL_DAY_TYPES),
                             figsize=(5 * len(ALL_DAY_TYPES), 5), sharey=True)
    if len(ALL_DAY_TYPES) == 1:
        axes = [axes]
    for ax, dt in zip(axes, ALL_DAY_TYPES):
        for d in dirs:
            sub = hmeans[(hmeans["direction"] == d)
                         & (hmeans["day_type"] == dt)].sort_values("hour")
            if not sub.empty:
                ax.plot(sub["hour"], sub["p85"], color=colors[d],
                        label=cg["dir_labels"][d], lw=2, marker="o", ms=3)
        ax.set_xlabel("Ora"); ax.set_title(DAY_TYPE_LABELS[dt], fontsize=10)
        ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=6)
        ax.legend(fontsize=7); ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
    axes[0].set_ylabel("V85 medio (km/h)")
    fig.suptitle(f"V85 per Ora — {cg['label']}", fontsize=13, y=1.02)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_v85_spatial_nolim(segments, carriage, day_type="Inv. Feriale"):
    """V85 spatial profile — NO speed-limit line."""
    cg = CARRIAGE_GROUPS[carriage]
    results = {}
    periods = {"Notte (22–05)": NIGHT_HOURS, "Punta AM (07–08)": AM_PEAK,
               "Mezzog. (12–13)": MIDDAY, "Punta PM (17–18)": PM_PEAK}
    pcols = ["#7570b3", "#d95f02", "#1b9e77", "#e7298a"]
    for d in cg["directions"]:
        sub = segments[(segments["carriage"] == carriage)
                       & (segments["direction"] == d)
                       & (segments["day_type"] == day_type)]
        if sub.empty:
            continue
        fig, ax = plt.subplots(figsize=(12, 5))
        for (pn, ph), pc in zip(periods.items(), pcols):
            ps = sub[sub["hour"].isin(ph)]
            prof = ps.groupby("seg_idx").agg(
                v85=("p85", "mean"), dist=("cum_dist_mid", "first")).sort_values("dist")
            ax.plot(prof["dist"] / 1000, prof["v85"], color=pc, label=pn, lw=1.8)
        ax.set_xlabel("Progressiva (km)"); ax.set_ylabel("V85 (km/h)")
        label = cg["dir_labels"][d]
        ax.set_title(f"Profilo Spaziale V85 — {DAY_TYPE_LABELS[day_type]} {label}")
        ax.legend(fontsize=8); ax.grid(True, alpha=.3)
        results[label] = fig_to_base64(fig)
    return results


def chart_night_nolim(segments, carriage, day_type="Inv. Feriale"):
    """Night analysis — NO speed-limit reference line."""
    cg = CARRIAGE_GROUPS[carriage]
    dirs = cg["directions"]; ncols = len(dirs)
    fig, axes = plt.subplots(2, ncols, figsize=(7 * ncols, 10))
    if ncols == 1:
        axes = axes.reshape(2, 1)
    sub_all = segments[(segments["carriage"] == carriage)
                       & (segments["day_type"] == day_type)]
    for ci, d in enumerate(dirs):
        sd = sub_all[sub_all["direction"] == d]
        if sd.empty:
            continue
        night = sd[sd["hour"].isin(NIGHT_HOURS)]["avg_speed"]
        day = sd[~sd["hour"].isin(NIGHT_HOURS)]["avg_speed"]
        ax = axes[0, ci]
        ax.hist(day, bins=40, alpha=.6, color="#1f77b4", label="Diurno (06–21)", density=True)
        ax.hist(night, bins=40, alpha=.6, color="#9467bd", label="Notturno (22–05)", density=True)
        ax.set_xlabel("Velocità media (km/h)"); ax.set_ylabel("Densità")
        ax.set_title(f"Distribuzione — {cg['dir_labels'][d]}"); ax.legend(fontsize=8)
        ax2 = axes[1, ci]
        sn = sd[sd["hour"].isin(NIGHT_HOURS)]
        prof = sn.groupby("seg_idx").agg(
            v85=("p85", "mean"), avg=("avg_speed", "mean"),
            dist=("cum_dist_mid", "first")).sort_values("dist")
        ax2.plot(prof["dist"] / 1000, prof["v85"], color="#9467bd", lw=2, label="V85 notturno")
        ax2.plot(prof["dist"] / 1000, prof["avg"], color="#1f77b4", lw=2, label="Vel. media notturna")
        ax2.set_xlabel("Progressiva (km)"); ax2.set_ylabel("Velocità (km/h)")
        ax2.set_title(f"Profilo Notturno — {cg['dir_labels'][d]}")
        ax2.legend(fontsize=8); ax2.grid(True, alpha=.3)
    fig.suptitle(f"Analisi Notturna — {cg['label']} {DAY_TYPE_LABELS[day_type]}",
                 fontsize=12, y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


def chart_temporal_by_tratta(segments, carriage):
    """Separate temporal charts per tratta (for laterali).
    Returns dict { tratta_name: { 'speed_by_hour': b64, 'v85_by_hour': b64 } }
    """
    cg = CARRIAGE_GROUPS[carriage]
    colors = CARRIAGE_COLORS[carriage]
    result = {}
    for d in cg["directions"]:
        color = colors[d]
        # Vel. media per ora
        fig1, axes1 = plt.subplots(1, len(ALL_DAY_TYPES),
                                   figsize=(5 * len(ALL_DAY_TYPES), 4.5), sharey=True)
        if len(ALL_DAY_TYPES) == 1:
            axes1 = [axes1]
        for ax, dt in zip(axes1, ALL_DAY_TYPES):
            sub = segments[(segments["carriage"] == carriage)
                           & (segments["direction"] == d)
                           & (segments["day_type"] == dt)]
            if sub.empty:
                continue
            hm = sub.groupby("hour")["avg_speed"].mean().reset_index().sort_values("hour")
            ax.plot(hm["hour"], hm["avg_speed"], color=color, lw=2, marker="o", ms=3)
            ax.set_xlabel("Ora"); ax.set_title(DAY_TYPE_LABELS[dt], fontsize=10)
            ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=6)
            ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
        axes1[0].set_ylabel("Velocità media (km/h)")
        fig1.suptitle(f"Velocità Media per Ora — {d}", fontsize=12, y=1.02)
        fig1.tight_layout()
        speed_b64 = fig_to_base64(fig1)
        # V85 per ora
        hmeans = hourly_means(segments[segments["carriage"] == carriage], "p85")
        fig2, axes2 = plt.subplots(1, len(ALL_DAY_TYPES),
                                   figsize=(5 * len(ALL_DAY_TYPES), 4.5), sharey=True)
        if len(ALL_DAY_TYPES) == 1:
            axes2 = [axes2]
        for ax, dt in zip(axes2, ALL_DAY_TYPES):
            sub = hmeans[(hmeans["direction"] == d)
                         & (hmeans["day_type"] == dt)].sort_values("hour")
            if not sub.empty:
                ax.plot(sub["hour"], sub["p85"], color=color, lw=2, marker="o", ms=3)
            ax.set_xlabel("Ora"); ax.set_title(DAY_TYPE_LABELS[dt], fontsize=10)
            ax.set_xticks(range(24)); ax.set_xticklabels(HOUR_SHORT, fontsize=6)
            ax.grid(True, alpha=.3); ax.set_xlim(-.5, 23.5)
        axes2[0].set_ylabel("V85 (km/h)")
        fig2.suptitle(f"V85 per Ora — {d}", fontsize=12, y=1.02)
        fig2.tight_layout()
        result[d] = {"speed_by_hour": speed_b64, "v85_by_hour": fig_to_base64(fig2)}
    return result


# ================================================================
# BLOCK C — TABLES (Direzione/Tratta first col, grouped by direction)
# ================================================================

def build_tables_ridotto(segments, summaries):
    """Tables with Direzione/Tratta as first column, grouped by direction."""
    tables = {}
    for carriage in CARRIAGE_GROUPS:
        cg = CARRIAGE_GROUPS[carriage]
        dirs = cg["directions"]
        slug = carriage.lower().replace(". ", "_").replace(" ", "_")
        is_lat = carriage.startswith("Lat.")
        col1 = "Tratta" if is_lat else "Direzione"

        # Overview: direction first, rows grouped by direction
        rows = []
        for d in dirs:
            for dt in ALL_DAY_TYPES:
                ss = summaries[(summaries["carriage"] == carriage)
                               & (summaries["direction"] == d)
                               & (summaries["day_type"] == dt)]
                if ss.empty:
                    continue
                rows.append({
                    col1: cg["dir_labels"][d],
                    "Periodo": DAY_TYPE_LABELS[dt],
                    "Vel. Media 24h": f"{ss['harm_avg_speed'].mean():.1f}",
                    "Vel. Punta AM": f"{ss[ss['hour'].isin(AM_PEAK)]['harm_avg_speed'].mean():.1f}",
                    "Vel. Punta PM": f"{ss[ss['hour'].isin(PM_PEAK)]['harm_avg_speed'].mean():.1f}",
                    "Vel. Notturna": f"{ss[ss['hour'].isin(NIGHT_HOURS)]['harm_avg_speed'].mean():.1f}",
                })
        tables[f"overview_{slug}"] = pd.DataFrame(rows)

        # Exceedance: direction first, rows grouped by direction, no "Nota"
        exc_rows = []
        for d in dirs:
            for dt in ["Inv. Feriale", "Est. Feriale"]:
                sub = segments[(segments["carriage"] == carriage)
                               & (segments["direction"] == d)
                               & (segments["day_type"] == dt)]
                if sub.empty:
                    continue
                total_m = sub.drop_duplicates("seg_idx")["seg_distance"].sum()

                def _pct(col, _s=sub, _t=total_m):
                    v = []
                    for h in range(24):
                        hs = _s[_s["hour"] == h]
                        v.append(hs.loc[hs[col] > hs["speedLimit"],
                                        "seg_distance"].sum() / _t * 100)
                    return np.mean(v)

                exc_rows.append({
                    col1: cg["dir_labels"][d],
                    "Periodo": DAY_TYPE_LABELS[dt],
                    "% km Vel.Media > Limite": f"{_pct('avg_speed'):.1f}%",
                    "% km V85 > Limite": f"{_pct('p85'):.1f}%",
                    "V85 massimo (km/h)": f"{sub['p85'].max():.0f}",
                })
        tables[f"exceedance_{slug}"] = pd.DataFrame(exc_rows)

    return tables


# ================================================================
# BLOCK D — TRATTE WITHOUT INTERSECTION SEGMENTS
# ================================================================

def detect_intersection_segments(segments, direction, threshold_ratio=0.80):
    """Flag segments whose mean V85 drops below threshold_ratio * corridor median.
    These are likely signalised intersections causing forced slowdowns.
    """
    c = segments[(segments["carriage"] == "Centrale")
                 & (segments["direction"] == direction)
                 & (segments["day_type"] == "Inv. Feriale")]
    seg = c.groupby("seg_idx").agg(
        v85_mean=("p85", "mean"),
        cum_dist_start=("cum_dist_start", "first"),
    ).sort_values("cum_dist_start")
    threshold = seg["v85_mean"].median() * threshold_ratio
    return set(seg[seg["v85_mean"] < threshold].index.tolist())


def compute_tratte_no_intersections(segments, n_bkps=5):
    """Compute tratte excluding intersection segments from segmentation signal.
    Intersections are re-assigned to nearest tratta afterwards but don't
    influence the statistics.
    """
    day_type = "Inv. Feriale"
    c = segments[(segments["carriage"] == "Centrale")
                 & (segments["day_type"] == day_type)]

    tratte_rows, tratte_seg_map = [], {}
    intersection_info = {}

    for direction in ["Ostia", "Centro"]:
        sub = c[c["direction"] == direction]
        intersections = detect_intersection_segments(segments, direction)
        intersection_info[direction] = intersections

        seg = sub.groupby("seg_idx").agg(
            cum_dist_start=("cum_dist_start", "first"),
            cum_dist_end=("cum_dist_end", "first"),
            seg_distance=("seg_distance", "first"),
            speedLimit=("speedLimit", "first"),
            v85_mean=("p85", "mean"),
            avg_speed=("avg_speed", "mean"),
        ).sort_values("cum_dist_start").reset_index()

        seg_clean = seg[~seg["seg_idx"].isin(intersections)].reset_index(drop=True)

        seg_hours = sub.pivot_table(
            index="seg_idx", columns="hour", values="p85", aggfunc="mean")
        signal = seg_hours.loc[seg_clean["seg_idx"].tolist()].values

        min_sz = max(2, min(3, len(seg_clean) // (n_bkps + 1)))
        algo = rpt.Binseg(model="l2", min_size=min_sz).fit(signal)
        bkps = algo.predict(n_bkps=n_bkps)

        # Build tratte from clean segments
        start, tratta_ranges = 0, []
        for t_id, end in enumerate(bkps, 1):
            sl = seg_clean.iloc[start:end]
            t_segs = set(sl["seg_idx"].tolist())
            tratta_ranges.append((t_id, t_segs))
            for _, r in sl.iterrows():
                tratte_seg_map[(direction, int(r["seg_idx"]))] = t_id
            start = end

        # Re-assign intersections to nearest tratta
        for iseg in sorted(intersections):
            ir = seg[seg["seg_idx"] == iseg]
            if ir.empty:
                continue
            idist = ir["cum_dist_start"].iloc[0]
            best_t, best_d = 1, float("inf")
            for t_id, t_segs in tratta_ranges:
                tr = seg[seg["seg_idx"].isin(t_segs)]
                if tr.empty:
                    continue
                d = abs(idist - tr["cum_dist_start"].mean())
                if d < best_d:
                    best_d, best_t = d, t_id
            tratte_seg_map[(direction, int(iseg))] = best_t
            tratta_ranges[best_t - 1][1].add(iseg)

        # Summary rows (stats from clean segments only)
        for t_id, t_segs in tratta_ranges:
            sl = seg[seg["seg_idx"].isin(t_segs)].sort_values("cum_dist_start")
            if sl.empty:
                continue
            sl_c = sl[~sl["seg_idx"].isin(intersections)]
            if sl_c.empty:
                sl_c = sl
            w = sl_c["seg_distance"].values
            v85_avg = np.average(sl_c["v85_mean"], weights=w)
            v85_std = np.sqrt(np.average((sl_c["v85_mean"] - v85_avg) ** 2, weights=w))
            cv = v85_std / v85_avg * 100 if v85_avg > 0 else 0
            tratte_rows.append({
                "direction": direction, "tratta": t_id,
                "cum_dist_start": sl["cum_dist_start"].iloc[0],
                "cum_dist_end": sl["cum_dist_end"].iloc[-1],
                "length_km": (sl["cum_dist_end"].iloc[-1] - sl["cum_dist_start"].iloc[0]) / 1000,
                "v85_mean": round(v85_avg, 1),
                "v85_min": round(sl_c["v85_mean"].min(), 1),
                "v85_max": round(sl_c["v85_mean"].max(), 1),
                "v85_cv": round(cv, 1),
                "n_segments": len(sl),
                "n_intersections": len(sl[sl["seg_idx"].isin(intersections)]),
            })

    return pd.DataFrame(tratte_rows), tratte_seg_map, intersection_info


# ================================================================
# BLOCK E — MAPS (V85 24h + V85 notturno only, no progressives)
# ================================================================

def generate_static_maps_ridotto(segments):
    """Only V85 24h + V85 notturno, no progressive annotations."""
    RIDOTTO_STATIC.mkdir(parents=True, exist_ok=True)
    paths = []; day_type = "Inv. Feriale"
    configs = [
        ("p85", None,        "V85 (24h)",    "v85_24h",   "mean"),
        ("p85", NIGHT_HOURS, "V85 Notturno", "v85_night", "mean"),
    ]
    for carriage in CARRIAGE_GROUPS:
        cg = CARRIAGE_GROUPS[carriage]
        dirs = list(cg["directions"])
        slug = carriage.lower().replace(". ", "_").replace(" ", "_")
        if carriage == "Lat. Centro":
            dirs = [d for d in dirs if d != "Svincolo Malafede"]
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            for vcol, hfilter, label, prefix, agg in configs:
                period = "22–05" if hfilter else "tutte le ore"
                title = f"{label} — {cg['label']} {cg['dir_labels'][d]} ({period})"
                fname = f"{prefix}_{slug}_{d_slug}.png"
                gdf = _build_speed_gdf(segments, carriage, d, day_type, vcol,
                                       hfilter, agg_func=agg)
                if gdf.empty:
                    continue
                fig, ax = plt.subplots(figsize=(10, 14))
                norm = mcolors.TwoSlopeNorm(vmin=20, vcenter=60, vmax=120)
                cmap = plt.cm.RdYlGn_r
                gdf.plot(ax=ax, color="#d0d0d0", linewidth=9, zorder=1, alpha=0.8)
                gdf.plot(ax=ax, column=vcol, cmap=cmap, norm=norm,
                         linewidth=5, legend=False, zorder=2)
                has_tiles = _add_basemap(ax, crs="EPSG:4326")
                if has_tiles:
                    ax.set_axis_off()
                sm = plt.cm.ScalarMappable(cmap=cmap, norm=norm); sm.set_array([])
                cbar = fig.colorbar(sm, ax=ax, shrink=0.4, pad=0.03, aspect=30)
                cbar.set_label("km/h", fontsize=10)
                ax.set_title(title, fontsize=12, fontweight="bold", pad=12)
                fig.tight_layout()
                out = RIDOTTO_STATIC / fname
                fig.savefig(str(out), dpi=150, bbox_inches="tight", facecolor="white")
                plt.close(fig); paths.append(str(out))
                print(f"  - {fname}")
    return paths


def generate_maps_ridotto(segments):
    """Interactive: V85 24h + V85 notturno, no progressive/top10/variability."""
    RIDOTTO_MAPS.mkdir(parents=True, exist_ok=True)
    map_files = {}; day_type = "Inv. Feriale"
    configs = [
        ("p85", None,        "V85 (24h)",    "v85_24h",   "mean", 20, 130),
        ("p85", NIGHT_HOURS, "V85 Notturno", "v85_night", "mean", 20, 130),
    ]
    for carriage in CARRIAGE_GROUPS:
        cg = CARRIAGE_GROUPS[carriage]
        dirs = list(cg["directions"])
        slug = carriage.lower().replace(". ", "_").replace(" ", "_")
        if carriage == "Lat. Centro":
            dirs = [d for d in dirs if d != "Svincolo Malafede"]
        for d in dirs:
            d_slug = d.lower().replace(" ", "_").replace("-", "_")
            for vcol, hfilter, label, prefix, agg, vmin, vmax in configs:
                period = "22–05" if hfilter else "tutte le ore"
                title = f"{label} — {cg['label']} {cg['dir_labels'][d]} ({period})"
                fname = f"{prefix}_{slug}_{d_slug}.html"
                sub = segments[segments["carriage"] == carriage]
                if hfilter:
                    sub = sub[sub["hour"].isin(hfilter)]
                gdf = _seg_geodataframe(sub, carriage, d, day_type, vcol, agg)
                if gdf.empty:
                    continue
                m = make_folium_map(gdf, vcol, title, vmin, vmax,
                                    add_progressive=False)
                m.save(str(RIDOTTO_MAPS / fname))
                map_files[fname] = title
    return map_files


# ================================================================
# BLOCK F — TRATTE PROFILE CHART (intersections marked)
# ================================================================

def chart_tratte_profile_ridotto(segments, tratte_df, intersection_info):
    """V85 spatial profile colour-coded by tratta. Intersections marked with X."""
    import matplotlib.patches as mpatches
    day_type = "Inv. Feriale"
    c = segments[(segments["carriage"] == "Centrale")
                 & (segments["day_type"] == day_type)]

    fig, axes = plt.subplots(2, 1, figsize=(16, 12), sharex=False)
    for ax_i, direction in enumerate(["Ostia", "Centro"]):
        ax = axes[ax_i]
        sub = c[c["direction"] == direction]
        seg = sub.groupby("seg_idx").agg(
            v85=("p85", "mean"), dist=("cum_dist_mid", "first"),
            seg_distance=("seg_distance", "first"),
        ).sort_values("dist")
        td = tratte_df[tratte_df["direction"] == direction].sort_values("tratta")
        inters = intersection_info.get(direction, set())

        ax.plot(seg["dist"] / 1000, seg["v85"], color="#888", lw=0.8, alpha=0.5)
        for _, tr in td.iterrows():
            t_id = int(tr["tratta"])
            col = TRATTA_COLORS[(t_id - 1) % len(TRATTA_COLORS)]
            t_s, t_e = tr["cum_dist_start"] / 1000, tr["cum_dist_end"] / 1000
            ax.axvspan(t_s, t_e, alpha=0.12, color=col)
            y_label = 118 if ax_i == 0 else seg["v85"].max() + 3
            ax.text((t_s + t_e) / 2, y_label, f"T{t_id}",
                    ha="center", fontsize=9, fontweight="bold", color=col)
        # Intersection X markers
        for iseg in inters:
            if iseg in seg.index:
                r = seg.loc[iseg]
                ax.plot(r["dist"] / 1000, r["v85"], "x",
                        color="red", ms=8, markeredgewidth=2, zorder=5)
        # Non-intersection coloured dots
        for _, tr in td.iterrows():
            t_id = int(tr["tratta"])
            col = TRATTA_COLORS[(t_id - 1) % len(TRATTA_COLORS)]
            mask = ((seg["dist"] / 1000 >= tr["cum_dist_start"] / 1000)
                    & (seg["dist"] / 1000 <= tr["cum_dist_end"] / 1000)
                    & (~seg.index.isin(inters)))
            ax.scatter(seg.loc[mask, "dist"] / 1000, seg.loc[mask, "v85"],
                       color=col, s=20, zorder=4)
        ax.set_ylabel("V85 (km/h)"); ax.set_xlabel("Progressiva (km)")
        ax.set_title(f"Dir. {direction}", fontsize=12, fontweight="bold")
        ax.grid(True, alpha=.2)

    cross_patch = mpatches.Patch(color="red", label="Intersezioni semaforiche (escluse)")
    fig.legend(handles=[cross_patch], loc="upper right", fontsize=9)
    fig.suptitle("Profilo V85 con Tratte Omogenee (intersezioni escluse)",
                 fontsize=14, fontweight="bold", y=1.01)
    fig.tight_layout()
    return fig_to_base64(fig)


# ================================================================
# BLOCK G — HTML REPORT ASSEMBLY
# ================================================================

def _df_html(df):
    return df.to_html(index=False, classes="data-table", border=0)

def _ifr(fname, title):
    return (f'<h4>{title}</h4>'
            f'<iframe src="maps/{fname}" width="100%" height="600" '
            f'style="border:1px solid #ddd;border-radius:4px;"></iframe>')

CSS = """
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:'Segoe UI',Tahoma,Geneva,Verdana,sans-serif;line-height:1.6;
  color:#333;background:#f9f9f9;max-width:1200px;margin:0 auto;padding:20px}
header{background:linear-gradient(135deg,#1a237e,#283593);color:#fff;
  padding:40px;border-radius:8px;margin-bottom:30px;text-align:center}
header h1{font-size:2em;margin-bottom:8px}
header h2{font-size:1.3em;font-weight:300;opacity:.9}
header p{margin-top:12px;opacity:.8;font-size:.95em}
section{background:#fff;padding:30px;border-radius:8px;margin-bottom:20px;
  box-shadow:0 2px 4px rgba(0,0,0,.1)}
h2{color:#1a237e;border-bottom:2px solid #e8eaf6;padding-bottom:8px;margin-bottom:20px}
h3{color:#283593;margin:20px 0 10px 0} h4{color:#3949ab;margin:15px 0 8px 0}
img{max-width:100%;height:auto;border-radius:4px;margin:10px 0}
.data-table{width:100%;border-collapse:collapse;margin:15px 0;font-size:.9em}
.data-table th{background:#1a237e;color:#fff;padding:10px 12px;text-align:left}
.data-table td{padding:8px 12px;border-bottom:1px solid #e0e0e0}
.data-table tr:nth-child(even){background:#f5f5f5}
.method-box{background:#e3f2fd;border-left:4px solid #1976d2;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}
.insight-box{background:#e8f5e9;border-left:4px solid #4caf50;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}
.warning-box{background:#fff3e0;border-left:4px solid #ff9800;padding:15px 20px;
  margin:15px 0;border-radius:0 4px 4px 0}
nav{background:#fff;padding:20px 30px;border-radius:8px;margin-bottom:30px;
  box-shadow:0 2px 4px rgba(0,0,0,.1)}
nav h3{margin-bottom:10px;color:#1a237e}
nav ol{padding-left:20px} nav li{margin-bottom:4px}
nav a{color:#1565c0;text-decoration:none}
footer{text-align:center;padding:20px;color:#999;font-size:.85em}
"""

def _build_lateral_html(charts, tables, map_files, carriage, sec_num,
                        critical_html="", exclude_dirs=None):
    """Build HTML for a lateral carriage section."""
    cg = CARRIAGE_GROUPS[carriage]
    slug = carriage.lower().replace(". ", "_").replace(" ", "_")
    dirs = [d for d in cg["directions"]
            if not (exclude_dirs and d in exclude_dirs)]

    # Tables — filter excluded directions
    ov = tables[f"overview_{slug}"]
    exc = tables[f"exceedance_{slug}"]
    if exclude_dirs:
        for ex in exclude_dirs:
            lbl = cg["dir_labels"][ex]
            ov = ov[ov["Tratta"] != lbl]
            exc = exc[exc["Tratta"] != lbl]

    # Temporal per tratta
    temporal = charts.get(f"{carriage}_by_tratta", {})
    temporal_html = ""
    for d in dirs:
        tc = temporal.get(d, {})
        if tc:
            temporal_html += (
                f"<h4>{d} &mdash; Velocit&agrave; Media per Ora</h4>"
                f'<img src="data:image/png;base64,{tc["speed_by_hour"]}" style="max-width:100%;">'
                f"<h4>{d} &mdash; V85 per Ora</h4>"
                f'<img src="data:image/png;base64,{tc["v85_by_hour"]}" style="max-width:100%;">')

    # V85 spatial
    cdata = charts.get(carriage, {})
    v85_sp = cdata.get("v85_spatial", {})
    v85_html = ""
    for d in dirs:
        lbl = cg["dir_labels"][d]
        if lbl in v85_sp:
            v85_html += f'<h4>{lbl}</h4><img src="data:image/png;base64,{v85_sp[lbl]}" style="max-width:100%;">'

    night_html = ""
    if "night" in cdata:
        night_html = f'<img src="data:image/png;base64,{cdata["night"]}" style="max-width:100%;">'

    map_html = ""
    for d in dirs:
        d_slug = d.lower().replace(" ", "_").replace("-", "_")
        for pre, lab in [("v85_24h", "V85 (24h)"), ("v85_night", "V85 Notturno")]:
            fn = f"{pre}_{slug}_{d_slug}.html"
            if fn in map_files:
                map_html += _ifr(fn, f"{lab} — {d}")

    s = sec_num
    return f"""
<section id="carriage_{slug}">
<h2 style="color:#fff;background:#1a237e;padding:15px 20px;border-radius:6px;">
{cg['label']}</h2>

<h3>{s}.1 Panoramica</h3>
{_df_html(ov)}

<h3>{s}.2 Superamento Limiti (pesato per km)</h3>
{_df_html(exc)}

<h3>{s}.3 Profili Temporali</h3>
{temporal_html}

<h3>{s}.4 Profili Spaziali V85</h3>
{v85_html}
{critical_html}

<h3>{s}.5 Analisi Notturna</h3>
{night_html}

<h3>{s}.6 Mappe V85</h3>
{map_html}
</section>"""


def _build_tratte_html(charts, tratte_df, intersection_info, map_files):
    """Build HTML for tratte section."""
    profile_img = charts.get("tratte_profile", "")
    tables_html = ""
    for direction in ["Ostia", "Centro"]:
        td = tratte_df[tratte_df["direction"] == direction].sort_values("tratta")
        n_i = len(intersection_info.get(direction, set()))
        tables_html += f"<h4>Dir. {direction}</h4>\n"
        tables_html += ('<table class="data-table">'
                        '<tr><th>Tratta</th><th>Progressiva</th><th>Lunghezza</th>'
                        '<th>V85 medio</th><th>V85 min</th><th>V85 max</th><th>CV</th></tr>\n')
        for _, tr in td.iterrows():
            tables_html += (
                f'<tr><td><strong>T{tr["tratta"]}</strong></td>'
                f'<td>{tr["cum_dist_start"]:.0f}&ndash;{tr["cum_dist_end"]:.0f} m</td>'
                f'<td>{tr["length_km"]:.2f} km</td>'
                f'<td><strong>{tr["v85_mean"]:.0f}</strong> km/h</td>'
                f'<td>{tr["v85_min"]:.0f} km/h</td>'
                f'<td>{tr["v85_max"]:.0f} km/h</td>'
                f'<td>{tr["v85_cv"]:.1f}%</td></tr>\n')
        tables_html += "</table>\n"
        if n_i:
            tables_html += f"<p><em>{n_i} segmenti identificati come intersezioni semaforiche ed esclusi dal calcolo.</em></p>\n"

    tratte_maps = ""
    for direction in ["Ostia", "Centro"]:
        fn = f"tratte_v85_centrale_{direction.lower()}.html"
        if fn in map_files:
            tratte_maps += _ifr(fn, f"Tratte V85 — Dir. {direction}")

    return f"""
<section id="tratte_analysis">
<h2>5. Segmentazione Carreggiata Centrale in Tratte Omogenee V85</h2>

<p>La Carreggiata Centrale &egrave; stata suddivisa in <strong>tratte omogenee</strong>
sulla base del profilo spaziale del V85 (85&deg; percentile delle velocit&agrave;).
I segmenti corrispondenti a <strong>intersezioni semaforiche</strong> sono stati
identificati (V85 medio &lt; 80% della mediana del corridoio) ed esclusi dal
calcolo della segmentazione, in quanto non rappresentativi del comportamento
di guida in marcia libera.</p>

<div class="method-box">
<strong>Metodologia nel dettaglio:</strong>
<p>La segmentazione utilizza l&rsquo;algoritmo <strong>Binary Segmentation</strong>
(BinSeg), un metodo di <em>change-point detection</em>:</p>
<ol>
<li><strong>Segnale multivariato:</strong> per ogni segmento (escluse le intersezioni)
    si costruisce un vettore di 24 valori, corrispondenti al V85 medio per ciascuna
    ora del giorno. Questo cattura non solo il livello medio di velocit&agrave;
    ma anche il <em>profilo temporale</em> (ore di punta vs notte).</li>
<li><strong>Ricerca dei punti di cambio:</strong> l&rsquo;algoritmo divide ricorsivamente
    la sequenza cercando il punto che minimizza la somma dei costi (modello L2 = varianza)
    delle due partizioni risultanti. Ad ogni passo il taglio migliore viene selezionato
    e il processo si ripete sulle sotto-sequenze.</li>
<li><strong>Vincolo di dimensione minima:</strong> ogni tratta contiene almeno
    2&ndash;3 segmenti, evitando tratte troppo corte o artefatti.</li>
<li><strong>Numero di tratte:</strong> pu&ograve; essere inferiore a 9, a seconda
    della struttura dei dati. Il numero bilancia omogeneit&agrave; interna e
    significativit&agrave; delle differenze tra tratte adiacenti.</li>
</ol>
<p><strong>Coefficiente di Variazione (CV):</strong> misura l&rsquo;omogeneit&agrave;
interna. Valori &lt;5% = tratte molto omogenee; 5&ndash;10% = accettabile;
&gt;10% = variabilit&agrave; intrinseca (rampe, svincoli).</p>
<p><strong>Ri-assegnazione intersezioni:</strong> dopo la segmentazione, i segmenti
di intersezione vengono ri-assegnati alla tratta pi&ugrave; vicina, ma le loro
velocit&agrave; <em>non</em> influenzano le statistiche della tratta.</p>
</div>

<h3>5.1 Profilo Spaziale V85 con Tratte</h3>
<img src="data:image/png;base64,{profile_img}" style="max-width:100%;">

<h3>5.2 Tabella Riepilogativa Tratte</h3>
{tables_html}

<h3>5.3 Mappe Interattive Tratte</h3>
{tratte_maps}
</section>"""


def build_html_ridotto(charts, tables, map_files, tratte_df, intersection_info):
    """Full reduced HTML report."""
    cdata = charts["Centrale"]
    ov_html = _df_html(tables["overview_centrale"])
    exc_html = _df_html(tables["exceedance_centrale"])
    v85_sp = "".join(
        f'<h4>{lab}</h4><img src="data:image/png;base64,{b64}" style="max-width:100%;">'
        for lab, b64 in cdata["v85_spatial"].items())
    map_c = ""
    for d in ["Ostia", "Centro"]:
        for pre, lab in [("v85_24h", "V85 (24h)"), ("v85_night", "V85 Notturno")]:
            fn = f"{pre}_centrale_{d.lower()}.html"
            if fn in map_files:
                map_c += _ifr(fn, f"{lab} — Dir. {d}")

    sec3 = _build_lateral_html(charts, tables, map_files, "Lat. Ostia", 3,
        critical_html="""
<div class="warning-box">
<strong>Aree critiche &mdash; Pontina-Mezzocammino:</strong>
tra <strong>Via Carmelo Maestrini</strong> e <strong>Via Ercole Dei</strong> e in
prossimit&agrave; del <strong>GRA</strong>, con V85 massime tra <strong>107 e 113 km/h</strong>.
</div>
<div class="warning-box">
<strong>Aree critiche &mdash; Via di Acilia - Piazzale Colombo:</strong>
tra <strong>Via di Acilia</strong> e <strong>Via Pindaro</strong>, e tra
<strong>Via Demostene</strong> e <strong>Via Pericle</strong>, con V85 massime
tra <strong>110 e 130 km/h</strong>.
</div>""")

    sec4 = _build_lateral_html(charts, tables, map_files, "Lat. Centro", 4,
        exclude_dirs={"Svincolo Malafede"},
        critical_html="""
<div class="warning-box">
<strong>Aree critiche &mdash; Piazzale-Acilia:</strong>
V85 elevati nelle zone di transizione, con valori massimi significativi
nelle fasce orarie notturne.
</div>
<div class="warning-box">
<strong>Aree critiche &mdash; Alfonsine-Brandeliero:</strong>
velocit&agrave; elevate lungo il tratto rettilineo, con V85 che superano il limite
nelle ore notturne.
</div>""")

    sec5 = _build_tratte_html(charts, tratte_df, intersection_info, map_files)

    return f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Report Ridotto &mdash; Via Cristoforo Colombo</title>
<style>{CSS}</style>
</head>
<body>
<header>
<h1>Analisi della Distribuzione delle Velocit&agrave;</h1>
<h2>Via Cristoforo Colombo &mdash; Roma</h2>
<p>Report Ridotto &bull; Dati TomTom Speed Profiles</p>
</header>
<nav><h3>Indice</h3>
<ol>
<li><a href="#intro">Introduzione e Metodologia</a></li>
<li><a href="#carriage_centrale">Carreggiata Centrale</a></li>
<li><a href="#carriage_lat_ostia">Carreggiata Laterale Dir. Ostia</a></li>
<li><a href="#carriage_lat_centro">Carreggiata Laterale Dir. Centro</a></li>
<li><a href="#tratte_analysis">Segmentazione Centrale in Tratte V85</a></li>
</ol></nav>

<section id="intro">
<h2>1. Introduzione e Metodologia</h2>
<p>Il presente report analizza la distribuzione delle velocit&agrave; veicolari lungo
<strong>Via Cristoforo Colombo</strong> a Roma, utilizzando i dati TomTom Speed Profiles.</p>
<p>L&rsquo;arteria &egrave; composta da tre carreggiate distinte:</p>
<ul>
<li><strong>Carreggiata Centrale</strong> (Dir. Ostia e Dir. Centro) &mdash; ~16 km</li>
<li><strong>Laterale Dir. Ostia</strong> (Pontina-Mezzocammino + Via di Acilia-Piazzale)
    &mdash; ~13 km</li>
<li><strong>Laterale Dir. Centro</strong> (Piazzale-Acilia + Alfonsine-Brandeliero)
    &mdash; ~11 km</li>
</ul>
<p>I dati coprono <strong>4 periodi</strong>:</p>
<ul>
<li><strong>Invernale Feriale</strong>: 1&ndash;15 febbraio 2026 (lun&ndash;ven)</li>
<li><strong>Estiva Feriale</strong>: 15&ndash;31 luglio 2025 (lun&ndash;ven)</li>
<li><strong>Estiva Festivo</strong>: 15&ndash;31 luglio 2025 (sab&ndash;dom)</li>
<li><strong>Invernale Festiva</strong>: 1&ndash;15 febbraio 2026 (sab&ndash;dom)</li>
</ul>
<div class="method-box">
<strong>Metriche principali:</strong>
<ul>
<li><strong>Velocit&agrave; media armonica</strong> &mdash; media ponderata per lunghezza segmento</li>
<li><strong>V85 (85&deg; percentile)</strong> &mdash; velocit&agrave; non superata dall&rsquo;85% dei veicoli</li>
<li><strong>Deviazione standard</strong> &mdash; dispersione delle velocit&agrave;</li>
<li><strong>Tasso di superamento</strong> &mdash; % dei km con velocit&agrave; superiore al limite
    specifico del segmento</li>
</ul>
</div>
</section>

<section id="carriage_centrale">
<h2 style="color:#fff;background:#1a237e;padding:15px 20px;border-radius:6px;">
Carreggiata Centrale</h2>

<h3>2.1 Panoramica</h3>
{ov_html}

<h3>2.2 Superamento Limiti (pesato per km)</h3>
{exc_html}

<h3>2.3 Profili Temporali</h3>
<h4>Velocit&agrave; Media per Ora</h4>
<img src="data:image/png;base64,{cdata['speed_by_hour']}" style="max-width:100%;">
<h4>V85 per Ora</h4>
<img src="data:image/png;base64,{cdata['v85_by_hour']}" style="max-width:100%;">

<h3>2.4 Profili Spaziali V85</h3>
{v85_sp}

<div class="warning-box">
<strong>Aree critiche &mdash; Dir. Ostia:</strong>
<ul>
<li>Tra l&rsquo;altezza di <strong>Via Giovanni Soncelli</strong> e <strong>Via Carmelo Maestrini</strong></li>
<li>In prossimit&agrave; di <strong>Viale Don Pasquino Borghi</strong></li>
</ul>
</div>
<div class="warning-box">
<strong>Aree critiche &mdash; Dir. Centro:</strong>
<ul>
<li>In prossimit&agrave; di <strong>Viale Don Pasquino Borghi</strong></li>
<li>Ad altezza del <strong>GRA</strong></li>
<li>Ad altezza di <strong>Via Carmelo Maestrini</strong></li>
<li>Ad altezza di <strong>Via Armando Brasini</strong></li>
<li>In corrispondenza di <strong>Via del Risaro</strong> e <strong>Via del Canale della Lingua</strong></li>
</ul>
<p>In entrambe le direzioni si raggiungono valori prossimi ai <strong>130 km/h</strong>
nella fascia oraria pi&ugrave; critica.</p>
</div>

<h3>2.5 Analisi Notturna</h3>
<img src="data:image/png;base64,{cdata['night']}" style="max-width:100%;">

<h3>2.6 Mappe V85</h3>
{map_c}
</section>

{sec3}
{sec4}
{sec5}

<footer><p>TomTom Speed Profiles &mdash; Via Cristoforo Colombo &mdash; Report Ridotto</p></footer>
</body></html>"""


# ================================================================
# BLOCK H — WORD REPORT ASSEMBLY
# ================================================================

def _add_lateral_to_doc(doc, carriage, sec, charts, tables,
                        static_paths, img_w, exclude_dirs=None, crit_notes=None):
    cg = CARRIAGE_GROUPS[carriage]
    slug = carriage.lower().replace(". ", "_").replace(" ", "_")
    dirs = [d for d in cg["directions"]
            if not (exclude_dirs and d in exclude_dirs)]

    doc.add_heading(f"{sec}. {cg['label']}", level=1)
    doc.add_heading(f"{sec}.1 Panoramica", level=2)
    ov = tables[f"overview_{slug}"]
    if exclude_dirs:
        for ex in exclude_dirs:
            ov = ov[ov["Tratta"] != cg["dir_labels"][ex]]
    _add_table_to_doc(doc, ov)

    doc.add_heading(f"{sec}.2 Superamento Limiti", level=2)
    exc = tables[f"exceedance_{slug}"]
    if exclude_dirs:
        for ex in exclude_dirs:
            exc = exc[exc["Tratta"] != cg["dir_labels"][ex]]
    _add_table_to_doc(doc, exc)

    doc.add_heading(f"{sec}.3 Profili Temporali", level=2)
    temporal = charts.get(f"{carriage}_by_tratta", {})
    for d in dirs:
        tc = temporal.get(d, {})
        if tc:
            doc.add_heading(d, level=3)
            doc.add_picture(_b64_to_stream(tc["speed_by_hour"]), width=img_w)
            _add_caption(doc, f"Velocità media per ora — {d}")
            doc.add_picture(_b64_to_stream(tc["v85_by_hour"]), width=img_w)
            _add_caption(doc, f"V85 per ora — {d}")

    cdata = charts.get(carriage, {})
    v85_sp = cdata.get("v85_spatial", {})
    doc.add_heading(f"{sec}.4 Profili Spaziali V85", level=2)
    for d in dirs:
        lbl = cg["dir_labels"][d]
        if lbl in v85_sp:
            doc.add_picture(_b64_to_stream(v85_sp[lbl]), width=img_w)
            _add_caption(doc, f"Profilo V85 — {lbl}")

    if crit_notes:
        doc.add_heading("Aree critiche", level=3)
        for n in crit_notes:
            doc.add_paragraph(n)

    if "night" in cdata:
        doc.add_heading(f"{sec}.5 Analisi Notturna", level=2)
        doc.add_picture(_b64_to_stream(cdata["night"]), width=img_w)

    for sp in static_paths:
        if slug in sp and ("v85_24h" in sp or "v85_night" in sp):
            for d in dirs:
                ds = d.lower().replace(" ", "_").replace("-", "_")
                if ds in sp:
                    doc.add_picture(sp, width=img_w)
                    _add_caption(doc, Path(sp).stem.replace("_", " ").title())
                    break
    doc.add_page_break()


def build_word_ridotto(charts, tables, static_paths, tratte_df, inter_info):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    img_w = Inches(6.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # Title
    doc.add_paragraph(); doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Analisi della Distribuzione delle Velocità")
    r.bold = True; r.font.size = Pt(24); r.font.color.rgb = RGBColor(0x1A,0x23,0x7E)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Via Cristoforo Colombo — Roma"); r.bold = True
    r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x28,0x35,0x93)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Report Ridotto — Dati TomTom Speed Profiles")
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_page_break()

    # TOC
    doc.add_heading("Indice", level=1)
    for it in ["1. Introduzione", "2. Carreggiata Centrale",
               "3. Laterale Dir. Ostia", "4. Laterale Dir. Centro",
               "5. Segmentazione Tratte V85"]:
        doc.add_paragraph(it, style="List Number")
    doc.add_page_break()

    # 1. Introduzione (no speed limit references)
    doc.add_heading("1. Introduzione e Metodologia", level=1)
    doc.add_paragraph(
        "Il presente report analizza la distribuzione delle velocità veicolari lungo "
        "Via Cristoforo Colombo a Roma, utilizzando i dati TomTom Speed Profiles.")
    doc.add_paragraph("L'arteria è composta da tre carreggiate distinte:")
    for it in ["Carreggiata Centrale (Dir. Ostia e Dir. Centro) — ~16 km",
               "Laterale Dir. Ostia (2 tratti) — ~13 km",
               "Laterale Dir. Centro (2 tratti) — ~11 km"]:
        doc.add_paragraph(it, style="List Bullet")
    doc.add_paragraph("I dati coprono 4 periodi:")
    for it in ["Invernale Feriale: 1–15 febbraio 2026 (lun–ven)",
               "Estiva Feriale: 15–31 luglio 2025 (lun–ven)",
               "Estiva Festivo: 15–31 luglio 2025 (sab–dom)",
               "Invernale Festiva: 1–15 febbraio 2026 (sab–dom)"]:
        doc.add_paragraph(it, style="List Bullet")
    _add_method_box(doc,
        "Metriche: velocità media armonica, V85 (85° percentile), deviazione standard, "
        "tasso di superamento (% km oltre il limite specifico del segmento).")
    doc.add_page_break()

    # 2. Centrale
    cdata = charts["Centrale"]
    doc.add_heading("2. Carreggiata Centrale", level=1)
    doc.add_heading("2.1 Panoramica", level=2)
    _add_table_to_doc(doc, tables["overview_centrale"])
    doc.add_heading("2.2 Superamento Limiti", level=2)
    _add_table_to_doc(doc, tables["exceedance_centrale"])
    doc.add_heading("2.3 Profili Temporali", level=2)
    doc.add_picture(_b64_to_stream(cdata["speed_by_hour"]), width=img_w)
    _add_caption(doc, "Velocità media per ora (4 periodi)")
    doc.add_picture(_b64_to_stream(cdata["v85_by_hour"]), width=img_w)
    _add_caption(doc, "V85 per ora (4 periodi)")
    doc.add_heading("2.4 Profili Spaziali V85", level=2)
    for lab, b64 in cdata["v85_spatial"].items():
        doc.add_picture(_b64_to_stream(b64), width=img_w)
        _add_caption(doc, f"Profilo V85 — {lab}")
    doc.add_heading("Aree critiche", level=3)
    doc.add_paragraph(
        "Dir. Ostia: tra l'altezza di Via Giovanni Soncelli e Via Carmelo Maestrini; "
        "in prossimità di Viale Don Pasquino Borghi.")
    doc.add_paragraph(
        "Dir. Centro: prossimità Viale Don Pasquino Borghi; altezza GRA; "
        "altezza Via Carmelo Maestrini; altezza Via Armando Brasini; "
        "corrispondenza Via del Risaro e Via del Canale della Lingua.")
    doc.add_paragraph(
        "In entrambe le direzioni si raggiungono valori prossimi ai 130 km/h "
        "nella fascia oraria più critica.")
    doc.add_heading("2.5 Analisi Notturna", level=2)
    doc.add_picture(_b64_to_stream(cdata["night"]), width=img_w)
    for sp in static_paths:
        if "centrale" in sp and ("v85_24h" in sp or "v85_night" in sp):
            doc.add_picture(sp, width=img_w)
            _add_caption(doc, Path(sp).stem.replace("_", " ").title())
    doc.add_page_break()

    # 3. Laterale Ostia
    _add_lateral_to_doc(doc, "Lat. Ostia", 3, charts, tables, static_paths, img_w,
        crit_notes=[
            "Tratto Pontina-Mezzocammino: punti critici tra Via Carmelo Maestrini e "
            "Via Ercole Dei e in prossimità del GRA, con V85 massime tra 107 e 113 km/h.",
            "Tratto Via di Acilia - Piazzale Colombo: punti critici tra Via di Acilia e "
            "Via Pindaro, e tra Via Demostene e Via Pericle, con V85 massime tra 110 e 130 km/h.",
        ])

    # 4. Laterale Centro (no Svincolo Malafede)
    _add_lateral_to_doc(doc, "Lat. Centro", 4, charts, tables, static_paths, img_w,
        exclude_dirs={"Svincolo Malafede"},
        crit_notes=[
            "Tratto Piazzale-Acilia: V85 elevati nelle zone di transizione.",
            "Tratto Alfonsine-Brandeliero: velocità elevate lungo il tratto rettilineo, "
            "con V85 che superano il limite nelle ore notturne.",
        ])

    # 5. Tratte
    doc.add_heading("5. Segmentazione Centrale in Tratte Omogenee V85", level=1)
    doc.add_paragraph(
        "La Carreggiata Centrale è stata suddivisa in tratte omogenee sulla base "
        "del profilo spaziale del V85. I segmenti di intersezione semaforica sono "
        "stati esclusi dal calcolo della segmentazione.")
    _add_method_box(doc,
        "Metodologia: segmentazione multivariata (profilo V85 a 24 ore) con "
        "algoritmo Binary Segmentation (modello L2). Per ogni segmento non di "
        "intersezione si costruisce un vettore di 24 valori V85 (uno per ora). "
        "L'algoritmo divide ricorsivamente la sequenza cercando i punti di cambio "
        "che minimizzano la varianza interna. Ogni tratta contiene almeno 2-3 segmenti. "
        "Il numero di tratte può essere inferiore a 9. "
        "CV: <5% molto omogenea, 5-10% accettabile, >10% variabilità intrinseca. "
        "Le intersezioni sono ri-assegnate alla tratta più vicina senza influenzare le statistiche.")
    profile_b64 = charts.get("tratte_profile", "")
    if profile_b64:
        doc.add_picture(_b64_to_stream(profile_b64), width=img_w)
        _add_caption(doc, "Profilo V85 con tratte omogenee (intersezioni escluse)")
    for direction in ["Ostia", "Centro"]:
        td = tratte_df[tratte_df["direction"] == direction].sort_values("tratta")
        n_i = len(inter_info.get(direction, set()))
        doc.add_heading(f"Dir. {direction}", level=3)
        tbl = pd.DataFrame({
            "Tratta": [f"T{r['tratta']}" for _, r in td.iterrows()],
            "Progressiva": [f"{r['cum_dist_start']:.0f}–{r['cum_dist_end']:.0f} m"
                            for _, r in td.iterrows()],
            "Lungh.": [f"{r['length_km']:.2f} km" for _, r in td.iterrows()],
            "V85 medio": [f"{r['v85_mean']:.0f}" for _, r in td.iterrows()],
            "V85 min": [f"{r['v85_min']:.0f}" for _, r in td.iterrows()],
            "V85 max": [f"{r['v85_max']:.0f}" for _, r in td.iterrows()],
            "CV": [f"{r['v85_cv']:.1f}%" for _, r in td.iterrows()],
        })
        _add_table_to_doc(doc, tbl)
        if n_i:
            doc.add_paragraph(f"{n_i} segmenti di intersezione esclusi dal calcolo.")

    out = RIDOTTO_DIR / "report_ridotto_colombo.docx"
    doc.save(str(out))
    return str(out)


# ================================================================
# BLOCK I — MAIN
# ================================================================

def main():
    print("=" * 60)
    print("REPORT RIDOTTO — Via Cristoforo Colombo")
    print("=" * 60)

    print("\n1. Caricamento dati...")
    segments, summaries, headers = load_all_data()
    print(f"   {len(segments)} righe, {segments['seg_idx'].nunique()} segmenti")

    print("2. Statistiche segmento...")
    seg_stats = segment_peak_stats(segments)

    RIDOTTO_DIR.mkdir(parents=True, exist_ok=True)
    RIDOTTO_MAPS.mkdir(parents=True, exist_ok=True)
    RIDOTTO_STATIC.mkdir(parents=True, exist_ok=True)

    print("3. Tratte (senza intersezioni)...")
    tratte_df, tratte_seg_map, inter_info = compute_tratte_no_intersections(segments)
    for d in ["Ostia", "Centro"]:
        td = tratte_df[tratte_df["direction"] == d]
        ni = len(inter_info.get(d, set()))
        print(f"   Dir. {d}: {len(td)} tratte, {ni} intersezioni escluse")

    print("4. Grafici (senza linee limite)...")
    charts = {}
    for carriage in CARRIAGE_GROUPS:
        charts[carriage] = {
            "speed_by_hour": chart_speed_by_hour_nolim(summaries, carriage),
            "v85_by_hour": chart_v85_by_hour_nolim(segments, carriage),
            "v85_spatial": chart_v85_spatial_nolim(segments, carriage),
            "night": chart_night_nolim(segments, carriage),
        }
    for carriage in ["Lat. Ostia", "Lat. Centro"]:
        charts[f"{carriage}_by_tratta"] = chart_temporal_by_tratta(segments, carriage)
    charts["tratte_profile"] = chart_tratte_profile_ridotto(
        segments, tratte_df, inter_info)
    print("   Grafici generati.")

    print("5. Tabelle...")
    tables = build_tables_ridotto(segments, summaries)

    print("6. Mappe statiche...")
    static_paths = generate_static_maps_ridotto(segments)
    print(f"   {len(static_paths)} mappe statiche")

    print("7. Mappe interattive...")
    map_files = generate_maps_ridotto(segments)
    for direction in ["Ostia", "Centro"]:
        m = make_folium_tratte_map(segments, tratte_df, tratte_seg_map, direction)
        if m:
            fn = f"tratte_v85_centrale_{direction.lower()}.html"
            m.save(str(RIDOTTO_MAPS / fn))
            map_files[fn] = f"Tratte V85 — Dir. {direction}"
    print(f"   {len(map_files)} mappe interattive")

    print("8. Report HTML...")
    html = build_html_ridotto(charts, tables, map_files, tratte_df, inter_info)
    html_path = RIDOTTO_DIR / "report_ridotto_colombo.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write(html)
    print(f"   {html_path}")

    print("9. Report Word...")
    docx_path = build_word_ridotto(charts, tables, static_paths,
                                   tratte_df, inter_info)
    print(f"   {docx_path}")

    print(f"\n{'=' * 60}")
    print(f"HTML: {html_path}")
    print(f"DOCX: {docx_path}")
    print(f"Mappe: {RIDOTTO_MAPS}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
